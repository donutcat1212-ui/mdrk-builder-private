from __future__ import annotations

import os
import queue
import re
import sys
import threading
import tkinter as tk
import traceback
from collections.abc import Callable
from datetime import datetime
from pathlib import Path
from tempfile import TemporaryDirectory, gettempdir
from tkinter import filedialog, messagebox, scrolledtext, ttk

from docx import Document

from mdrk_builder.application.scanner import scan_patient_folder
from mdrk_builder.application.validation import (
    ACKNOWLEDGEABLE_CONFLICT_CODES,
    acknowledge_conflict,
    can_generate,
    clear_conflict_acknowledgements,
    current_issues,
    is_conflict_acknowledged,
)
from mdrk_builder.domain import (
    Episode,
    MdrkKind,
    ReviewIssue,
    ReviewSeverity,
    ScaleMeasurement,
    SourceDocument,
    SpecialistFinding,
    SpecialistRole,
)
from mdrk_builder.infrastructure.docx_writer import (
    SignatoryRow,
    canonical_template_path,
    write_mdrk_docx,
)
from mdrk_builder.ui.dialogs import (
    FindingDialog,
    IcfDomainDialog,
    ProcedureDialog,
    ScaleDialog,
    install_edit_shortcuts,
)
from mdrk_builder.ui.episode_adapter import (
    EpisodeFormData,
    apply_episode_form_data,
    episode_signatory_roles,
    format_date,
    format_datetime,
    parse_episode_folder,
    parse_episode_form_data,
    parse_optional_datetime,
    parse_optional_meeting_datetime,
    sections_for,
)


SEVERITY_LABELS = {
    ReviewSeverity.BLOCKING: "БЛОКИРУЕТ",
    ReviewSeverity.WARNING: "ПРЕДУПРЕЖДЕНИЕ",
    ReviewSeverity.INFO: "ИНФО",
}

MEETING_RESCAN_MESSAGE = (
    "Время заседания изменено. Нажмите «Сканировать», чтобы "
    "заново собрать данные на этот момент. Ручные правки будут заменены "
    "только после вашего подтверждения."
)

CONFLICT_FIELD_LABELS = {
    "identity_conflict_medical_record_number": "Номер ИБ",
    "mixed_hospitalizations_admission_date": "Дата и время поступления",
}


def _normalized_record_number(value: str) -> str:
    normalized = "".join(
        character
        for character in value.casefold().replace("№", "")
        if character.isalnum() or character == "/"
    )
    while normalized.startswith("скп"):
        normalized = normalized.removeprefix("скп")
    return normalized


class MdrkBuilderApp:
    def __init__(self, root: tk.Tk) -> None:
        self.root = root
        self.episode: Episode | None = None
        self._current_kind = MdrkKind.INITIAL
        self._scan_results: queue.Queue[tuple[Episode | None, Exception | None]] = queue.Queue()
        self._scan_thread: threading.Thread | None = None
        self._scan_folder: Path | None = None
        self._scanning = False
        self._setting_folder_field = False
        self._last_form_error = ""
        self._entry_variables: dict[str, tk.StringVar] = {}
        self._text_fields: dict[str, tk.Text] = {}
        self._scale_refs: list[tuple[int, int]] = []
        self._issue_refs: dict[str, ReviewIssue] = {}

        self.folder_var = tk.StringVar()
        self.kind_var = tk.StringVar(value=MdrkKind.INITIAL.value)
        self.status_var = tk.StringVar(value="Выберите папку эпизода")

        self._configure_window()
        install_edit_shortcuts(self.root)
        self._build_menu()
        self._build_layout()
        self.folder_var.trace_add("write", self._on_folder_field_changed)
        self._update_action_states()
        self.root.protocol("WM_DELETE_WINDOW", self._on_close)

    def _configure_window(self) -> None:
        self.root.title("МДРК — сборщик документов")
        self.root.geometry("1180x790")
        self.root.minsize(980, 660)
        style = ttk.Style(self.root)
        preferred = "winnative" if sys.platform == "win32" else "clam"
        if preferred in style.theme_names():
            style.theme_use(preferred)
        style.configure("Status.TLabel", relief="sunken", anchor="w", padding=(5, 2))

    def _build_menu(self) -> None:
        menu = tk.Menu(self.root)
        file_menu = tk.Menu(menu, tearoff=False)
        file_menu.add_command(label="Выбрать папку…", command=self._choose_folder, accelerator="Ctrl+O")
        file_menu.add_command(label="Сканировать", command=self._start_scan, accelerator="F5")
        file_menu.add_separator()
        file_menu.add_command(label="Создать DOCX…", command=self._generate, accelerator="Ctrl+S")
        file_menu.add_separator()
        file_menu.add_command(label="Выход", command=self._on_close)
        menu.add_cascade(label="Файл", menu=file_menu)

        help_menu = tk.Menu(menu, tearoff=False)
        help_menu.add_command(label="О программе", command=self._show_about)
        menu.add_cascade(label="Справка", menu=help_menu)
        self.root.config(menu=menu)
        self.root.bind("<Control-o>", lambda _event: self._choose_folder())
        self.root.bind("<Control-s>", lambda _event: self._generate())
        self.root.bind("<F5>", lambda _event: self._start_scan())

    def _build_layout(self) -> None:
        top = ttk.Frame(self.root, padding=6)
        top.pack(fill="x")
        ttk.Label(top, text="Папка эпизода:").grid(row=0, column=0, sticky="w")
        folder_entry = ttk.Entry(top, textvariable=self.folder_var)
        folder_entry.grid(row=0, column=1, sticky="ew", padx=(6, 4))
        ttk.Button(top, text="Обзор…", command=self._choose_folder).grid(row=0, column=2, padx=2)
        self.scan_button = ttk.Button(top, text="Сканировать", command=self._start_scan)
        self.scan_button.grid(row=0, column=3, padx=2)
        self.generate_button = ttk.Button(
            top,
            text="Создать DOCX…",
            command=self._generate,
            state="disabled",
        )
        self.generate_button.grid(row=0, column=4, padx=(8, 2))
        top.columnconfigure(1, weight=1)

        snapshot = ttk.LabelFrame(self.root, text="Снимок", padding=(8, 4))
        snapshot.pack(fill="x", padx=6, pady=(0, 5))
        ttk.Radiobutton(
            snapshot,
            text="МДРК-1 — исходный",
            value=MdrkKind.INITIAL.value,
            variable=self.kind_var,
            command=self._on_kind_changed,
        ).pack(side="left", padx=(0, 18))
        ttk.Radiobutton(
            snapshot,
            text="МДРК-2 — итоговый",
            value=MdrkKind.FINAL.value,
            variable=self.kind_var,
            command=self._on_kind_changed,
        ).pack(side="left")

        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill="both", expand=True, padx=6, pady=(0, 5))
        self._build_main_tab()
        self._build_sources_tab()
        self._build_icf_tab()
        self._build_scales_tab()
        self._build_procedures_tab()
        self._build_findings_tab()
        self._build_issues_tab()

        bottom = ttk.Frame(self.root)
        bottom.pack(fill="x", side="bottom")
        self.progress = ttk.Progressbar(bottom, mode="indeterminate", length=170)
        self.progress.pack(side="right", padx=(5, 6), pady=2)
        ttk.Label(bottom, textvariable=self.status_var, style="Status.TLabel").pack(
            side="left", fill="x", expand=True
        )

    def _build_main_tab(self) -> None:
        tab = ttk.Frame(self.notebook, padding=7)
        self.notebook.add(tab, text="Основное")
        metadata = ttk.LabelFrame(tab, text="Шапка документа", padding=7)
        metadata.pack(fill="x")
        fields = (
            ("full_name", "ФИО пациента"),
            ("record_number", "Номер ИБ"),
            ("birth_date", "Дата рождения"),
            ("sex", "Пол"),
            ("admission", "Поступление"),
            ("meeting", "Заседание"),
            ("department", "Отделение"),
            ("stage", "Этап реабилитации"),
            ("duration", "Койко-дни"),
        )
        for index, (key, label) in enumerate(fields):
            row, group = divmod(index, 2)
            column = group * 2
            ttk.Label(metadata, text=label).grid(row=row, column=column, sticky="w", padx=(0, 5), pady=3)
            variable = tk.StringVar()
            self._entry_variables[key] = variable
            ttk.Entry(metadata, textvariable=variable).grid(
                row=row, column=column + 1, sticky="ew", padx=(0, 12), pady=3
            )
        metadata.columnconfigure(1, weight=1)
        metadata.columnconfigure(3, weight=1)

        sections = ttk.Notebook(tab)
        sections.pack(fill="both", expand=True, pady=(7, 0))
        self._add_text_group(
            sections,
            "Клиника",
            (("clinical_diagnosis", "Клинический диагноз", 7), ("disease_history", "Анамнез заболевания", 7), ("life_history", "Анамнез жизни", 5)),
        )
        self._add_text_group(
            sections,
            "Исследования",
            (("laboratory_results", "Лабораторные исследования", 9), ("instrumental_results", "Инструментальные исследования", 9)),
        )
        self._add_text_group(
            sections,
            "План",
            (("rehabilitation_potential", "Реабилитационный потенциал", 2), ("limitations", "Ограничивающие факторы", 3), ("risks", "Факторы риска", 3), ("goal", "Цель", 4), ("tasks", "Задачи", 5)),
        )
        self._add_text_group(
            sections,
            "Режим и лечение",
            (("movement_regimen", "Двигательный режим", 2), ("diet", "Диета", 2), ("medication", "Медикаментозное лечение", 14)),
        )

    def _add_text_group(
        self, notebook: ttk.Notebook, title: str, fields: tuple[tuple[str, str, int], ...]
    ) -> None:
        frame = ttk.Frame(notebook, padding=6)
        notebook.add(frame, text=title)
        for row, (key, label, height) in enumerate(fields):
            ttk.Label(frame, text=label).grid(row=row * 2, column=0, sticky="w", pady=(3, 1))
            widget = scrolledtext.ScrolledText(frame, height=height, wrap="word", undo=True)
            widget.grid(row=row * 2 + 1, column=0, sticky="nsew", pady=(0, 4))
            frame.rowconfigure(row * 2 + 1, weight=max(1, height))
            self._text_fields[key] = widget
        frame.columnconfigure(0, weight=1)

    def _build_icf_tab(self) -> None:
        tab = ttk.Frame(self.notebook, padding=7)
        self.notebook.add(tab, text="МКФ")
        columns = ("code", "description", "role", "initial", "final", "dynamic", "note")
        self.icf_tree = self._create_tree_with_scrollbars(tab, columns, selectmode="extended")
        headings = {
            "code": "Код",
            "description": "Описание",
            "role": "Специалист",
            "initial": "Исх.",
            "final": "Повт.",
            "dynamic": "+/-",
            "note": "Уточнение",
        }
        widths = {"code": 80, "description": 330, "role": 210, "initial": 55, "final": 55, "dynamic": 45, "note": 170}
        for column in columns:
            self.icf_tree.heading(column, text=headings[column])
            self.icf_tree.column(column, width=widths[column], minwidth=40, anchor="w")
        self.icf_tree.bind("<Double-1>", lambda _event: self._edit_icf())
        self._bind_tree_delete(self.icf_tree, self._delete_icf)
        buttons = ttk.Frame(tab)
        buttons.pack(fill="x", pady=(6, 0))
        ttk.Button(buttons, text="Добавить…", command=self._add_icf).pack(side="left")
        ttk.Button(buttons, text="Изменить…", command=self._edit_icf).pack(side="left", padx=4)
        ttk.Button(buttons, text="Удалить", command=self._delete_icf).pack(side="left")

    def _build_sources_tab(self) -> None:
        tab = ttk.Frame(self.notebook, padding=7)
        self.notebook.add(tab, text="Источники")
        columns = ("role", "clinical_datetime", "document_type", "path")
        self.source_tree = self._create_tree_with_scrollbars(tab, columns)
        for column, heading, width in (
            ("role", "Роль", 250),
            ("clinical_datetime", "Клиническая дата", 155),
            ("document_type", "Тип документа", 190),
            ("path", "Путь", 610),
        ):
            self.source_tree.heading(column, text=heading)
            self.source_tree.column(column, width=width, minwidth=65, anchor="w")
        self.source_tree.tag_configure("used", background="#e3f3df")
        self.source_tree.tag_configure("excluded", background="#e4e4e4", foreground="#666666")
        ttk.Label(
            tab,
            text=(
                "Зелёным отмечены исходники, из которых взяты поля, заключения, шкалы, МКФ "
                "или процедуры. Серым — документы с другим номером ИБ; они видны, но не участвуют в сборке."
            ),
        ).pack(anchor="w", pady=(6, 0))

    def _build_procedures_tab(self) -> None:
        tab = ttk.Frame(self.notebook, padding=7)
        self.notebook.add(tab, text="Процедуры")
        columns = ("code", "name", "specialist", "count", "duration", "frequency")
        self.procedure_tree = self._create_tree_with_scrollbars(
            tab, columns, selectmode="extended"
        )
        headings = {
            "code": "Код",
            "name": "Процедура",
            "specialist": "Ответственный",
            "count": "Кол-во",
            "duration": "Мин.",
            "frequency": "Кратность",
        }
        widths = {"code": 120, "name": 420, "specialist": 220, "count": 70, "duration": 65, "frequency": 90}
        for column in columns:
            self.procedure_tree.heading(column, text=headings[column])
            self.procedure_tree.column(column, width=widths[column], minwidth=45, anchor="w")
        self.procedure_tree.bind("<Double-1>", lambda _event: self._edit_procedure())
        self._bind_tree_delete(self.procedure_tree, self._delete_procedure)
        buttons = ttk.Frame(tab)
        buttons.pack(fill="x", pady=(6, 0))
        ttk.Button(buttons, text="Добавить…", command=self._add_procedure).pack(side="left")
        ttk.Button(buttons, text="Изменить…", command=self._edit_procedure).pack(side="left", padx=4)
        ttk.Button(buttons, text="Удалить", command=self._delete_procedure).pack(side="left")

    def _build_scales_tab(self) -> None:
        tab = ttk.Frame(self.notebook, padding=7)
        self.notebook.add(tab, text="Шкалы")
        columns = ("role", "date", "name", "value", "source")
        self.scale_tree = self._create_tree_with_scrollbars(tab, columns, selectmode="extended")
        for column, heading, width in (
            ("role", "Специалист", 245),
            ("date", "Дата и время", 145),
            ("name", "Шкала/опросник", 410),
            ("value", "Результат", 210),
            ("source", "Источник", 260),
        ):
            self.scale_tree.heading(column, text=heading)
            self.scale_tree.column(column, width=width, minwidth=55, anchor="w")
        self.scale_tree.bind("<Double-1>", lambda _event: self._edit_scale())
        self._bind_tree_delete(self.scale_tree, self._delete_scale)
        buttons = ttk.Frame(tab)
        buttons.pack(fill="x", pady=(6, 0))
        ttk.Button(buttons, text="Добавить…", command=self._add_scale).pack(side="left")
        ttk.Button(buttons, text="Изменить…", command=self._edit_scale).pack(side="left", padx=4)
        ttk.Button(buttons, text="Удалить", command=self._delete_scale).pack(side="left")

    def _build_findings_tab(self) -> None:
        tab = ttk.Frame(self.notebook, padding=7)
        self.notebook.add(tab, text="Заключения")
        columns = ("role", "date", "scales", "conclusion")
        self.finding_tree = self._create_tree_with_scrollbars(
            tab, columns, selectmode="extended"
        )
        for column, heading, width in (
            ("role", "Специалист", 250),
            ("date", "Клиническая дата", 145),
            ("scales", "Шкал", 55),
            ("conclusion", "Заключение", 620),
        ):
            self.finding_tree.heading(column, text=heading)
            self.finding_tree.column(column, width=width, minwidth=45, anchor="w")
        self.finding_tree.bind("<Double-1>", lambda _event: self._edit_finding())
        self._bind_tree_delete(self.finding_tree, self._delete_finding)
        buttons = ttk.Frame(tab)
        buttons.pack(fill="x", pady=(6, 0))
        ttk.Button(buttons, text="Добавить…", command=self._add_finding).pack(side="left")
        ttk.Button(buttons, text="Изменить…", command=self._edit_finding).pack(side="left", padx=4)
        ttk.Button(buttons, text="Удалить", command=self._delete_finding).pack(side="left")

    def _build_issues_tab(self) -> None:
        tab = ttk.Frame(self.notebook, padding=7)
        self.issues_tab = tab
        self.notebook.add(tab, text="Предупреждения")
        columns = ("severity", "message", "field", "source")
        self.issue_tree = self._create_tree_with_scrollbars(tab, columns)
        for column, heading, width in (
            ("severity", "Уровень", 130),
            ("message", "Сообщение", 580),
            ("field", "Поле", 170),
            ("source", "Источник", 310),
        ):
            self.issue_tree.heading(column, text=heading)
            self.issue_tree.column(column, width=width, minwidth=60, anchor="w")
        self.issue_tree.tag_configure("blocking", background="#ffd6d6")
        self.issue_tree.tag_configure("warning", background="#fff4c2")
        self.issue_tree.tag_configure("info", background="#e6f1ff")
        buttons = ttk.Frame(tab)
        buttons.pack(fill="x", pady=(6, 0))
        ttk.Button(
            buttons,
            text="Подтвердить выбранный конфликт…",
            command=self._acknowledge_selected_conflict,
        ).pack(side="left")
        ttk.Button(
            buttons,
            text="Сбросить подтверждения",
            command=self._reset_conflict_acknowledgements,
        ).pack(side="left", padx=4)
        ttk.Button(
            buttons,
            text="Обновить после правок",
            command=self._refresh_issues,
        ).pack(side="left")

    @staticmethod
    def _create_tree_with_scrollbars(
        parent: ttk.Frame,
        columns: tuple[str, ...],
        *,
        selectmode: str = "browse",
    ) -> ttk.Treeview:
        container = ttk.Frame(parent)
        container.pack(fill="both", expand=True)
        tree = ttk.Treeview(
            container,
            columns=columns,
            show="headings",
            selectmode=selectmode,
        )
        vertical = ttk.Scrollbar(container, orient="vertical", command=tree.yview)
        horizontal = ttk.Scrollbar(container, orient="horizontal", command=tree.xview)
        tree.configure(yscrollcommand=vertical.set, xscrollcommand=horizontal.set)
        tree.grid(row=0, column=0, sticky="nsew")
        vertical.grid(row=0, column=1, sticky="ns")
        horizontal.grid(row=1, column=0, sticky="ew")
        container.rowconfigure(0, weight=1)
        container.columnconfigure(0, weight=1)
        return tree

    @staticmethod
    def _bind_tree_delete(tree: ttk.Treeview, command: Callable[[], None]) -> None:
        def delete_selected(_event: tk.Event) -> str:
            command()
            return "break"

        tree.bind("<Delete>", delete_selected)
        tree.bind("<KP_Delete>", delete_selected)

    def _set_folder_field(self, value: str) -> None:
        self._setting_folder_field = True
        try:
            self.folder_var.set(value)
        finally:
            self._setting_folder_field = False

    def _on_folder_field_changed(self, *_args: str) -> None:
        if self._setting_folder_field:
            return
        if self.episode is not None:
            self._invalidate_episode()
        self._update_action_states()
        self.status_var.set("Папка изменена. Выполните сканирование заново.")

    def _folder_field_matches(self, expected: Path) -> bool:
        try:
            return parse_episode_folder(self.folder_var.get()) == expected.resolve()
        except (OSError, ValueError):
            return False

    def _invalidate_episode(self) -> None:
        self.episode = None
        self._last_form_error = ""
        for variable in self._entry_variables.values():
            variable.set("")
        for widget in self._text_fields.values():
            widget.delete("1.0", "end")
        for tree in (
            self.source_tree,
            self.icf_tree,
            self.scale_tree,
            self.procedure_tree,
            self.finding_tree,
            self.issue_tree,
        ):
            self._clear_tree(tree)
        self._scale_refs = []
        self._issue_refs = {}
        self._update_action_states()

    def _choose_folder(self) -> None:
        selected = filedialog.askdirectory(title="Выберите папку эпизода")
        if selected:
            self._invalidate_episode()
            self._set_folder_field(selected)
            self.status_var.set("Папка выбрана. Нажмите «Сканировать».")

    def _start_scan(self) -> None:
        if self._scanning:
            return
        try:
            folder = parse_episode_folder(self.folder_var.get())
        except (OSError, ValueError) as exc:
            messagebox.showerror("Папка не найдена", str(exc))
            return
        try:
            meeting_variable = self._entry_variables.get("meeting")
            entered_meeting = parse_optional_meeting_datetime(
                meeting_variable.get() if meeting_variable is not None else ""
            )
            record_variable = self._entry_variables.get("record_number")
            entered_record_number = (
                record_variable.get().strip() if record_variable is not None else ""
            )
            admission_variable = self._entry_variables.get("admission")
            entered_admission = parse_optional_datetime(
                admission_variable.get() if admission_variable is not None else ""
            )
        except ValueError as exc:
            messagebox.showerror("Проверьте даты и реквизиты", str(exc))
            return
        if not folder.is_dir():
            messagebox.showerror("Папка не найдена", "Выберите существующую папку эпизода.")
            return
        if self.episode and not messagebox.askyesno(
            "Повторное сканирование",
            "Текущие ручные правки будут заменены результатами нового сканирования. Продолжить?",
        ):
            return
        scan_overrides: dict[str, datetime | str] = {}
        if self.episode is not None:
            record_number_changed = (
                bool(entered_record_number)
                and _normalized_record_number(entered_record_number)
                != _normalized_record_number(
                    self.episode.materialized_medical_record_number
                    or self.episode.identity.medical_record_number
                )
            )
            admission_changed = (
                entered_admission is not None
                and entered_admission
                != (
                    self.episode.materialized_admission_datetime
                    or self.episode.admission_datetime
                )
            )
            metadata_changed = record_number_changed or admission_changed
            if entered_record_number:
                scan_overrides["medical_record_number_override"] = entered_record_number
            if entered_admission is not None and (
                not record_number_changed or admission_changed
            ):
                scan_overrides["admission_datetime_override"] = entered_admission

            if metadata_changed:
                current_meeting = self.episode.meeting_at(self._current_kind)
                if entered_meeting is not None and entered_meeting != current_meeting:
                    override_name = (
                        "initial_meeting_at"
                        if self._current_kind is MdrkKind.INITIAL
                        else "final_meeting_at"
                    )
                    scan_overrides[override_name] = entered_meeting
            else:
                initial_meeting = self.episode.initial_meeting_at
                final_meeting = self.episode.final_meeting_at
                if self._current_kind is MdrkKind.INITIAL:
                    initial_meeting = entered_meeting
                else:
                    final_meeting = entered_meeting
                if initial_meeting is not None:
                    scan_overrides["initial_meeting_at"] = initial_meeting
                if final_meeting is not None:
                    scan_overrides["final_meeting_at"] = final_meeting
        elif entered_meeting is not None:
            override_name = (
                "initial_meeting_at"
                if self._current_kind is MdrkKind.INITIAL
                else "final_meeting_at"
            )
            scan_overrides[override_name] = entered_meeting
        self._invalidate_episode()
        self._set_folder_field(str(folder))
        self._scan_folder = folder
        self._set_scanning(True)
        self.status_var.set("Сканирование исходных документов…")

        def worker() -> None:
            try:
                self._scan_results.put(
                    (scan_patient_folder(folder, **scan_overrides), None)
                )
            except Exception as exc:  # delivered to the UI thread
                self._scan_results.put((None, exc))

        self._scan_thread = threading.Thread(target=worker, name="mdrk-folder-scan", daemon=False)
        self._scan_thread.start()
        self.root.after(100, self._poll_scan_result)

    def _poll_scan_result(self) -> None:
        try:
            episode, error = self._scan_results.get_nowait()
        except queue.Empty:
            if self._scanning:
                self.root.after(100, self._poll_scan_result)
            return
        self._set_scanning(False)
        scan_folder = self._scan_folder
        self._scan_folder = None
        if error is not None:
            self._invalidate_episode()
            self.status_var.set("Сканирование завершилось ошибкой")
            messagebox.showerror("Ошибка сканирования", str(error))
            return
        if (
            episode is None
            or scan_folder is None
            or not self._folder_field_matches(scan_folder)
            or episode.folder.resolve() != scan_folder
        ):
            self._invalidate_episode()
            self.status_var.set("Результат отброшен: папка изменилась во время сканирования.")
            return
        self.episode = episode
        self._set_folder_field(str(episode.folder))
        self._populate_from_episode()
        self._update_action_states()
        self.status_var.set(
            f"Готово: {len(episode.sources)} источников, {len(episode.icf_domains)} доменов, "
            f"{len(episode.procedures)} процедур"
        )

    def _set_scanning(self, value: bool) -> None:
        self._scanning = value
        self._update_action_states()
        if value:
            self.progress.start(12)
        else:
            self.progress.stop()

    def _update_action_states(self) -> None:
        self.scan_button.configure(state="disabled" if self._scanning else "normal")
        can_use_episode = self.episode is not None and not self._scanning
        self.generate_button.configure(state="normal" if can_use_episode else "disabled")

    def _populate_from_episode(self) -> None:
        if not self.episode:
            return
        episode = self.episode
        values = {
            "full_name": episode.identity.full_name,
            "record_number": episode.identity.medical_record_number,
            "birth_date": format_date(episode.identity.birth_date),
            "sex": episode.identity.sex,
            "admission": format_datetime(episode.admission_datetime),
            "meeting": format_datetime(episode.meeting_at(self._current_kind)),
            "department": episode.department,
            "stage": episode.stage,
            "duration": "" if episode.course_duration_days is None else str(episode.course_duration_days),
        }
        for key, value in values.items():
            self._entry_variables[key].set(value)
        current_sections = sections_for(episode, self._current_kind)
        for key, widget in self._text_fields.items():
            widget.delete("1.0", "end")
            widget.insert("1.0", getattr(current_sections, key))
        self._refresh_all_trees()

    def _parsed_form_data(self) -> EpisodeFormData:
        entry_values = {
            key: variable.get() for key, variable in self._entry_variables.items()
        }
        section_values = {
            key: widget.get("1.0", "end-1c")
            for key, widget in self._text_fields.items()
        }
        return parse_episode_form_data(entry_values, section_values)

    def _apply_form(self, kind: MdrkKind | None = None) -> bool:
        if not self.episode:
            messagebox.showwarning("Нет данных", "Сначала просканируйте папку эпизода.")
            return False
        target_kind = kind or self._current_kind
        try:
            form = self._parsed_form_data()
        except ValueError as exc:
            self._last_form_error = str(exc)
            self._render_form_error()
            messagebox.showerror("Проверьте поля", str(exc))
            return False
        if form.meeting_at != self.episode.meeting_at(target_kind):
            self._last_form_error = MEETING_RESCAN_MESSAGE
            self.status_var.set("Время заседания изменено: нужно повторное сканирование.")
            messagebox.showerror("Нужно повторное сканирование", MEETING_RESCAN_MESSAGE)
            return False
        apply_episode_form_data(self.episode, target_kind, form)
        self._last_form_error = ""
        return True

    def _selected_kind(self) -> MdrkKind:
        return MdrkKind(self.kind_var.get())

    def _on_kind_changed(self) -> None:
        requested_kind = self._selected_kind()
        if self.episode:
            try:
                form = self._parsed_form_data()
            except ValueError as exc:
                self.kind_var.set(self._current_kind.value)
                self._last_form_error = str(exc)
                self.status_var.set("Снимок не переключён: исправьте поля.")
                messagebox.showerror("Снимок не переключён", str(exc))
                return
            if form.meeting_at != self.episode.meeting_at(self._current_kind):
                self.kind_var.set(self._current_kind.value)
                self._last_form_error = MEETING_RESCAN_MESSAGE
                self.status_var.set(
                    "Снимок не переключён: нужно повторное сканирование."
                )
                messagebox.showerror(
                    "Нужно повторное сканирование",
                    MEETING_RESCAN_MESSAGE,
                )
                return
            apply_episode_form_data(self.episode, self._current_kind, form)
        self._current_kind = requested_kind
        if self.episode:
            self._populate_from_episode()

    def _generate(self) -> None:
        if self._scanning or not self.episode:
            return
        if not self._folder_field_matches(self.episode.folder):
            self._invalidate_episode()
            self.status_var.set("Папка не совпадает со сканированным эпизодом.")
            messagebox.showerror(
                "Нужно повторное сканирование",
                "Папка в поле не совпадает с папкой загруженного эпизода.",
            )
            return
        if not self._apply_form(self._current_kind):
            return
        kind = self._current_kind
        self._refresh_issues()
        if not can_generate(self.episode, kind):
            self.notebook.select(self.issues_tab)
            messagebox.showerror(
                "Документ не создан",
                "Остались блокирующие проблемы. Исправьте обязательные поля на вкладке "
                "«Предупреждения».",
            )
            return
        warnings = [
            issue for issue in current_issues(self.episode, kind) if issue.severity is ReviewSeverity.WARNING
        ]
        if warnings and not messagebox.askyesno(
            "Требуется проверка",
            f"Осталось предупреждений: {len(warnings)}. Создать редактируемый DOCX с этими данными?",
        ):
            return
        default_name = self._default_output_name(kind)
        output = filedialog.asksaveasfilename(
            title="Сохранить МДРК",
            defaultextension=".docx",
            filetypes=(("Документ Word", "*.docx"),),
            initialfile=default_name,
        )
        if not output:
            return
        signatories = [
            SignatoryRow(role=role.display_name)
            for role in episode_signatory_roles(self.episode, kind)
        ]
        try:
            created = write_mdrk_docx(self.episode, kind, Path(output), signatories=signatories)
        except Exception as exc:
            messagebox.showerror("Не удалось создать DOCX", str(exc))
            self.status_var.set("Ошибка генерации DOCX")
            return
        self.status_var.set(f"DOCX создан: {created}")
        messagebox.showinfo("Готово", f"Документ создан:\n{created}\n\nПроверьте его перед использованием.")

    def _default_output_name(self, kind: MdrkKind) -> str:
        patient = self.episode.identity.full_name if self.episode else "пациент"
        safe_patient = re.sub(r"[^0-9A-Za-zА-Яа-яЁё_-]+", "_", patient).strip("_") or "пациент"
        number = "1" if kind is MdrkKind.INITIAL else "2"
        return f"МДРК_{number}_{safe_patient}.docx"

    def _refresh_all_trees(self) -> None:
        self._refresh_sources()
        self._refresh_icf()
        self._refresh_scales()
        self._refresh_procedures()
        self._refresh_findings()
        self._refresh_issues()

    @staticmethod
    def _clear_tree(tree: ttk.Treeview) -> None:
        tree.delete(*tree.get_children())

    def _refresh_icf(self) -> None:
        self._clear_tree(self.icf_tree)
        if not self.episode:
            return
        for index, domain in enumerate(self.episode.icf_domains):
            self.icf_tree.insert(
                "",
                "end",
                iid=str(index),
                values=(
                    domain.code,
                    domain.description,
                    domain.specialist.display_name,
                    domain.initial.display() if domain.initial else "",
                    domain.final.display() if domain.final else "",
                    domain.dynamic_marker if domain.dynamic_marker is not None else "?",
                    domain.note,
                ),
            )

    def _refresh_sources(self) -> None:
        self._clear_tree(self.source_tree)
        if not self.episode:
            return
        used_paths = set(self.episode.initial_field_sources.values())
        used_paths.update(self.episode.field_sources.values())
        used_paths.update(finding.source for finding in self.episode.findings if finding.source)
        used_paths.update(
            measurement.source
            for finding in self.episode.findings
            for measurement in finding.scales
            if measurement.source
        )
        used_paths.update(domain.initial_source for domain in self.episode.icf_domains if domain.initial_source)
        used_paths.update(domain.final_source for domain in self.episode.icf_domains if domain.final_source)
        used_paths.update(procedure.source for procedure in self.episode.procedures if procedure.source)
        for index, source in enumerate(self.episode.sources):
            tags = (
                ("excluded",)
                if not self.episode.source_is_active(source)
                else (("used",) if source.path in used_paths else ())
            )
            self.source_tree.insert(
                "",
                "end",
                iid=str(index),
                tags=tags,
                values=(
                    source.role.display_name,
                    format_datetime(source.clinical_datetime),
                    source.document_type,
                    str(source.path),
                ),
            )

    def _refresh_procedures(self) -> None:
        self._clear_tree(self.procedure_tree)
        if not self.episode:
            return
        for index, procedure in enumerate(self.episode.procedures):
            self.procedure_tree.insert(
                "",
                "end",
                iid=str(index),
                values=(
                    procedure.code,
                    procedure.name,
                    procedure.specialist,
                    "" if procedure.actual_count is None else procedure.actual_count,
                    "" if procedure.duration_minutes is None else procedure.duration_minutes,
                    procedure.frequency,
                ),
            )

    def _refresh_scales(self) -> None:
        self._clear_tree(self.scale_tree)
        self._scale_refs = []
        if not self.episode:
            return
        for finding_index, finding in enumerate(self.episode.findings):
            for scale_index, measurement in enumerate(finding.scales):
                row_index = len(self._scale_refs)
                self._scale_refs.append((finding_index, scale_index))
                self.scale_tree.insert(
                    "",
                    "end",
                    iid=str(row_index),
                    values=(
                        measurement.specialist.display_name,
                        format_datetime(measurement.measured_at),
                        measurement.name,
                        measurement.value,
                        str(measurement.source) if measurement.source else "",
                    ),
                )

    def _refresh_findings(self) -> None:
        self._clear_tree(self.finding_tree)
        if not self.episode:
            return
        for index, finding in enumerate(self.episode.findings):
            conclusion = " ".join(finding.conclusion.split())
            self.finding_tree.insert(
                "",
                "end",
                iid=str(index),
                values=(
                    finding.role.display_name,
                    format_datetime(finding.source_datetime),
                    len(finding.scales),
                    conclusion,
                ),
            )

    def _refresh_issues(self) -> None:
        if not self.episode:
            self._clear_tree(self.issue_tree)
            self._issue_refs = {}
            return
        if not self._apply_form_without_messages():
            self._render_form_error()
            return
        self._clear_tree(self.issue_tree)
        self._issue_refs = {}
        for index, issue in enumerate(current_issues(self.episode, self._current_kind)):
            row_id = str(index)
            self._issue_refs[row_id] = issue
            self.issue_tree.insert(
                "",
                "end",
                iid=row_id,
                tags=(issue.severity.value,),
                values=(
                    SEVERITY_LABELS[issue.severity],
                    issue.message,
                    issue.field,
                    str(issue.source) if issue.source else "",
                ),
            )

    def _selected_issue(self) -> ReviewIssue | None:
        selected = self.issue_tree.selection()
        if not selected:
            return None
        return self._issue_refs.get(selected[0])

    def _acknowledge_selected_conflict(self) -> None:
        if not self.episode:
            messagebox.showwarning("Нет данных", "Сначала просканируйте папку эпизода.")
            return
        issue = self._selected_issue()
        if issue is None:
            messagebox.showwarning(
                "Конфликт не выбран",
                "Выберите блокирующую строку с номером ИБ или датой поступления.",
            )
            return
        if issue.code not in ACKNOWLEDGEABLE_CONFLICT_CODES:
            messagebox.showwarning(
                "Нельзя подтвердить",
                "Эта проблема требует исправления и не может быть проигнорирована.",
            )
            return
        if not self._apply_form(self._current_kind):
            return
        if is_conflict_acknowledged(self.episode, issue.code):
            messagebox.showinfo(
                "Уже подтверждено",
                "Текущее значение уже подтверждено вручную.",
            )
            return

        label = CONFLICT_FIELD_LABELS[issue.code]
        value = (
            self.episode.identity.medical_record_number.strip()
            if issue.code == "identity_conflict_medical_record_number"
            else format_datetime(self.episode.admission_datetime)
        )
        confirmed = messagebox.askyesno(
            "Подтвердить ручное значение",
            f"{issue.message}\n\n"
            f"Использовать проверенное значение из формы:\n{label}: {value}\n\n"
            "Исходный конфликт останется в предупреждениях. "
            "Остальные блокирующие проверки продолжат действовать.",
        )
        if not confirmed:
            return
        try:
            acknowledge_conflict(self.episode, issue.code)
        except ValueError as exc:
            messagebox.showerror("Не удалось подтвердить", str(exc))
            return
        self._refresh_issues()
        self.status_var.set(f"Подтверждено вручную: {label} = {value}")

    def _reset_conflict_acknowledgements(self) -> None:
        if not self.episode or not self.episode.acknowledged_conflicts:
            messagebox.showinfo("Подтверждений нет", "Ручных подтверждений конфликтов нет.")
            return
        if not messagebox.askyesno(
            "Сбросить подтверждения",
            "Вернуть блокировку для вручную подтверждённых конфликтов?",
        ):
            return
        clear_conflict_acknowledgements(self.episode)
        self._refresh_issues()
        self.status_var.set("Ручные подтверждения конфликтов сброшены.")

    def _render_form_error(self) -> None:
        self._clear_tree(self.issue_tree)
        self._issue_refs = {}
        self.issue_tree.insert(
            "",
            "end",
            iid="form_error",
            tags=(ReviewSeverity.BLOCKING.value,),
            values=(
                SEVERITY_LABELS[ReviewSeverity.BLOCKING],
                self._last_form_error or "Неверные данные в форме",
                "ui.form",
                "",
            ),
        )
        self.status_var.set("Проверьте поля: обнаружена блокирующая ошибка.")

    def _apply_form_without_messages(self) -> bool:
        if not self.episode:
            return False
        try:
            form = self._parsed_form_data()
        except ValueError as exc:
            self._last_form_error = str(exc)
            return False
        if form.meeting_at != self.episode.meeting_at(self._current_kind):
            self._last_form_error = MEETING_RESCAN_MESSAGE
            return False
        apply_episode_form_data(self.episode, self._current_kind, form)
        self._last_form_error = ""
        return True

    @staticmethod
    def _selected_index(tree: ttk.Treeview) -> int | None:
        selection = tree.selection()
        return int(selection[0]) if selection else None

    @staticmethod
    def _selected_indices(tree: ttk.Treeview) -> list[int]:
        return sorted(int(item_id) for item_id in tree.selection())

    def _add_icf(self) -> None:
        if not self.episode:
            return
        dialog = IcfDomainDialog(self.root)
        if dialog.result:
            self.episode.icf_domains.append(dialog.result)
            self._refresh_icf()
            self._refresh_issues()

    def _edit_icf(self) -> None:
        if not self.episode or (index := self._selected_index(self.icf_tree)) is None:
            return
        dialog = IcfDomainDialog(self.root, self.episode.icf_domains[index])
        if dialog.result:
            self.episode.icf_domains[index] = dialog.result
            self._refresh_icf()
            self._refresh_issues()

    def _delete_icf(self) -> None:
        if not self.episode or not (indices := self._selected_indices(self.icf_tree)):
            return
        noun = (
            "выбранный домен МКФ"
            if len(indices) == 1
            else f"выбранные домены МКФ ({len(indices)})"
        )
        if messagebox.askyesno("Удалить домен", f"Удалить {noun}?"):
            for index in reversed(indices):
                self.episode.icf_domains.pop(index)
            self._refresh_icf()
            self._refresh_issues()

    def _add_procedure(self) -> None:
        if not self.episode:
            return
        dialog = ProcedureDialog(self.root)
        if dialog.result:
            self.episode.procedures.append(dialog.result)
            self._refresh_procedures()
            self._refresh_issues()

    def _add_scale(self) -> None:
        if not self.episode:
            return
        dialog = ScaleDialog(self.root)
        if dialog.result:
            self._append_scale(dialog.result)
            self._refresh_scales()
            self._refresh_findings()
            self._refresh_issues()

    def _edit_scale(self) -> None:
        if not self.episode or (row_index := self._selected_index(self.scale_tree)) is None:
            return
        finding_index, scale_index = self._scale_refs[row_index]
        measurement = self.episode.findings[finding_index].scales[scale_index]
        dialog = ScaleDialog(self.root, measurement)
        if not dialog.result:
            return
        if dialog.result.specialist is self.episode.findings[finding_index].role:
            self.episode.findings[finding_index].scales[scale_index] = dialog.result
        else:
            self.episode.findings[finding_index].scales.pop(scale_index)
            self._append_scale(dialog.result)
        self._refresh_scales()
        self._refresh_findings()
        self._refresh_issues()

    def _delete_scale(self) -> None:
        if not self.episode or not (row_indices := self._selected_indices(self.scale_tree)):
            return
        noun = (
            "выбранное измерение шкалы"
            if len(row_indices) == 1
            else f"выбранные измерения шкал ({len(row_indices)})"
        )
        if messagebox.askyesno("Удалить измерение", f"Удалить {noun}?"):
            references = sorted(
                (self._scale_refs[row_index] for row_index in row_indices),
                reverse=True,
            )
            for finding_index, scale_index in references:
                self.episode.findings[finding_index].scales.pop(scale_index)
            self._refresh_scales()
            self._refresh_findings()
            self._refresh_issues()

    def _append_scale(self, measurement: ScaleMeasurement) -> None:
        if not self.episode:
            return
        finding = next(
            (item for item in reversed(self.episode.findings) if item.role is measurement.specialist),
            None,
        )
        if finding is None:
            finding = SpecialistFinding(role=measurement.specialist)
            self.episode.findings.append(finding)
        finding.scales.append(measurement)

    def _edit_procedure(self) -> None:
        if not self.episode or (index := self._selected_index(self.procedure_tree)) is None:
            return
        dialog = ProcedureDialog(self.root, self.episode.procedures[index])
        if dialog.result:
            self.episode.procedures[index] = dialog.result
            self._refresh_procedures()
            self._refresh_issues()

    def _delete_procedure(self) -> None:
        if not self.episode or not (indices := self._selected_indices(self.procedure_tree)):
            return
        noun = (
            "выбранную процедуру"
            if len(indices) == 1
            else f"выбранные процедуры ({len(indices)})"
        )
        if messagebox.askyesno("Удалить процедуру", f"Удалить {noun}?"):
            for index in reversed(indices):
                self.episode.procedures.pop(index)
            self._refresh_procedures()
            self._refresh_issues()

    def _add_finding(self) -> None:
        if not self.episode:
            return
        dialog = FindingDialog(self.root)
        if dialog.result:
            self.episode.findings.append(dialog.result)
            self._refresh_findings()
            self._refresh_scales()
            self._refresh_issues()

    def _edit_finding(self) -> None:
        if not self.episode or (index := self._selected_index(self.finding_tree)) is None:
            return
        dialog = FindingDialog(self.root, self.episode.findings[index])
        if dialog.result:
            for measurement in dialog.result.scales:
                measurement.specialist = dialog.result.role
            self.episode.findings[index] = dialog.result
            self._refresh_findings()
            self._refresh_scales()
            self._refresh_issues()

    def _delete_finding(self) -> None:
        if not self.episode or not (indices := self._selected_indices(self.finding_tree)):
            return
        scale_count = sum(len(self.episode.findings[index].scales) for index in indices)
        question = (
            "Удалить выбранное заключение специалиста?"
            if len(indices) == 1
            else f"Удалить выбранные заключения специалистов ({len(indices)})?"
        )
        if scale_count:
            question += (
                f"\n\nВместе с ним будут удалены измерения шкал: {scale_count}."
            )
        if messagebox.askyesno("Удалить заключение", question):
            for index in reversed(indices):
                self.episode.findings.pop(index)
            self._refresh_findings()
            self._refresh_scales()
            self._refresh_issues()

    def _show_about(self) -> None:
        messagebox.showinfo(
            "О программе",
            "МДРК Builder 0.1.4\n\nЛокальная подготовка редактируемых МДРК-1 и МДРК-2.\n"
            "Программа не отправляет документы в интернет и не заменяет проверку специалистом.",
        )

    def _on_close(self) -> None:
        if self._scanning:
            messagebox.showwarning(
                "Сканирование выполняется",
                "Дождитесь завершения сканирования: сейчас закрытие может оставить Microsoft Word открытым.",
            )
            return
        self.root.destroy()


def _generate_smoke_document(directory: Path) -> Path:
    _write_smoke_report("phase=template")
    template = canonical_template_path()
    if not template.is_file():
        raise FileNotFoundError(f"Канонический шаблон не найден: {template}")

    episode = Episode(folder=directory)
    episode.identity.full_name = "Тестов Тест Тестович"
    episode.identity.medical_record_number = "SMOKE-1"
    episode.admission_datetime = datetime(2026, 1, 1, 9, 0)
    episode.initial_meeting_at = datetime(2026, 1, 2, 8, 0)
    episode.initial_sections.clinical_diagnosis = "Тестовый диагноз"
    episode.sections.clinical_diagnosis = "Тестовый диагноз"
    episode.sources.append(
        SourceDocument(directory / "smoke-source.docx", role=SpecialistRole.NEUROLOGIST)
    )
    _write_smoke_report("phase=write_docx")
    output = write_mdrk_docx(
        episode,
        MdrkKind.INITIAL,
        directory / "smoke-output.docx",
    )
    if not output.is_file() or output.stat().st_size == 0:
        raise RuntimeError("Тестовый DOCX не был создан")
    _write_smoke_report("phase=reopen_docx")
    reopened = Document(output)
    if not reopened.paragraphs:
        raise RuntimeError("Тестовый DOCX не содержит ожидаемых абзацев")
    return output


def smoke_test(*, include_ui: bool = False) -> int:
    with TemporaryDirectory(prefix="mdrk-builder-smoke-") as temporary:
        _generate_smoke_document(Path(temporary))

    if include_ui:
        _write_smoke_report("phase=tk_init")
        root = tk.Tk()
        try:
            root.withdraw()
            MdrkBuilderApp(root)
            _write_smoke_report("phase=app_constructed")
            _assert_consistent_geometry_managers(root)
            root.update_idletasks()
            root.update()
            _write_smoke_report("phase=idle_updated")
        finally:
            root.destroy()
            _write_smoke_report("phase=ui_destroyed")
    return 0


def _assert_consistent_geometry_managers(widget: tk.Misc) -> None:
    if widget.pack_slaves() and widget.grid_slaves():
        raise RuntimeError(
            f"В одном контейнере смешаны pack и grid: {widget.winfo_pathname(widget.winfo_id())}"
        )
    for child in widget.winfo_children():
        _assert_consistent_geometry_managers(child)


def _write_smoke_report(message: str, *, reset: bool = False) -> None:
    report_path = os.environ.get("MDRK_BUILDER_SMOKE_REPORT")
    if not report_path:
        return
    mode = "w" if reset else "a"
    try:
        with Path(report_path).open(mode, encoding="utf-8") as report:
            report.write(f"{message}\n")
    except OSError:
        return


def _run_smoke(*, include_ui: bool) -> int:
    _write_smoke_report("phase=start", reset=True)
    try:
        result = smoke_test(include_ui=include_ui)
        _write_smoke_report("status=ok")
        return result
    except Exception:
        try:
            _write_smoke_report(traceback.format_exc())
        except OSError:
            pass
        return 1


def _write_crash_report(error_text: str) -> Path | None:
    base = Path(os.environ.get("LOCALAPPDATA", gettempdir()))
    report = base / "MDRK Builder" / "logs" / "startup-error.log"
    try:
        report.parent.mkdir(parents=True, exist_ok=True)
        report.write_text(error_text, encoding="utf-8")
    except OSError:
        return None
    return report


def _run_gui() -> int:
    root: tk.Tk | None = None
    try:
        root = tk.Tk()
        MdrkBuilderApp(root)
        root.mainloop()
        return 0
    except Exception:
        report = _write_crash_report(traceback.format_exc())
        message = "Программа не смогла запуститься."
        if report is not None:
            message += f"\n\nТехнический отчёт сохранён:\n{report}"
        if root is not None:
            try:
                messagebox.showerror("Ошибка запуска МДРК Builder", message, parent=root)
            except Exception:
                pass
            try:
                root.destroy()
            except Exception:
                pass
        return 1


def main(argv: list[str] | None = None) -> int:
    arguments = list(sys.argv[1:] if argv is None else argv)
    if "--smoke-test-ui" in arguments:
        return _run_smoke(include_ui=True)
    if "--smoke-test" in arguments:
        return _run_smoke(include_ui=False)
    return _run_gui()


if __name__ == "__main__":
    raise SystemExit(main())
