from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Twips
from docx.styles.style import BaseStyle

from .docx_layout import (
    BOTTOM_MARGIN_DXA,
    FOOTER_DISTANCE_DXA,
    HEADER_DISTANCE_DXA,
    LEFT_MARGIN_DXA,
    PAGE_HEIGHT_DXA,
    PAGE_WIDTH_DXA,
    RIGHT_MARGIN_DXA,
    TOP_MARGIN_DXA,
)
from .docx_output import resolve_docx_output_path, save_sanitized_docx_atomically


TEMPLATE_FILENAME = "canonical_mdrk_template.docx"
FONT_NAME = "Times New Roman"
FONT_SIZE_PT = 12
TABLE_FONT_SIZE_PT = 10

STYLE_BODY = "MDRK Body"
STYLE_TITLE = "MDRK Title"
STYLE_MEETING = "MDRK Meeting"
STYLE_SECTION = "MDRK Section"
STYLE_LABEL = "MDRK Label"
STYLE_TABLE = "MDRK Table"
STYLE_TABLE_HEADER = "MDRK Table Header"
STYLE_MCF_CODE = "MDRK MCF Code"
STYLE_TASK = "MDRK Task"
STYLE_WARNING = "MDRK Warning"


def canonical_template_path() -> Path:
    return Path(__file__).resolve().parent.parent / "resources" / TEMPLATE_FILENAME


def create_canonical_template(output_path: Path) -> Path:
    output_path = resolve_docx_output_path(output_path)

    document = Document()
    _configure_page(document)
    configure_mdrk_styles(document)
    _configure_metadata(document)

    document.add_paragraph("Канонический шаблон МДРК", style=STYLE_TITLE)
    document.add_paragraph(
        "Служебный ресурс. Содержимое заменяется генератором документа.",
        style=STYLE_BODY,
    )
    return save_sanitized_docx_atomically(document, output_path)


def _configure_page(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Twips(PAGE_WIDTH_DXA)
    section.page_height = Twips(PAGE_HEIGHT_DXA)
    section.left_margin = Twips(LEFT_MARGIN_DXA)
    section.right_margin = Twips(RIGHT_MARGIN_DXA)
    section.top_margin = Twips(TOP_MARGIN_DXA)
    section.bottom_margin = Twips(BOTTOM_MARGIN_DXA)
    section.header_distance = Twips(HEADER_DISTANCE_DXA)
    section.footer_distance = Twips(FOOTER_DISTANCE_DXA)
    section.gutter = Twips(0)
    section.different_first_page_header_footer = False


def configure_mdrk_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    _set_font(normal)
    _set_paragraph_format(normal, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    body = _paragraph_style(styles, STYLE_BODY, normal)
    _set_font(body)
    _set_paragraph_format(body, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    title = _paragraph_style(styles, STYLE_TITLE, body)
    _set_font(title)
    _set_paragraph_format(
        title,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        keep_with_next=True,
    )

    meeting = _paragraph_style(styles, STYLE_MEETING, body)
    _set_font(meeting)
    _set_paragraph_format(
        meeting,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        keep_with_next=True,
    )

    section = _paragraph_style(styles, STYLE_SECTION, body)
    _set_font(section)
    _set_paragraph_format(
        section,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        space_before_pt=3,
        keep_with_next=True,
    )

    table = _paragraph_style(styles, STYLE_TABLE, body)
    _set_font(table, size_pt=TABLE_FONT_SIZE_PT)
    _set_paragraph_format(table, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    table_header = _paragraph_style(styles, STYLE_TABLE_HEADER, table)
    _set_font(table_header, bold=True, size_pt=TABLE_FONT_SIZE_PT)
    _set_paragraph_format(
        table_header,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        keep_with_next=True,
    )

    mcf_code = _paragraph_style(styles, STYLE_MCF_CODE, table)
    _set_font(mcf_code, size_pt=TABLE_FONT_SIZE_PT)
    _set_paragraph_format(mcf_code, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    try:
        task_base = styles["List Number"]
    except KeyError:
        task_base = body
    task = _paragraph_style(styles, STYLE_TASK, task_base)
    _set_font(task)
    _set_paragraph_format(task, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    warning = _paragraph_style(styles, STYLE_WARNING, body)
    _set_font(warning, bold=True)
    _set_paragraph_format(warning, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    label = _character_style(styles, STYLE_LABEL)
    _set_font(label, bold=True)


def _paragraph_style(styles: object, name: str, base: BaseStyle) -> BaseStyle:
    try:
        style = styles[name]  # type: ignore[index]
    except KeyError:
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)  # type: ignore[attr-defined]
    style.base_style = base
    return style


def _character_style(styles: object, name: str) -> BaseStyle:
    try:
        return styles[name]  # type: ignore[index]
    except KeyError:
        return styles.add_style(name, WD_STYLE_TYPE.CHARACTER)  # type: ignore[attr-defined]


def _set_font(
    style: BaseStyle,
    *,
    bold: bool = False,
    size_pt: int = FONT_SIZE_PT,
) -> None:
    style.font.name = FONT_NAME
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)
    run_properties = style.element.get_or_add_rPr()
    fonts = run_properties.get_or_add_rFonts()
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attribute}"), FONT_NAME)


def _set_paragraph_format(
    style: BaseStyle,
    *,
    alignment: WD_ALIGN_PARAGRAPH,
    space_before_pt: int = 0,
    keep_with_next: bool = False,
) -> None:
    paragraph = style.paragraph_format
    paragraph.alignment = alignment
    paragraph.space_before = Pt(space_before_pt)
    paragraph.space_after = Pt(0)
    paragraph.line_spacing = 1.0
    paragraph.keep_with_next = keep_with_next
    paragraph.widow_control = True


def _configure_metadata(document: Document) -> None:
    properties = document.core_properties
    properties.author = "MDRK Builder"
    properties.last_modified_by = "MDRK Builder"
    properties.title = "Канонический шаблон МДРК"
    properties.subject = "Локальное формирование МДРК"
    properties.comments = "Санитизированный служебный ресурс без данных пациента"
    properties.keywords = ""
    properties.category = ""
