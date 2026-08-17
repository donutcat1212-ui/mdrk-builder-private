from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from mdrk_builder.infrastructure.docx_output import (
    resolve_docx_output_path,
    save_sanitized_docx_atomically,
)


TEMPLATE_FILENAME = "discharge_summary_template.docx"
FOOTER_PATIENT_PLACEHOLDER = "{{PATIENT_FULL_NAME}}"
FOOTER_RECORD_PLACEHOLDER = "{{MEDICAL_RECORD_NUMBER}}"
ACKNOWLEDGEMENT_TEMPLATE = (
    "Выписной эпикриз получен на руки. С рекомендациями ознакомлен(а)."
)
SIGNING_DATE_TEMPLATE = '"__" __________ 20__ г. ___ час. ___ мин.'


def discharge_summary_template_path() -> Path:
    return Path(__file__).resolve().parent.parent / "resources" / TEMPLATE_FILENAME


def create_discharge_summary_template(
    reference_path: Path,
    output_path: Path,
) -> Path:
    reference = reference_path.resolve()
    if not reference.is_file():
        raise FileNotFoundError(reference)
    output = resolve_docx_output_path(
        output_path,
        template_path=reference,
        source_paths=(reference,),
    )
    document = Document(reference)
    _retain_letterhead_only(document)
    _replace_footer(document)
    _configure_metadata(document)
    return save_sanitized_docx_atomically(document, output)


def _retain_letterhead_only(document: DocxDocument) -> None:
    body = document._element.body
    letterhead = next((child for child in body if child.tag == qn("w:tbl")), None)
    if letterhead is None:
        raise ValueError("reference discharge template has no letterhead table")
    for child in list(body):
        if child is not letterhead and child.tag != qn("w:sectPr"):
            body.remove(child)


def _replace_footer(document: DocxDocument) -> None:
    for section in document.sections:
        footer = section.footer
        for table in list(footer.tables):
            table._element.getparent().remove(table._element)
        paragraphs = footer.paragraphs
        paragraph = paragraphs[0] if paragraphs else footer.add_paragraph()
        paragraph.clear()
        paragraph.alignment = 0
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(
            f"{FOOTER_PATIENT_PLACEHOLDER}, МКП №{FOOTER_RECORD_PLACEHOLDER}"
        )
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(31, 78, 121)
        fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
        for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
            fonts.set(qn(f"w:{attribute}"), "Times New Roman")
        for extra in paragraphs[1:]:
            extra._element.getparent().remove(extra._element)


def _configure_metadata(document: DocxDocument) -> None:
    properties = document.core_properties
    properties.author = "MDRK Builder"
    properties.last_modified_by = "MDRK Builder"
    properties.title = "Шаблон выписного эпикриза"
    properties.subject = "Локальное формирование выписного эпикриза"
    properties.comments = "Санитизированный служебный ресурс без данных пациента"
    properties.keywords = ""
    properties.category = ""
