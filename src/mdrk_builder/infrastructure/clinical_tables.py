from __future__ import annotations

from collections import OrderedDict
from collections.abc import Sequence

from docx.document import Document as DocxDocument
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import _Cell, _Row

from mdrk_builder.domain import IcfDomain, IcfQualifier, MdrkKind, Procedure, SpecialistRole

from .docx_layout import (
    MCF_FINAL_WIDTHS,
    MCF_INITIAL_WIDTHS,
    PROCEDURE_WIDTHS,
    compact_header_cell,
    configure_table,
    mark_header_row,
    set_cell_shading,
    set_cell_text,
)
from .docx_template import (
    STYLE_MCF_CODE,
    STYLE_TABLE,
    STYLE_TABLE_HEADER,
    STYLE_WARNING,
)


MCF_QUALIFIER_FILL = "BFBFBF"


def render_icf_profile(
    document: DocxDocument,
    kind: MdrkKind,
    domains: Sequence[IcfDomain],
    *,
    repeat_missing_final: bool = True,
) -> None:
    grouped: OrderedDict[str, list[IcfDomain]] = OrderedDict()
    for domain in domains:
        grouped.setdefault(_mcf_category(domain.code), []).append(domain)

    category_header_rows = sum(
        1 if _is_personal_factor(items[0].code) else 2
        for items in grouped.values()
    )
    row_count = 2 + category_header_rows + len(domains)
    widths = MCF_FINAL_WIDTHS if kind is MdrkKind.FINAL else MCF_INITIAL_WIDTHS
    table = document.add_table(rows=row_count, cols=len(widths))
    configure_table(table, widths)

    title = table.rows[0].cells[0].merge(table.rows[0].cells[-1])
    set_cell_text(
        title,
        "МКФ категориальный профиль",
        style=STYLE_TABLE_HEADER,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        keep_with_next=True,
    )

    header = table.rows[1]
    category = header.cells[0].merge(header.cells[1])
    classifier = header.cells[2].merge(header.cells[10])
    set_cell_text(
        category,
        "МКФ категории",
        style=STYLE_TABLE_HEADER,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        keep_with_next=True,
    )
    set_cell_text(
        classifier,
        "МКФ классификатор",
        style=STYLE_TABLE_HEADER,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        keep_with_next=True,
    )
    for cell in header.cells[11:14]:
        set_cell_text(
            cell,
            "",
            style=STYLE_TABLE_HEADER,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            keep_with_next=True,
        )
    set_cell_text(
        header.cells[14],
        "+/-" if kind is MdrkKind.FINAL else "",
        style=STYLE_TABLE_HEADER,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        keep_with_next=True,
    )
    for cell in header.cells[11:13]:
        compact_header_cell(cell)

    for row in table.rows[:2]:
        mark_header_row(row)

    row_index = 2
    for category_name, category_domains in grouped.items():
        sample_code = category_domains[0].code
        if _is_personal_factor(sample_code):
            _fill_mcf_personal_factor_header(table.rows[row_index], category_name)
            row_index += 1
        else:
            header_row = table.rows[row_index]
            scale_row = table.rows[row_index + 1]
            if _is_environment_factor(sample_code):
                _fill_mcf_environment_headers(
                    header_row,
                    scale_row,
                    category_name,
                    kind,
                )
            else:
                _fill_mcf_problem_headers(
                    header_row,
                    scale_row,
                    category_name,
                    data_label=(
                        "Данные"
                        if sample_code.strip().casefold().startswith("s")
                        else "Ответственный специалист МДРК"
                    ),
                )
            row_index += 2
        environment_group = _is_environment_factor(sample_code)
        for domain_index, domain in enumerate(category_domains):
            domain_row = table.rows[row_index]
            _fill_mcf_domain_row(
                domain_row,
                domain,
                kind,
                repeat_missing_final=repeat_missing_final,
            )
            if environment_group and domain_index < len(category_domains) - 1:
                for cell in domain_row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.keep_with_next = True
            row_index += 1


def render_final_icf_profile(
    document: DocxDocument,
    domains: Sequence[IcfDomain],
    *,
    repeat_missing_final: bool = True,
) -> None:
    render_icf_profile(
        document,
        MdrkKind.FINAL,
        domains,
        repeat_missing_final=repeat_missing_final,
    )


def render_completed_program(
    document: DocxDocument,
    procedures: Sequence[Procedure],
) -> None:
    table = document.add_table(
        rows=max(1, len(procedures)) + 1,
        cols=len(PROCEDURE_WIDTHS),
    )
    configure_table(table, PROCEDURE_WIDTHS)
    headers = (
        "Реабилитационные процедуры",
        "Ответственный специалист",
        "Количество",
        "Продолжительность в мин.",
        "Кратность",
    )
    for cell, value in zip(table.rows[0].cells, headers, strict=True):
        set_cell_text(
            cell,
            value,
            style=STYLE_TABLE_HEADER,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            keep_with_next=True,
        )
    compact_header_cell(table.rows[0].cells[2])
    mark_header_row(table.rows[0])

    if not procedures:
        empty = table.rows[1].cells[0].merge(table.rows[1].cells[-1])
        set_cell_text(
            empty,
            "Мероприятия не представлены",
            style=STYLE_WARNING,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
        )
        return

    for row, procedure in zip(table.rows[1:], procedures, strict=True):
        procedure_name = " ".join(
            item for item in (procedure.code, procedure.name) if item
        ).strip()
        values = (
            procedure_name,
            procedure.specialist,
            "" if procedure.actual_count is None else procedure.actual_count,
            "" if procedure.duration_minutes is None else procedure.duration_minutes,
            procedure.frequency,
        )
        for cell, value in zip(row.cells, values, strict=True):
            set_cell_text(
                cell,
                value,
                style=STYLE_TABLE,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP,
            )


def _fill_mcf_problem_headers(
    header_row: _Row,
    scale_row: _Row,
    category_name: str,
    *,
    data_label: str,
) -> None:
    header_cells = header_row.cells
    category_blank = header_cells[0].merge(header_cells[1])
    problems = header_cells[6].merge(header_cells[10])
    set_cell_text(category_blank, "", style=STYLE_TABLE_HEADER, keep_with_next=True)
    for cell in header_cells[2:6]:
        set_cell_text(cell, "", style=STYLE_TABLE_HEADER, keep_with_next=True)
    set_cell_text(
        problems,
        "Проблемы",
        style=STYLE_TABLE_HEADER,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        keep_with_next=True,
    )
    for cell in header_cells[11:13]:
        set_cell_text(cell, "", style=STYLE_TABLE_HEADER, keep_with_next=True)
    set_cell_text(
        header_cells[13],
        data_label,
        style=STYLE_TABLE_HEADER,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        keep_with_next=True,
    )
    set_cell_text(header_cells[14], "", style=STYLE_TABLE_HEADER, keep_with_next=True)
    for cell in header_cells[2:13]:
        compact_header_cell(cell)

    scale_cells = scale_row.cells
    category = scale_cells[0].merge(scale_cells[1])
    set_cell_text(
        category,
        category_name,
        style=STYLE_TABLE_HEADER,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        keep_with_next=True,
    )
    for index, cell in enumerate(scale_cells[2:], start=2):
        value = str(index - 6) if 6 <= index <= 10 else ""
        set_cell_text(
            cell,
            value,
            style=STYLE_TABLE_HEADER,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            keep_with_next=True,
        )
    for cell in scale_cells[2:13]:
        compact_header_cell(cell)


def _fill_mcf_environment_headers(
    header_row: _Row,
    scale_row: _Row,
    category_name: str,
    kind: MdrkKind,
) -> None:
    header_cells = header_row.cells
    category_blank = header_cells[0].merge(header_cells[1])
    positive = header_cells[2].merge(header_cells[5])
    barriers = header_cells[7].merge(header_cells[10])
    set_cell_text(category_blank, "", style=STYLE_TABLE_HEADER, keep_with_next=True)
    set_cell_text(
        positive,
        "Позитивные\nфакторы",
        style=STYLE_TABLE_HEADER,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        keep_with_next=True,
    )
    set_cell_text(header_cells[6], "", style=STYLE_TABLE_HEADER, keep_with_next=True)
    set_cell_text(
        barriers,
        "Барьеры",
        style=STYLE_TABLE_HEADER,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        keep_with_next=True,
    )
    for cell in header_cells[11:13]:
        set_cell_text(cell, "", style=STYLE_TABLE_HEADER, keep_with_next=True)
    set_cell_text(
        header_cells[13],
        "Уточнение",
        style=STYLE_TABLE_HEADER,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        keep_with_next=True,
    )
    set_cell_text(
        header_cells[14],
        "+/-" if kind is MdrkKind.FINAL else "",
        style=STYLE_TABLE_HEADER,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        keep_with_next=True,
    )
    for cell in header_cells[2:13]:
        compact_header_cell(cell)

    scale_cells = scale_row.cells
    category = scale_cells[0].merge(scale_cells[1])
    set_cell_text(
        category,
        category_name,
        style=STYLE_TABLE_HEADER,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        keep_with_next=True,
    )
    values = ("4+", "3+", "2+", "1+", "0", "1", "2", "3", "4")
    for cell, value in zip(scale_cells[2:11], values, strict=True):
        set_cell_text(
            cell,
            value,
            style=STYLE_TABLE_HEADER,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            keep_with_next=True,
        )
    for cell in scale_cells[11:]:
        set_cell_text(cell, "", style=STYLE_TABLE_HEADER, keep_with_next=True)
    for cell in scale_cells[2:13]:
        compact_header_cell(cell)


def _fill_mcf_personal_factor_header(row: _Row, category_name: str) -> None:
    band = row.cells[0].merge(row.cells[-1])
    set_cell_text(
        band,
        category_name,
        style=STYLE_TABLE_HEADER,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        keep_with_next=True,
    )


def _fill_mcf_domain_row(
    row: _Row,
    domain: IcfDomain,
    kind: MdrkKind,
    *,
    repeat_missing_final: bool,
) -> None:
    cells = row.cells
    set_cell_text(
        cells[0],
        domain.code,
        style=STYLE_MCF_CODE,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    if _is_personal_factor(domain.code):
        description = cells[1].merge(cells[-1])
        set_cell_text(
            description,
            domain.description,
            style=STYLE_TABLE,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP,
        )
        return
    set_cell_text(
        cells[1],
        domain.description,
        style=STYLE_TABLE,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP,
    )
    for cell in cells[2:11]:
        set_cell_text(
            cell,
            "",
            style=STYLE_TABLE,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )
    _shade_initial_qualifier(cells, domain.initial)
    initial = domain.initial.display() if domain.initial is not None else ""
    final_qualifier = domain.final
    if (
        kind is MdrkKind.FINAL
        and final_qualifier is None
        and repeat_missing_final
    ):
        final_qualifier = domain.initial
    final = (
        final_qualifier.display()
        if kind is MdrkKind.FINAL and final_qualifier is not None
        else ""
    )
    marker = domain.dynamic_marker if kind is MdrkKind.FINAL else ""
    set_cell_text(
        cells[11],
        initial,
        style=STYLE_TABLE,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    set_cell_text(
        cells[12],
        final,
        style=STYLE_TABLE,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    set_cell_text(
        cells[13],
        _mcf_responsible(domain),
        style=STYLE_TABLE,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    )
    set_cell_text(
        cells[14],
        marker or "",
        style=STYLE_TABLE,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )


def _mcf_category(code: str) -> str:
    normalized = code.strip().casefold().replace("е", "e")
    if normalized.startswith("b"):
        return "Структура/функция"
    if normalized.startswith("s"):
        return "Структуры организма"
    if normalized.startswith("d"):
        return "Активность/участие"
    if normalized.startswith("e"):
        return "Факторы окружающей среды"
    if normalized.startswith("pf"):
        return "Персональные факторы"
    return "Другие домены"


def _is_personal_factor(code: str) -> bool:
    return code.strip().casefold().replace(" ", "").startswith("pf")


def _is_environment_factor(code: str) -> bool:
    return code.strip().casefold().replace("е", "e").startswith("e")


def _mcf_responsible(domain: IcfDomain) -> str:
    if domain.note.strip():
        return domain.note.strip()
    if domain.specialist is SpecialistRole.OTHER:
        return ""
    return domain.specialist.display_name


def _shade_initial_qualifier(
    cells: Sequence[_Cell],
    qualifier: IcfQualifier | None,
) -> None:
    if qualifier is None:
        return
    zero_column = 6
    value = qualifier.value
    if qualifier.facilitator:
        qualifier_columns = range(zero_column - value, zero_column + 1)
    else:
        qualifier_columns = range(zero_column, zero_column + value + 1)
    for column in qualifier_columns:
        set_cell_shading(cells[column], MCF_QUALIFIER_FILL)
