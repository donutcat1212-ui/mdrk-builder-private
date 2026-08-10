from __future__ import annotations

from collections.abc import Sequence

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell, _Row, Table


PAGE_WIDTH_DXA = 11906
PAGE_HEIGHT_DXA = 16838
LEFT_MARGIN_DXA = 1701
RIGHT_MARGIN_DXA = 850
TOP_MARGIN_DXA = 1134
BOTTOM_MARGIN_DXA = 1134
HEADER_DISTANCE_DXA = 708
FOOTER_DISTANCE_DXA = 708

TABLE_WIDTH_DXA = 9345
MCF_INITIAL_WIDTHS = (
    716,
    2573,
    430,
    430,
    430,
    430,
    316,
    316,
    316,
    316,
    317,
    429,
    222,
    1651,
    453,
)
MCF_FINAL_WIDTHS = (
    716,
    2130,
    430,
    430,
    430,
    430,
    316,
    316,
    316,
    316,
    316,
    429,
    666,
    1651,
    453,
)
# Backward-compatible public name for the initial form geometry.
MCF_WIDTHS = MCF_INITIAL_WIDTHS
SCALE_INITIAL_WIDTHS = (7000, 2345)
SCALE_FINAL_WIDTHS = (5655, 1845, 1845)
ORDINARY_SCALE_WIDTHS = (1700, 5400, 2245)
PROCEDURE_WIDTHS = (2787, 1672, 1394, 1950, 1542)
SIGNATURE_WIDTHS = (4440, 2700, 2205)

CELL_MARGIN_VERTICAL_DXA = 40
CELL_MARGIN_HORIZONTAL_DXA = 60
BORDER_SIZE_EIGHTH_POINTS = 4


def _replace_child(
    parent: object, tag: str, *, successors: Sequence[str] = ()
) -> OxmlElement:
    qualified = qn(tag)
    for child in list(parent):  # type: ignore[arg-type]
        if child.tag == qualified:
            parent.remove(child)  # type: ignore[attr-defined]
    element = OxmlElement(tag)
    parent.insert_element_before(element, *successors)  # type: ignore[attr-defined]
    return element


def _set_dxa(element: OxmlElement, value: int) -> None:
    element.set(qn("w:w"), str(value))
    element.set(qn("w:type"), "dxa")


def configure_table(table: Table, widths: Sequence[int]) -> None:
    """Apply the deterministic table geometry shared by every generated document."""

    if len(widths) != len(table.columns):
        raise ValueError(f"expected {len(table.columns)} widths, got {len(widths)}")
    if sum(widths) != TABLE_WIDTH_DXA:
        raise ValueError(f"table widths must sum to {TABLE_WIDTH_DXA}, got {sum(widths)}")

    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_properties = table._tbl.tblPr

    _set_dxa(
        _replace_child(
            table_properties,
            "w:tblW",
            successors=(
                "w:jc",
                "w:tblCellSpacing",
                "w:tblInd",
                "w:tblBorders",
                "w:shd",
                "w:tblLayout",
                "w:tblCellMar",
                "w:tblLook",
                "w:tblCaption",
                "w:tblDescription",
                "w:tblPrChange",
            ),
        ),
        TABLE_WIDTH_DXA,
    )
    _set_dxa(
        _replace_child(
            table_properties,
            "w:tblInd",
            successors=(
                "w:tblBorders",
                "w:shd",
                "w:tblLayout",
                "w:tblCellMar",
                "w:tblLook",
                "w:tblCaption",
                "w:tblDescription",
                "w:tblPrChange",
            ),
        ),
        0,
    )
    layout = _replace_child(
        table_properties,
        "w:tblLayout",
        successors=(
            "w:tblCellMar",
            "w:tblLook",
            "w:tblCaption",
            "w:tblDescription",
            "w:tblPrChange",
        ),
    )
    layout.set(qn("w:type"), "fixed")
    _set_table_cell_margins(table_properties)
    _set_table_borders(table_properties)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    for row in table.rows:
        set_cant_split(row)
        _remove_fixed_height(row)
        for cell, width in zip(row.cells, widths, strict=True):
            cell_width = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            _set_dxa(cell_width, width)


def _set_table_cell_margins(table_properties: object) -> None:
    margins = _replace_child(
        table_properties,
        "w:tblCellMar",
        successors=("w:tblLook", "w:tblCaption", "w:tblDescription", "w:tblPrChange"),
    )
    for side, value in (
        ("top", CELL_MARGIN_VERTICAL_DXA),
        ("left", CELL_MARGIN_HORIZONTAL_DXA),
        ("bottom", CELL_MARGIN_VERTICAL_DXA),
        ("right", CELL_MARGIN_HORIZONTAL_DXA),
    ):
        element = OxmlElement(f"w:{side}")
        _set_dxa(element, value)
        margins.append(element)


def _set_table_borders(table_properties: object) -> None:
    borders = _replace_child(
        table_properties,
        "w:tblBorders",
        successors=(
            "w:shd",
            "w:tblLayout",
            "w:tblCellMar",
            "w:tblLook",
            "w:tblCaption",
            "w:tblDescription",
            "w:tblPrChange",
        ),
    )
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(BORDER_SIZE_EIGHTH_POINTS))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        borders.append(border)


def _remove_fixed_height(row: _Row) -> None:
    row_properties = row._tr.get_or_add_trPr()
    for height in list(row_properties.findall(qn("w:trHeight"))):
        row_properties.remove(height)


def set_repeat_table_header(row: _Row) -> None:
    row_properties = row._tr.get_or_add_trPr()
    header = row_properties.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        row_properties.append(header)
    header.set(qn("w:val"), "true")


def set_cant_split(row: _Row) -> None:
    row_properties = row._tr.get_or_add_trPr()
    marker = row_properties.find(qn("w:cantSplit"))
    if marker is None:
        marker = OxmlElement("w:cantSplit")
        row_properties.append(marker)
    marker.set(qn("w:val"), "true")


def set_cell_text(
    cell: _Cell,
    text: object,
    *,
    style: str,
    alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
    vertical_alignment: WD_CELL_VERTICAL_ALIGNMENT = WD_CELL_VERTICAL_ALIGNMENT.CENTER,
    keep_with_next: bool = False,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.style = style
    paragraph.alignment = alignment
    paragraph.paragraph_format.keep_with_next = keep_with_next
    paragraph.add_run("" if text is None else str(text))
    cell.vertical_alignment = vertical_alignment


def set_cell_shading(cell: _Cell, fill: str) -> None:
    """Apply one explicit solid OOXML cell fill using an RGB hex value."""

    normalized = fill.strip().upper().removeprefix("#")
    if len(normalized) != 6 or any(character not in "0123456789ABCDEF" for character in normalized):
        raise ValueError("fill must be a six-digit RGB hex value")

    cell_properties = cell._tc.get_or_add_tcPr()
    for existing in list(cell_properties.findall(qn("w:shd"))):
        cell_properties.remove(existing)
    shading = _replace_child(
        cell_properties,
        "w:shd",
        successors=(
            "w:noWrap",
            "w:tcMar",
            "w:textDirection",
            "w:tcFitText",
            "w:vAlign",
            "w:hideMark",
            "w:headers",
            "w:cellIns",
            "w:cellDel",
            "w:cellMerge",
            "w:tcPrChange",
        ),
    )
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), normalized)


def set_cell_no_wrap(cell: _Cell) -> None:
    properties = cell._tc.get_or_add_tcPr()
    marker = properties.find(qn("w:noWrap"))
    if marker is None:
        marker = OxmlElement("w:noWrap")
        properties.insert_element_before(
            marker,
            "w:tcMar",
            "w:textDirection",
            "w:tcFitText",
            "w:vAlign",
            "w:hideMark",
            "w:headers",
            "w:cellIns",
            "w:cellDel",
            "w:cellMerge",
            "w:tcPrChange",
        )
    marker.set(qn("w:val"), "true")


def set_cell_horizontal_margins(cell: _Cell, *, left: int, right: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = _replace_child(
        properties,
        "w:tcMar",
        successors=(
            "w:textDirection",
            "w:tcFitText",
            "w:vAlign",
            "w:hideMark",
            "w:headers",
            "w:cellIns",
            "w:cellDel",
            "w:cellMerge",
            "w:tcPrChange",
        ),
    )
    for side, value in (("left", left), ("right", right)):
        element = OxmlElement(f"w:{side}")
        _set_dxa(element, value)
        margins.append(element)


def mark_header_row(row: _Row) -> None:
    set_repeat_table_header(row)
    set_cant_split(row)
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.keep_with_next = True
