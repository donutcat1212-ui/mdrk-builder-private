from __future__ import annotations

from datetime import date, datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from mdrk_builder.domain import ReverseSheetDraft, ReverseSheetRow

from .docx_output import resolve_docx_output_path, save_sanitized_docx_atomically


_GRID_WIDTHS = (2739, 1420, 1826, 81, 1746, 1826)
_MONTHS = {
    1: "января",
    2: "февраля",
    3: "марта",
    4: "апреля",
    5: "мая",
    6: "июня",
    7: "июля",
    8: "августа",
    9: "сентября",
    10: "октября",
    11: "ноября",
    12: "декабря",
}


def _set_cell_borders(cell) -> None:
    properties = cell._tc.get_or_add_tcPr()
    borders = properties.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), "000000")


def _set_cell_margins(cell, *, top: int = 20, left: int = 35, bottom: int = 20, right: int = 35) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, value in (("top", top), ("start", left), ("bottom", bottom), ("end", right)):
        element = margins.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _format_date(value: date | None) -> str:
    return value.strftime("%d.%m.%Y") if value else ""


def _format_datetime(value: datetime | None) -> str:
    return value.strftime("%d.%m.%Y %H:%M") if value else ""


def _age_on(birth_date: date | None, at: date | None) -> int | None:
    if birth_date is None or at is None or birth_date > at:
        return None
    return at.year - birth_date.year - ((at.month, at.day) < (birth_date.month, birth_date.day))


def _birth_line(birth_date: date | None, admission: datetime | None) -> str:
    if birth_date is None:
        return ""
    value = f"«{birth_date.day:02d}» {_MONTHS[birth_date.month]} {birth_date.year}г."
    age = _age_on(birth_date, admission.date() if admission else None)
    return f"{value}({_age_label(age)})" if age is not None else value


def _age_label(value: int) -> str:
    if value % 10 == 1 and value % 100 != 11:
        suffix = "год"
    elif value % 10 in {2, 3, 4} and value % 100 not in {12, 13, 14}:
        suffix = "года"
    else:
        suffix = "лет"
    return f"{value} {suffix}"


def _set_run_font(run, *, size: float = 7.5, bold: bool = False, underline: bool = False) -> None:
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.underline = underline


def _set_cell_text(cell, text: str, *, size: float = 7.5, bold: bool = False, alignment=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(1.45)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(7.7 if size <= 7 else 10.7)
    run = paragraph.add_run(text)
    _set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _add_identity_line(
    paragraph,
    label: str,
    value: str,
    *,
    bold: bool = False,
    final: bool = False,
) -> None:
    label_run = paragraph.add_run(label)
    _set_run_font(label_run, size=10)
    value_run = paragraph.add_run(value)
    _set_run_font(value_run, size=10, bold=bold, underline=bool(value))
    if not final:
        paragraph.add_run("\n")


def _mark_repeat_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def _configure_table(document, row_count: int):
    table = document.add_table(rows=row_count, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    properties = table._tbl.tblPr
    position = OxmlElement("w:tblpPr")
    position.set(qn("w:leftFromText"), "567")
    position.set(qn("w:vertAnchor"), "page")
    position.set(qn("w:horzAnchor"), "page")
    position.set(qn("w:tblpX"), "567")
    position.set(qn("w:tblpY"), "567")
    properties.insert(0, position)
    width = properties.first_child_found_in("w:tblW")
    if width is not None:
        width.set(qn("w:w"), "9638")
        width.set(qn("w:type"), "dxa")
    layout = properties.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in _GRID_WIDTHS:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Pt(26.55)
        for index, cell in enumerate(row.cells):
            cell.width = Inches(_GRID_WIDTHS[index] / 1440)
            _set_cell_margins(cell, top=0, left=15, bottom=0, right=15)
    return table


def _render_data_row(cells, row: ReverseSheetRow | None) -> None:
    values = (
        row.intervention if row else "",
        _format_date(row.appointment_date) if row else "",
        "",
        _format_datetime(row.performed_at) if row else "",
        row.performer if row else "",
    )
    sizes = (8, 10, 8, 10, 10)
    alignments = (
        WD_ALIGN_PARAGRAPH.JUSTIFY,
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.JUSTIFY,
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.LEFT,
    )
    for cell, value, size, alignment in zip(cells, values, sizes, alignments, strict=True):
        _set_cell_borders(cell)
        _set_cell_text(cell, value, size=size, alignment=alignment)


def _rows_with_manual_gaps(
    rows: list[ReverseSheetRow],
    *,
    gap_rows: int = 2,
) -> list[ReverseSheetRow | None]:
    arranged: list[ReverseSheetRow | None] = []
    previous_key: date | None | object = object()
    for row in rows:
        block_key = row.appointment_date or (
            row.performed_at.date() if row.performed_at is not None else None
        )
        if arranged and block_key != previous_key:
            arranged.extend([None] * gap_rows)
        arranged.append(row)
        previous_key = block_key
    return arranged


def write_reverse_sheet_docx(
    draft: ReverseSheetDraft,
    output_path: Path,
    *,
    minimum_data_rows: int = 24,
) -> Path:
    source_paths = {
        path
        for path in [draft.header_source, *(row.source for row in draft.rows)]
        if path is not None
    }
    output = resolve_docx_output_path(output_path, source_paths=source_paths)

    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = 7_560_310
    section.page_height = 10_692_130
    section.left_margin = 1_080_135
    section.right_margin = 539_750
    section.top_margin = 720_090
    section.bottom_margin = 720_090
    document.core_properties.title = "Оборотная сторона раздела — лист назначений и их выполнение"
    document.core_properties.author = "MDRK Builder"
    document.core_properties.subject = "Редактируемый проект медицинской формы"
    document.core_properties.keywords = ""
    document.core_properties.comments = ""

    arranged_rows = _rows_with_manual_gaps(draft.rows)
    data_rows = max(minimum_data_rows, len(arranged_rows))
    table = _configure_table(document, 3 + data_rows)

    title_left = table.cell(0, 0).merge(table.cell(0, 3))
    _set_cell_text(title_left, "", size=8)
    title = table.cell(0, 4).merge(table.cell(0, 5))
    _set_cell_text(
        title,
        'оборотная сторона раздела\n"Лист назначений и их выполнение"',
        size=10,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    table.rows[0].height = Pt(22.7)

    identity_cell = table.cell(1, 0).merge(table.cell(1, 5))
    identity_cell.text = ""
    identity_paragraph = identity_cell.paragraphs[0]
    identity_paragraph.paragraph_format.space_before = Pt(1.45)
    identity_paragraph.paragraph_format.space_after = Pt(0)
    identity_paragraph.paragraph_format.line_spacing = Pt(10.7)
    _add_identity_line(
        identity_paragraph,
        "Фамилия, имя, отчество (при наличии) ",
        draft.identity.full_name,
        bold=True,
    )
    _add_identity_line(identity_paragraph, "Дата рождения: ", _birth_line(draft.identity.birth_date, draft.admission_datetime))
    _add_identity_line(
        identity_paragraph,
        "Медицинская карта пациента, получающего медицинскую помощь в стационарных условиях, "
        "в условиях дневного стационара № ",
        draft.identity.medical_record_number,
        bold=True,
    )
    admission = draft.admission_datetime
    admission_value = ""
    if admission is not None:
        admission_value = (
            f"«{admission.day:02d}» {_MONTHS[admission.month]} {admission.year}г., "
            f"время {admission.hour:02d} час. {admission.minute:02d} мин."
        )
    _add_identity_line(
        identity_paragraph,
        "Дата и время поступления ",
        admission_value,
        final=True,
    )
    identity_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    table.rows[1].height = Pt(69.95)

    header_cells = [
        table.cell(2, 0),
        table.cell(2, 1),
        table.cell(2, 2),
        table.cell(2, 3).merge(table.cell(2, 4)),
        table.cell(2, 5),
    ]
    headers = (
        "Медицинское вмешательство",
        "Дата назначения",
        "Подпись лечащего врача\n(врача-специалиста),\nсделавшего назначение",
        "Дата и время исполнения\nназначения",
        "Фамилия, имя, отчество\n(при наличии) и подпись\nмедицинского работника,\n"
        "ответственного за\nисполнение назначения",
    )
    for cell, value in zip(header_cells, headers, strict=True):
        _set_cell_borders(cell)
        alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if cell is header_cells[0] else WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_text(cell, value, size=7, alignment=alignment)
    table.rows[2].height = Pt(47.15)
    _mark_repeat_header(table.rows[2])

    for index in range(data_rows):
        row_cells = [
            table.cell(index + 3, 0),
            table.cell(index + 3, 1),
            table.cell(index + 3, 2),
            table.cell(index + 3, 3).merge(table.cell(index + 3, 4)),
            table.cell(index + 3, 5),
        ]
        _render_data_row(
            row_cells,
            arranged_rows[index] if index < len(arranged_rows) else None,
        )

    return save_sanitized_docx_atomically(document, output)
