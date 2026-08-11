from __future__ import annotations

import re
from collections import OrderedDict
from collections.abc import Sequence
from dataclasses import dataclass
from datetime import date, datetime
from os import replace
from pathlib import Path
from tempfile import NamedTemporaryFile

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.table import _Cell

from mdrk_builder.application.snapshot import ScaleRow, Snapshot, build_snapshot
from mdrk_builder.application.validation import current_issues
from mdrk_builder.domain import (
    Episode,
    IcfDomain,
    IcfQualifier,
    MdrkKind,
    ReviewIssue,
    ReviewSeverity,
    ScaleMeasurement,
    SpecialistFinding,
    SpecialistRole,
)

from .docx_layout import (
    MCF_FINAL_WIDTHS,
    MCF_INITIAL_WIDTHS,
    ORDINARY_SCALE_WIDTHS,
    PROCEDURE_WIDTHS,
    SCALE_FINAL_WIDTHS,
    SCALE_INITIAL_WIDTHS,
    SIGNATURE_WIDTHS,
    configure_table,
    mark_header_row,
    set_cant_split,
    set_cell_horizontal_margins,
    set_cell_no_wrap,
    set_cell_shading,
    set_cell_text,
)
from .docx_template import (
    STYLE_BODY,
    STYLE_LABEL,
    STYLE_MCF_CODE,
    STYLE_MEETING,
    STYLE_SECTION,
    STYLE_TABLE,
    STYLE_TABLE_HEADER,
    STYLE_TASK,
    STYLE_TITLE,
    STYLE_WARNING,
    canonical_template_path,
)


CONSILIUM_TITLE = (
    "Консилиум мультидисциплинарной реабилитационной команды в составе "
    "заведующего отделением, врача фрм, специалиста по физической реабилитации, "
    "медицинского психолога/нейропсихолога, медицинского логопеда и специалиста "
    "по эргореабилитации"
)
MCF_QUALIFIER_FILL = "BFBFBF"

_MONTHS = {
    1: "января",
    2: "февраля",
    3: "марта",
    4: "апреля",
    5: "мая",
    6: "июня",
    7: "июля",
    8: "августа",
    9: "сентября",
    10: "октября",
    11: "ноября",
    12: "декабря",
}


class DocumentGenerationBlockedError(ValueError):
    def __init__(self, issues: Sequence[ReviewIssue]) -> None:
        self.issues = tuple(issues)
        details = "; ".join(issue.message for issue in self.issues)
        super().__init__(f"Формирование документа заблокировано: {details}")


@dataclass(frozen=True, slots=True)
class SignatoryRow:
    role: str | SpecialistRole
    full_name: str = ""

    @property
    def display_role(self) -> str:
        return self.role.display_name if isinstance(self.role, SpecialistRole) else self.role


SIGNATORY_ROSTER: tuple[SignatoryRow, ...] = (
    SignatoryRow(SpecialistRole.FRM),
    SignatoryRow(SpecialistRole.PHYSICAL_THERAPIST),
    SignatoryRow(SpecialistRole.NEUROPSYCHOLOGIST),
    SignatoryRow(SpecialistRole.PATHOPSYCHOLOGIST),
    SignatoryRow(SpecialistRole.LOGOPEDIST),
    SignatoryRow(SpecialistRole.OCCUPATIONAL_THERAPIST),
    SignatoryRow("Консультанты"),
    SignatoryRow("Заведующий отделением"),
)


def _signatory_role_key(row: SignatoryRow) -> str:
    if isinstance(row.role, SpecialistRole):
        return {
            SpecialistRole.FRM: "frm",
            SpecialistRole.NEUROLOGIST: "frm",
            SpecialistRole.PHYSICAL_THERAPIST: "physical_therapist",
            SpecialistRole.OCCUPATIONAL_THERAPIST: "occupational_therapist",
            SpecialistRole.LOGOPEDIST: "logopedist",
            SpecialistRole.NEUROPSYCHOLOGIST: "neuropsychologist",
            SpecialistRole.PATHOPSYCHOLOGIST: "pathopsychologist",
            SpecialistRole.OTHER: "consultants",
        }[row.role]

    normalized = " ".join(re.sub(r"[^0-9a-zа-яё]+", " ", row.display_role.casefold()).split())
    if "заведующ" in normalized:
        return "department_head"
    if "консульт" in normalized:
        return "consultants"
    if "патопсих" in normalized:
        return "pathopsychologist"
    if "нейропсих" in normalized:
        return "neuropsychologist"
    if "эрго" in normalized:
        return "occupational_therapist"
    if "логоп" in normalized:
        return "logopedist"
    if "фрм" in normalized or "невролог" in normalized or (
        "врач" in normalized and "реабилитац" in normalized
    ):
        return "frm"
    if any(token in normalized for token in ("физическ", "физический терапевт", "лфк")):
        return "physical_therapist"
    return normalized


def write_mdrk_docx(
    episode: Episode,
    kind: MdrkKind,
    output_path: Path,
    *,
    signatories: Sequence[SignatoryRow] = (),
    template_path: Path | None = None,
) -> Path:
    """Create one editable MDRK snapshot without changing any source document."""

    blocking = [
        issue
        for issue in current_issues(episode, kind)
        if issue.severity is ReviewSeverity.BLOCKING
    ]
    if blocking:
        raise DocumentGenerationBlockedError(blocking)

    template = (template_path or canonical_template_path()).resolve()
    if not template.is_file():
        raise FileNotFoundError(
            f"Канонический шаблон не найден: {template}. "
            "Запустите tools/build_canonical_template.py."
        )

    output = output_path.resolve()
    if output.suffix.casefold() != ".docx":
        raise ValueError("output_path must use the .docx extension")
    if output == template:
        raise ValueError("output_path must not overwrite the canonical template")
    source_paths = {source.path.resolve() for source in episode.sources}
    if output in source_paths:
        raise ValueError("output_path must not overwrite an immutable source document")

    snapshot = build_snapshot(episode, kind)
    document = Document(template)
    _clear_body(document)
    _set_output_metadata(document, kind)
    _DocumentRenderer(document, episode, snapshot, tuple(signatories)).render()

    output.parent.mkdir(parents=True, exist_ok=True)
    temporary_path: Path | None = None
    try:
        with NamedTemporaryFile(
            prefix=f".{output.stem}-",
            suffix=".docx",
            dir=output.parent,
            delete=False,
        ) as temporary:
            temporary_path = Path(temporary.name)
        document.save(temporary_path)
        replace(temporary_path, output)
        temporary_path = None
    finally:
        if temporary_path is not None:
            temporary_path.unlink(missing_ok=True)
    return output


class _DocumentRenderer:
    def __init__(
        self,
        document: DocxDocument,
        episode: Episode,
        snapshot: Snapshot,
        signatories: tuple[SignatoryRow, ...],
    ) -> None:
        self.document = document
        self.episode = episode
        self.snapshot = snapshot
        self.signatories = signatories

    def render(self) -> None:
        self._render_title_block()
        self._render_clinical_diagnosis()
        self._render_rehabilitation_course()
        self._render_history()
        self._render_diagnostics()
        self._render_specialists()
        self._render_outcomes()
        self._render_plan()

    def _render_title_block(self) -> None:
        self.document.add_paragraph(CONSILIUM_TITLE, style=STYLE_TITLE)
        self.document.add_paragraph(
            _format_meeting(self.snapshot.meeting_at),
            style=STYLE_MEETING,
        )
        self._add_blank_paragraph()
        identity = self.episode.identity
        record = self.document.add_paragraph(style=STYLE_BODY)
        record.add_run("Номер ИБ: ")
        record_value = record.add_run(identity.medical_record_number)
        record_value.bold = True
        record_value.underline = True

        department = self.episode.department.strip()
        if department and department[-1] not in ".!?":
            department += "."
        self.document.add_paragraph(department, style=STYLE_BODY)
        self._add_blank_paragraph()

        full_name = self.document.add_paragraph(style=STYLE_BODY)
        full_name_label = full_name.add_run("ФИО пациента: ")
        full_name_label.bold = True
        full_name_value = full_name.add_run(identity.full_name)
        full_name_value.bold = True
        full_name_value.underline = True

        self._add_original_birth_line(identity.birth_date)

        sex = self.document.add_paragraph(style=STYLE_BODY)
        sex_label = sex.add_run("Пол: ")
        sex_label.bold = True
        sex.add_run(identity.sex)
        self._add_blank_paragraph()

    def _render_clinical_diagnosis(self) -> None:
        self._add_section_heading(1, "Клинический диагноз")
        self._add_text_or_missing(
            self.snapshot.sections.clinical_diagnosis,
            "Клинический диагноз не указан",
        )
        self._add_labeled("Реабилитационный диагноз: ", "")
        if self.snapshot.icf_domains:
            self._render_mcf_table()
        else:
            self.document.add_paragraph("Профиль МКФ не представлен", style=STYLE_WARNING)

    def _render_rehabilitation_course(self) -> None:
        self._add_section_heading(2, "Сведения о реабилитации")
        self._add_labeled("2.1. Этап медицинской реабилитации: ", self.episode.stage)
        duration = (
            f"{self.episode.course_duration_days} дней"
            if self.episode.course_duration_days is not None
            else ""
        )
        self._add_labeled("2.2. Длительность курса медицинской реабилитации: ", duration)

    def _render_history(self) -> None:
        self._add_section_heading(3, "Анамнез заболевания")
        self._add_text_or_missing(
            self.snapshot.sections.disease_history,
            "Анамнез заболевания не указан",
        )
        self._add_section_heading(4, "Анамнез жизни")
        self._add_text_or_missing(
            self.snapshot.sections.life_history,
            "Анамнез жизни не указан",
        )

    def _render_diagnostics(self) -> None:
        self._add_section_heading(5, "Результаты диагностических мероприятий")
        laboratory = self.snapshot.sections.laboratory_results.strip()
        instrumental = self.snapshot.sections.instrumental_results.strip()
        if laboratory:
            self._add_labeled("5.1. Лабораторные исследования: ", "")
            self._add_multiline(laboratory)
        if instrumental:
            self._add_labeled("5.2. Инструментальные исследования: ", "")
            self._add_multiline(instrumental)
        if not laboratory and not instrumental:
            self.document.add_paragraph("Данные не представлены", style=STYLE_WARNING)

    def _render_specialists(self) -> None:
        self._add_section_heading(6, "Результаты осмотров участников МДРК и консультантов")
        findings = {finding.role: finding for finding in self.snapshot.findings}
        scales_by_role: dict[SpecialistRole, list[ScaleRow]] = {}
        for row in self.snapshot.scale_rows:
            scales_by_role.setdefault(row.role, []).append(row)

        rendered = False
        physician_roles = {SpecialistRole.FRM, SpecialistRole.NEUROLOGIST}
        physician_findings = [
            findings[role] for role in physician_roles if role in findings
        ]
        physician_rows = [
            row for role in physician_roles for row in scales_by_role.get(role, [])
        ]
        if physician_findings or physician_rows:
            rendered = True
            finding = max(
                physician_findings,
                key=lambda item: (
                    bool(item.conclusion.strip()),
                    item.source_datetime or datetime.min,
                ),
                default=None,
            )
            self._render_specialist_block(
                SpecialistRole.FRM,
                finding,
                _deduplicate_physician_scale_rows(physician_rows),
            )

        for role in SpecialistRole:
            if role in physician_roles:
                continue
            finding = findings.get(role)
            scale_rows = scales_by_role.get(role, [])
            if finding is None and not scale_rows:
                continue
            rendered = True
            self._render_specialist_block(role, finding, scale_rows)
        if not rendered:
            self.document.add_paragraph("Заключения специалистов не представлены", style=STYLE_WARNING)

    def _render_specialist_block(
        self,
        role: SpecialistRole,
        finding: SpecialistFinding | None,
        scale_rows: Sequence[ScaleRow],
    ) -> None:
        physician = role in {SpecialistRole.FRM, SpecialistRole.NEUROLOGIST}
        source_datetime = finding.source_datetime if finding is not None else None
        if source_datetime is None:
            source_datetime = _first_scale_datetime(scale_rows, self.snapshot.kind)
        heading = _specialist_result_heading(role, source_datetime)
        paragraph = self.document.add_paragraph(style=STYLE_BODY)
        paragraph.paragraph_format.keep_with_next = True
        paragraph.add_run(heading)
        if physician:
            self._add_blank_paragraph(keep_with_next=True)

        if scale_rows:
            if physician and self.snapshot.kind is MdrkKind.INITIAL:
                self._render_initial_physician_scale_table(scale_rows)
            else:
                self._render_scale_table(scale_rows)

        conclusion = self.document.add_paragraph(style=STYLE_BODY)
        conclusion.paragraph_format.keep_with_next = bool(
            finding is not None and finding.conclusion.strip()
        )
        conclusion.add_run("Заключение: ")
        if finding is not None and finding.conclusion.strip():
            conclusion.add_run(finding.conclusion.strip())

    def _render_outcomes(self) -> None:
        self._add_section_value(
            7,
            "Реабилитационный потенциал",
            self.snapshot.sections.rehabilitation_potential,
        )
        self._add_section_value(
            8,
            "Факторы, ограничивающие проведение реабилитационных мероприятий",
            self.snapshot.sections.limitations,
        )
        self._add_section_value(
            9,
            "Факторы риска проведения реабилитационных мероприятий",
            self.snapshot.sections.risks,
        )
        self._add_section_value(10, "Цель на этап медицинской реабилитации", self.snapshot.goal)

        self._add_section_heading(11, "Задачи медицинской реабилитации")
        task_lines = _split_tasks(self.snapshot.tasks)
        if not task_lines:
            self.document.add_paragraph("Задачи не указаны", style=STYLE_WARNING)
        elif self.snapshot.kind is MdrkKind.FINAL and len(task_lines) == 1:
            self.document.add_paragraph(task_lines[0], style=STYLE_BODY)
        else:
            for task in task_lines:
                self.document.add_paragraph(task, style=STYLE_TASK)

    def _render_plan(self) -> None:
        self.document.add_paragraph(
            "12. Индивидуальный план медицинской реабилитации:",
            style=STYLE_BODY,
        )
        self._add_blank_paragraph()
        self._add_original_plan_heading("Режим и питание:")
        self.document.add_paragraph(
            f"Двигательный режим: {self.snapshot.sections.movement_regimen}",
            style=STYLE_BODY,
        )
        self.document.add_paragraph(
            f"Диета: {self.snapshot.sections.diet}",
            style=STYLE_BODY,
        )
        self._add_blank_paragraph()
        self._add_original_plan_heading("Медикаментозное лечение:")
        medication = self.snapshot.sections.medication.strip()
        if medication:
            self._add_multiline(medication)
        self._add_blank_paragraph()
        self._add_original_plan_heading("Реабилитационные мероприятия:")
        self._add_blank_paragraph(keep_with_next=True)
        self._render_procedure_table()
        self._add_signature_table_separator()
        self._render_signature_table()

    def _add_signature_table_separator(self) -> None:
        """Keep adjacent procedure/signature tables distinct in Word renderers."""

        self._add_blank_paragraph()
        self._add_blank_paragraph(keep_with_next=True)

    def _render_initial_physician_scale_table(
        self,
        rows: Sequence[ScaleRow],
    ) -> None:
        table = self.document.add_table(
            rows=len(rows) + 1,
            cols=len(ORDINARY_SCALE_WIDTHS),
        )
        configure_table(table, ORDINARY_SCALE_WIDTHS)
        headers = (
            "Дата и время\nрасчета шкалы",
            "Шкала/опросник",
            "Результат расчета",
        )
        for cell, value in zip(table.rows[0].cells, headers, strict=True):
            set_cell_text(
                cell,
                value,
                style=STYLE_TABLE,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                keep_with_next=True,
            )
        mark_header_row(table.rows[0])

        for table_row, scale_row in zip(table.rows[1:], rows, strict=True):
            measurement = scale_row.initial
            values = (
                _format_original_scale_datetime(
                    measurement.measured_at if measurement is not None else None
                ),
                scale_row.name,
                measurement.value if measurement is not None else "",
            )
            for cell, value in zip(table_row.cells, values, strict=True):
                set_cell_text(
                    cell,
                    value,
                    style=STYLE_TABLE,
                    alignment=WD_ALIGN_PARAGRAPH.LEFT,
                    vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP,
                )

    def _render_mcf_table(self) -> None:
        grouped: OrderedDict[str, list[IcfDomain]] = OrderedDict()
        for domain in self.snapshot.icf_domains:
            grouped.setdefault(_mcf_category(domain.code), []).append(domain)

        category_header_rows = sum(
            1 if _is_personal_factor(domains[0].code) else 2
            for domains in grouped.values()
        )
        row_count = 2 + category_header_rows + len(self.snapshot.icf_domains)
        widths = (
            MCF_FINAL_WIDTHS
            if self.snapshot.kind is MdrkKind.FINAL
            else MCF_INITIAL_WIDTHS
        )
        table = self.document.add_table(rows=row_count, cols=len(widths))
        configure_table(table, widths)

        title = table.rows[0].cells[0].merge(table.rows[0].cells[-1])
        set_cell_text(
            title,
            "МКФ категориальный профиль",
            style=STYLE_TABLE_HEADER,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            keep_with_next=True,
        )

        header = table.rows[1]
        category = header.cells[0].merge(header.cells[1])
        classifier = header.cells[2].merge(header.cells[10])
        set_cell_text(category, "МКФ категории", style=STYLE_TABLE_HEADER, alignment=WD_ALIGN_PARAGRAPH.CENTER, keep_with_next=True)
        set_cell_text(classifier, "МКФ классификатор", style=STYLE_TABLE_HEADER, alignment=WD_ALIGN_PARAGRAPH.CENTER, keep_with_next=True)
        set_cell_text(header.cells[11], "", style=STYLE_TABLE_HEADER, alignment=WD_ALIGN_PARAGRAPH.CENTER, keep_with_next=True)
        set_cell_text(header.cells[12], "", style=STYLE_TABLE_HEADER, alignment=WD_ALIGN_PARAGRAPH.CENTER, keep_with_next=True)
        set_cell_text(header.cells[13], "", style=STYLE_TABLE_HEADER, alignment=WD_ALIGN_PARAGRAPH.CENTER, keep_with_next=True)
        set_cell_text(
            header.cells[14],
            "+/-" if self.snapshot.kind is MdrkKind.FINAL else "",
            style=STYLE_TABLE_HEADER,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            keep_with_next=True,
        )
        for cell in header.cells[11:13]:
            _compact_header_cell(cell)

        for row in table.rows[:2]:
            mark_header_row(row)

        row_index = 2
        for category_name, domains in grouped.items():
            sample_code = domains[0].code
            if _is_personal_factor(sample_code):
                self._fill_mcf_personal_factor_header(
                    table.rows[row_index],
                    category_name,
                )
                row_index += 1
            else:
                header_row = table.rows[row_index]
                scale_row = table.rows[row_index + 1]
                if _is_environment_factor(sample_code):
                    self._fill_mcf_environment_headers(
                        header_row,
                        scale_row,
                        category_name,
                    )
                else:
                    self._fill_mcf_problem_headers(
                        header_row,
                        scale_row,
                        category_name,
                        data_label=(
                            "Данные"
                            if sample_code.strip().casefold().startswith("s")
                            else "Ответственный специалист МДРК"
                        ),
                    )
                row_index += 2
            environment_group = _is_environment_factor(sample_code)
            for domain_index, domain in enumerate(domains):
                domain_row = table.rows[row_index]
                self._fill_mcf_domain_row(domain_row, domain)
                if environment_group and domain_index < len(domains) - 1:
                    for cell in domain_row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph.paragraph_format.keep_with_next = True
                row_index += 1

    def _fill_mcf_problem_headers(
        self,
        header_row: object,
        scale_row: object,
        category_name: str,
        *,
        data_label: str,
    ) -> None:
        header_cells = header_row.cells  # type: ignore[attr-defined]
        category_blank = header_cells[0].merge(header_cells[1])
        problems = header_cells[6].merge(header_cells[10])
        set_cell_text(category_blank, "", style=STYLE_TABLE_HEADER, keep_with_next=True)
        for cell in header_cells[2:6]:
            set_cell_text(cell, "", style=STYLE_TABLE_HEADER, keep_with_next=True)
        set_cell_text(
            problems,
            "Проблемы",
            style=STYLE_TABLE_HEADER,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            keep_with_next=True,
        )
        for cell in header_cells[11:13]:
            set_cell_text(cell, "", style=STYLE_TABLE_HEADER, keep_with_next=True)
        set_cell_text(
            header_cells[13],
            data_label,
            style=STYLE_TABLE_HEADER,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            keep_with_next=True,
        )
        set_cell_text(header_cells[14], "", style=STYLE_TABLE_HEADER, keep_with_next=True)
        for cell in header_cells[2:13]:
            _compact_header_cell(cell)

        scale_cells = scale_row.cells  # type: ignore[attr-defined]
        category = scale_cells[0].merge(scale_cells[1])
        set_cell_text(
            category,
            category_name,
            style=STYLE_TABLE_HEADER,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            keep_with_next=True,
        )
        for index, cell in enumerate(scale_cells[2:], start=2):
            value = str(index - 6) if 6 <= index <= 10 else ""
            set_cell_text(
                cell,
                value,
                style=STYLE_TABLE_HEADER,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                keep_with_next=True,
            )
        for cell in scale_cells[2:13]:
            _compact_header_cell(cell)

    def _fill_mcf_environment_headers(
        self,
        header_row: object,
        scale_row: object,
        category_name: str,
    ) -> None:
        header_cells = header_row.cells  # type: ignore[attr-defined]
        category_blank = header_cells[0].merge(header_cells[1])
        positive = header_cells[2].merge(header_cells[5])
        barriers = header_cells[7].merge(header_cells[10])
        set_cell_text(category_blank, "", style=STYLE_TABLE_HEADER, keep_with_next=True)
        set_cell_text(
            positive,
            "Позитивные\nфакторы",
            style=STYLE_TABLE_HEADER,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            keep_with_next=True,
        )
        set_cell_text(header_cells[6], "", style=STYLE_TABLE_HEADER, keep_with_next=True)
        set_cell_text(
            barriers,
            "Барьеры",
            style=STYLE_TABLE_HEADER,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            keep_with_next=True,
        )
        for cell in header_cells[11:13]:
            set_cell_text(cell, "", style=STYLE_TABLE_HEADER, keep_with_next=True)
        set_cell_text(
            header_cells[13],
            "Уточнение",
            style=STYLE_TABLE_HEADER,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            keep_with_next=True,
        )
        set_cell_text(
            header_cells[14],
            "+/-" if self.snapshot.kind is MdrkKind.FINAL else "",
            style=STYLE_TABLE_HEADER,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            keep_with_next=True,
        )
        for cell in header_cells[2:13]:
            _compact_header_cell(cell)

        scale_cells = scale_row.cells  # type: ignore[attr-defined]
        category = scale_cells[0].merge(scale_cells[1])
        set_cell_text(
            category,
            category_name,
            style=STYLE_TABLE_HEADER,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            keep_with_next=True,
        )
        values = ("4+", "3+", "2+", "1+", "0", "1", "2", "3", "4")
        for cell, value in zip(scale_cells[2:11], values, strict=True):
            set_cell_text(
                cell,
                value,
                style=STYLE_TABLE_HEADER,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                keep_with_next=True,
            )
        for cell in scale_cells[11:]:
            set_cell_text(cell, "", style=STYLE_TABLE_HEADER, keep_with_next=True)
        for cell in scale_cells[2:13]:
            _compact_header_cell(cell)

    def _fill_mcf_personal_factor_header(
        self,
        row: object,
        category_name: str,
    ) -> None:
        cells = row.cells  # type: ignore[attr-defined]
        band = cells[0].merge(cells[-1])
        set_cell_text(
            band,
            category_name,
            style=STYLE_TABLE_HEADER,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            keep_with_next=True,
        )

    def _fill_mcf_domain_row(self, row: object, domain: IcfDomain) -> None:
        cells = row.cells  # type: ignore[attr-defined]
        set_cell_text(cells[0], domain.code, style=STYLE_MCF_CODE, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        if _is_personal_factor(domain.code):
            description = cells[1].merge(cells[-1])
            set_cell_text(
                description,
                domain.description,
                style=STYLE_TABLE,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP,
            )
            return
        set_cell_text(
            cells[1],
            domain.description,
            style=STYLE_TABLE,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP,
        )
        for cell in cells[2:11]:
            set_cell_text(cell, "", style=STYLE_TABLE, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _shade_initial_qualifier(cells, domain.initial)
        initial = domain.initial.display() if domain.initial is not None else ""
        final = (
            domain.final.display()
            if self.snapshot.kind is MdrkKind.FINAL and domain.final is not None
            else ""
        )
        marker = domain.dynamic_marker if self.snapshot.kind is MdrkKind.FINAL else ""
        set_cell_text(cells[11], initial, style=STYLE_TABLE, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[12], final, style=STYLE_TABLE, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[13], _mcf_responsible(domain), style=STYLE_TABLE, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(cells[14], marker or "", style=STYLE_TABLE, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    def _render_scale_table(self, rows: Sequence[ScaleRow]) -> None:
        final_mode = self.snapshot.kind is MdrkKind.FINAL
        widths = SCALE_FINAL_WIDTHS if final_mode else SCALE_INITIAL_WIDTHS
        table = self.document.add_table(rows=len(rows) + 1, cols=len(widths))
        configure_table(table, widths)

        initial_date = _common_measurement_datetime([row.initial for row in rows])
        current_date = _common_measurement_datetime([row.current for row in rows])
        headers = ["Шкала/опросник", _format_scale_header(initial_date, "Исходно")]
        if final_mode:
            headers.append(_format_scale_header(current_date, "Повторно"))
        for cell, value in zip(table.rows[0].cells, headers, strict=True):
            set_cell_text(cell, value, style=STYLE_TABLE_HEADER, alignment=WD_ALIGN_PARAGRAPH.LEFT, keep_with_next=True)
        for cell in table.rows[0].cells[1:]:
            _compact_header_cell(cell)
        mark_header_row(table.rows[0])

        for table_row, scale_row in zip(table.rows[1:], rows, strict=True):
            set_cell_text(table_row.cells[0], scale_row.name, style=STYLE_TABLE, alignment=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_text(
                table_row.cells[1],
                _format_scale_value(scale_row.initial, initial_date),
                style=STYLE_TABLE,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
            )
            if final_mode:
                set_cell_text(
                    table_row.cells[2],
                    _format_scale_value(scale_row.current, current_date),
                    style=STYLE_TABLE,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER,
                )

    def _render_procedure_table(self) -> None:
        procedures = self.episode.procedures
        table = self.document.add_table(rows=max(1, len(procedures)) + 1, cols=len(PROCEDURE_WIDTHS))
        configure_table(table, PROCEDURE_WIDTHS)
        headers = (
            "Реабилитационные процедуры",
            "Ответственный специалист",
            "Количество",
            "Продолжительность в мин.",
            "Кратность",
        )
        for cell, value in zip(table.rows[0].cells, headers, strict=True):
            set_cell_text(cell, value, style=STYLE_TABLE_HEADER, alignment=WD_ALIGN_PARAGRAPH.CENTER, keep_with_next=True)
        _compact_header_cell(table.rows[0].cells[2])
        mark_header_row(table.rows[0])

        if not procedures:
            empty = table.rows[1].cells[0].merge(table.rows[1].cells[-1])
            set_cell_text(empty, "Мероприятия не представлены", style=STYLE_WARNING, alignment=WD_ALIGN_PARAGRAPH.LEFT)
            return

        for row, procedure in zip(table.rows[1:], procedures, strict=True):
            procedure_name = " ".join(item for item in (procedure.code, procedure.name) if item).strip()
            values = (
                procedure_name,
                procedure.specialist,
                "" if procedure.actual_count is None else procedure.actual_count,
                "" if procedure.duration_minutes is None else procedure.duration_minutes,
                procedure.frequency,
            )
            for cell, value in zip(row.cells, values, strict=True):
                set_cell_text(
                    cell,
                    value,
                    style=STYLE_TABLE,
                    alignment=WD_ALIGN_PARAGRAPH.LEFT,
                    vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP,
                )

    def _render_signature_table(self) -> None:
        rows = self._resolved_signatories()
        table = self.document.add_table(rows=len(rows) + 1, cols=len(SIGNATURE_WIDTHS))
        configure_table(table, SIGNATURE_WIDTHS)
        for cell, value in zip(table.rows[0].cells, ("Специалист МДРК", "ФИО", "Подпись"), strict=True):
            set_cell_text(cell, value, style=STYLE_TABLE_HEADER, alignment=WD_ALIGN_PARAGRAPH.CENTER, keep_with_next=True)
        mark_header_row(table.rows[0])

        for table_row, signatory in zip(table.rows[1:], rows, strict=True):
            set_cell_text(table_row.cells[0], signatory.display_role, style=STYLE_TABLE, alignment=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_text(table_row.cells[1], signatory.full_name, style=STYLE_TABLE, alignment=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_text(table_row.cells[2], "", style=STYLE_TABLE, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    def _resolved_signatories(self) -> tuple[SignatoryRow, ...]:
        rows = list(SIGNATORY_ROSTER)
        positions = {_signatory_role_key(row): index for index, row in enumerate(rows)}
        extras: list[SignatoryRow] = []
        for supplied in self.signatories:
            key = _signatory_role_key(supplied)
            if key not in positions:
                extras.append(supplied)
                continue
            if supplied.full_name.strip():
                index = positions[key]
                rows[index] = SignatoryRow(rows[index].role, supplied.full_name.strip())

        consultant_index = positions["consultants"]
        return tuple((*rows[:consultant_index], *extras, *rows[consultant_index:]))

    def _add_section_heading(self, number: int, title: str) -> None:
        self.document.add_paragraph(f"{number}. {title}", style=STYLE_SECTION)

    def _add_section_value(self, number: int, title: str, value: str) -> None:
        paragraph = self.document.add_paragraph(style=STYLE_SECTION)
        paragraph.add_run(f"{number}. {title}: ")
        paragraph.add_run(value or "")

    def _add_labeled(self, label: str, value: object) -> None:
        paragraph = self.document.add_paragraph(style=STYLE_BODY)
        if value is None or value == "":
            paragraph.paragraph_format.keep_with_next = True
        label_run = paragraph.add_run(label)
        label_run.style = STYLE_LABEL
        paragraph.add_run("" if value is None else str(value))

    def _add_blank_paragraph(self, *, keep_with_next: bool = False) -> None:
        paragraph = self.document.add_paragraph(style=STYLE_BODY)
        paragraph.paragraph_format.keep_with_next = keep_with_next

    def _add_original_plan_heading(self, text: str) -> None:
        paragraph = self.document.add_paragraph(style=STYLE_BODY)
        run = paragraph.add_run(text)
        run.bold = True
        run.underline = True

    def _add_original_birth_line(self, birth_date: date | None) -> None:
        paragraph = self.document.add_paragraph(style=STYLE_BODY)
        label = paragraph.add_run("Дата рождения: ")
        label.bold = True
        if birth_date is None:
            return
        paragraph.add_run(f"«{birth_date.day:02d}» ")
        month = paragraph.add_run(_MONTHS[birth_date.month])
        month.underline = True
        paragraph.add_run(" ")
        year = paragraph.add_run(str(birth_date.year))
        year.underline = True
        paragraph.add_run("г.")
        if self.snapshot.meeting_at is None:
            return
        age = _patient_age(birth_date, self.snapshot.meeting_at)
        opening = paragraph.add_run(" (")
        opening.bold = True
        age_run = paragraph.add_run(f"{age} {_age_word(age)}")
        age_run.bold = True
        age_run.underline = True
        closing = paragraph.add_run(")")
        closing.bold = True

    def _add_text_or_missing(self, value: str, missing_text: str) -> None:
        if value.strip():
            self._add_multiline(value)
        else:
            self.document.add_paragraph(missing_text, style=STYLE_WARNING)

    def _add_multiline(self, value: str) -> None:
        lines = [line.strip() for line in value.splitlines() if line.strip()]
        for line in lines or [value.strip()]:
            self.document.add_paragraph(line, style=STYLE_BODY)


def _clear_body(document: DocxDocument) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _set_output_metadata(document: DocxDocument, kind: MdrkKind) -> None:
    properties = document.core_properties
    properties.author = "MDRK Builder"
    properties.last_modified_by = "MDRK Builder"
    properties.title = "МДРК" if kind is MdrkKind.INITIAL else "МДРК итоговая"
    properties.subject = "Локально сформированный редактируемый документ"
    properties.comments = ""
    properties.keywords = ""


def _format_meeting(value: datetime | None) -> str:
    if value is None:
        return "Дата и время заседания не указаны"
    return (
        f'"{value.day:02d}" {_MONTHS[value.month]} {value.year} г. '
        f"время: {value.hour:02d} час. {value.minute:02d} мин."
    )


def _format_short_date(value: datetime | None) -> str:
    return value.strftime("%d.%m") if value is not None else ""


def _format_short_datetime(value: datetime | None) -> str:
    return value.strftime("%d.%m.%Y %H:%M") if value is not None else ""


def _patient_age(birth_date: date, meeting_at: datetime) -> int:
    return meeting_at.year - birth_date.year - (
        (meeting_at.month, meeting_at.day) < (birth_date.month, birth_date.day)
    )


def _age_word(age: int) -> str:
    if age % 100 in range(11, 15):
        return "лет"
    if age % 10 == 1:
        return "год"
    if age % 10 in range(2, 5):
        return "года"
    return "лет"


def _format_patient_birth(birth_date: date | None, meeting_at: datetime | None) -> str:
    if birth_date is None:
        return ""
    result = f"«{birth_date.day:02d}» {_MONTHS[birth_date.month]} {birth_date.year}г."
    if meeting_at is None:
        return result
    age = _patient_age(birth_date, meeting_at)
    return f"{result} ({age} {_age_word(age)})"


def _format_original_scale_datetime(value: datetime | None) -> str:
    if value is None:
        return ""
    return f"{value.day:02d} {_MONTHS[value.month]} {value.year}\n{value:%H:%M}"


def _first_scale_datetime(
    rows: Sequence[ScaleRow],
    kind: MdrkKind,
) -> datetime | None:
    for row in rows:
        measurement = row.current if kind is MdrkKind.FINAL else row.initial
        if measurement is not None and measurement.measured_at is not None:
            return measurement.measured_at
    return None


def _specialist_result_heading(
    role: SpecialistRole,
    source_datetime: datetime | None,
) -> str:
    if role in {SpecialistRole.FRM, SpecialistRole.NEUROLOGIST}:
        specialist = "врача физической и реабилитационной медицины"
        date_text = (
            f"{source_datetime.day:02d} {_MONTHS[source_datetime.month]} "
            f"{source_datetime.year} {source_datetime:%H:%M}"
            if source_datetime is not None
            else "дата, время"
        )
    else:
        specialist = {
            SpecialistRole.PHYSICAL_THERAPIST: "специалиста по физической реабилитации",
            SpecialistRole.LOGOPEDIST: "медицинского логопеда",
            SpecialistRole.NEUROPSYCHOLOGIST: "медицинского психолога/нейропсихолога",
            SpecialistRole.PATHOPSYCHOLOGIST: "медицинского психолога/патопсихолога",
            SpecialistRole.OCCUPATIONAL_THERAPIST: "специалиста по эргореабилитации",
            SpecialistRole.OTHER: "консультанта",
        }[role]
        date_text = (
            source_datetime.strftime("%d.%m.%Y, %H:%M")
            if source_datetime is not None
            else "дата, время"
        )
    return f"Результат осмотра {specialist} ({date_text}):"


def _deduplicate_physician_scale_rows(
    rows: Sequence[ScaleRow],
) -> tuple[ScaleRow, ...]:
    selected: dict[str, ScaleRow] = {}
    for row in rows:
        key = " ".join(row.name.casefold().split())
        existing = selected.get(key)
        if existing is None:
            selected[key] = row
            continue
        existing_measurement = existing.current or existing.initial
        candidate_measurement = row.current or row.initial
        if (
            candidate_measurement is not None
            and candidate_measurement.measured_at is not None
            and (
                existing_measurement is None
                or existing_measurement.measured_at is None
                or candidate_measurement.measured_at > existing_measurement.measured_at
            )
        ):
            selected[key] = row

    priority_tokens = (
        "ривермид",
        "рэнкин",
        "nrs 2002",
        "скф",
        "реабилитационной маршрутизации",
        "бартел",
    )

    def sort_key(row: ScaleRow) -> tuple[int, str]:
        normalized = row.name.casefold()
        priority = next(
            (
                index
                for index, token in enumerate(priority_tokens)
                if token in normalized
            ),
            len(priority_tokens),
        )
        return priority, normalized

    return tuple(sorted(selected.values(), key=sort_key))


def _mcf_category(code: str) -> str:
    normalized = code.strip().casefold().replace("е", "e")
    if normalized.startswith("b"):
        return "Структура/функция"
    if normalized.startswith("s"):
        return "Структуры организма"
    if normalized.startswith("d"):
        return "Активность/участие"
    if normalized.startswith("e"):
        return "Факторы окружающей среды"
    if normalized.startswith("pf"):
        return "Персональные факторы"
    return "Другие домены"


def _is_personal_factor(code: str) -> bool:
    return code.strip().casefold().replace(" ", "").startswith("pf")


def _is_environment_factor(code: str) -> bool:
    return code.strip().casefold().replace("е", "e").startswith("e")


def _mcf_responsible(domain: IcfDomain) -> str:
    if domain.note.strip():
        return domain.note.strip()
    if domain.specialist is SpecialistRole.OTHER:
        return ""
    return domain.specialist.display_name


def _shade_initial_qualifier(
    cells: Sequence[_Cell], qualifier: IcfQualifier | None
) -> None:
    """Shade the cumulative path from zero to the initial ICF qualifier."""

    if qualifier is None:
        return
    zero_column = 6
    value = qualifier.value
    if qualifier.facilitator:
        qualifier_columns = range(zero_column - value, zero_column + 1)
    else:
        qualifier_columns = range(zero_column, zero_column + value + 1)
    for column in qualifier_columns:
        set_cell_shading(cells[column], MCF_QUALIFIER_FILL)


def _compact_header_cell(cell: _Cell) -> None:
    set_cell_horizontal_margins(cell, left=0, right=0)
    set_cell_no_wrap(cell)


def _common_measurement_datetime(
    measurements: Sequence[ScaleMeasurement | None],
) -> datetime | None:
    dates = {item.measured_at for item in measurements if item is not None and item.measured_at is not None}
    return dates.pop() if len(dates) == 1 else None


def _format_scale_header(value: datetime | None, fallback: str) -> str:
    return _format_short_datetime(value) if value is not None else fallback


def _format_scale_value(
    measurement: ScaleMeasurement | None,
    shared_datetime: datetime | None,
) -> str:
    if measurement is None:
        return ""
    if measurement.measured_at is None or measurement.measured_at == shared_datetime:
        return measurement.value
    return f"{measurement.value}\n{_format_short_datetime(measurement.measured_at)}"


def _split_tasks(value: str) -> list[str]:
    lines = [line.strip() for line in value.splitlines() if line.strip()]
    cleaned = [re.sub(r"^(?:[-•–—]|\d+[.)])\s*", "", line).strip() for line in lines]
    return [line for line in cleaned if line]
