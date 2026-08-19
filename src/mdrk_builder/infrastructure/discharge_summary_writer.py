from __future__ import annotations

from collections.abc import Sequence
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from mdrk_builder.domain import (
    DischargeScaleRow,
    DischargeSummaryDraft,
    ReviewIssue,
)
from mdrk_builder.infrastructure.clinical_tables import (
    render_completed_program,
    render_final_icf_profile,
)
from mdrk_builder.infrastructure.discharge_summary_template import (
    ACKNOWLEDGEMENT_TEMPLATE,
    SIGNING_DATE_TEMPLATE,
    discharge_summary_template_path,
)
from mdrk_builder.infrastructure.docx_layout import (
    configure_table,
    mark_header_row,
    set_cell_text,
)
from mdrk_builder.infrastructure.docx_output import (
    resolve_docx_output_path,
    save_sanitized_docx_atomically,
)
from mdrk_builder.infrastructure.docx_template import (
    STYLE_BODY,
    STYLE_LABEL,
    STYLE_MCF_CODE,
    STYLE_SECTION,
    STYLE_TABLE,
    STYLE_TABLE_HEADER,
    STYLE_TITLE,
    STYLE_WARNING,
    configure_mdrk_styles,
)


BODY_FONT_SIZE_PT = 10
SCALE_WIDTHS = (1900, 5900, 1545)


class DischargeSummaryGenerationBlockedError(ValueError):
    def __init__(self, issues: Sequence[ReviewIssue]) -> None:
        self.issues = tuple(issues)
        details = "; ".join(issue.message for issue in self.issues)
        super().__init__(f"Формирование выписного эпикриза заблокировано: {details}")


def write_discharge_summary_docx(
    draft: DischargeSummaryDraft,
    output_path: Path,
    *,
    template_path: Path | None = None,
    ignore_issues: bool = False,
) -> Path:
    blocking = list(draft.blocking_issues())
    if blocking and not ignore_issues:
        raise DischargeSummaryGenerationBlockedError(blocking)

    template = (template_path or discharge_summary_template_path()).resolve()
    if not template.is_file():
        raise FileNotFoundError(
            f"Шаблон выписного эпикриза не найден: {template}. "
            "Запустите tools/build_discharge_summary_template.py."
        )
    output = resolve_docx_output_path(
        output_path,
        template_path=template,
        source_paths=draft.immutable_sources(),
    )

    document = Document(template)
    _retain_letterhead(document)
    configure_mdrk_styles(document)
    _configure_body_styles(document)
    _set_footer(document, draft)
    _set_metadata(document)
    _DischargeSummaryRenderer(document, draft).render()
    return save_sanitized_docx_atomically(document, output)


class _DischargeSummaryRenderer:
    def __init__(self, document: DocxDocument, draft: DischargeSummaryDraft) -> None:
        self.document = document
        self.draft = draft

    def render(self) -> None:
        self._title_and_header()
        self._diagnoses()
        self._admission_state()
        self._team_results()
        self._medical_results()
        self._treatment()
        self._discharge_state()
        self._closing()

    def _title_and_header(self) -> None:
        title = self.document.add_paragraph(style=STYLE_TITLE)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("ВЫПИСНОЙ ЭПИКРИЗ")
        run.bold = True
        self._multiline(self.draft.header_text)

    def _diagnoses(self) -> None:
        self._section("Заключительный клинический диагноз")
        self._multiline(self.draft.clinical_diagnosis)
        self._section("Реабилитационный диагноз")
        if self.draft.icf_domains:
            render_final_icf_profile(
                self.document,
                self.draft.icf_domains,
                repeat_missing_final=False,
            )
        else:
            self._multiline("")

    def _admission_state(self) -> None:
        self._section("Состояние при поступлении")
        self._labeled("Жалобы", self.draft.complaints)
        self._labeled("Анамнез заболевания", self.draft.disease_history)
        self._labeled("Анамнез жизни", self.draft.life_history)
        self._labeled(
            "Пациентом представлены необходимые для госпитализации документы",
            self.draft.provided_documents,
        )
        self._labeled("Физикальное обследование", self.draft.physical_exam)
        self._labeled("Неврологический статус", self.draft.neurological_status)
        self._labeled("Локальный статус", self.draft.local_status)
        self._section("Шкалы при поступлении")
        self._scale_table(self.draft.admission_scale_rows)

    def _team_results(self) -> None:
        self._section("Проведенные обследования, лечение, медицинская реабилитация")
        for finding in self.draft.team_findings:
            self._labeled(
                f"Заключение: {finding.role.display_name}",
                finding.conclusion,
            )
        self._labeled("Консультации узких специалистов", self.draft.other_consultations)

    def _medical_results(self) -> None:
        self._section("Результаты медицинского обследования")
        self._labeled("Лабораторные исследования", self.draft.laboratory_results)
        self._labeled("Инструментальные исследования", self.draft.instrumental_results)

    def _treatment(self) -> None:
        self._manual_block(
            "Применение лекарственных препаратов (включая химиотерапию, "
            "вакцинацию), медицинских изделий, лечебного питания",
            self.draft.medications,
            blank_lines=3,
        )
        self._labeled("Двигательный режим", self.draft.movement_regimen)
        self._labeled("Диета", self.draft.diet)
        self._manual_block(
            "Трансфузии (переливания) донорской крови и (или) ее компонентов",
            self.draft.transfusions,
            blank_lines=2,
        )
        self._labeled(
            "Оперативные вмешательства (операции), включая сведения об "
            "анестезиологическом пособии",
            self.draft.operations,
        )
        self._section("Медицинские вмешательства")
        self._labeled("Проведенная программа медицинской реабилитации", "")
        render_completed_program(self.document, self.draft.completed_procedures)
        self._labeled("Дополнительные сведения", self.draft.additional_information)

    def _discharge_state(self) -> None:
        self._section("Шкалы при выписке")
        self._scale_table(self.draft.discharge_scale_rows)
        self._section("Состояние при выписке")
        self._multiline(self.draft.discharge_condition)
        self._manual_block(
            "Неврологический статус",
            self.draft.discharge_neurological_status,
            blank_lines=2,
        )
        self._labeled(
            "Факторы риска проведения реабилитационных мероприятий",
            self.draft.risks,
        )
        self._labeled(
            "Факторы, ограничивающие проведение реабилитационных мероприятий",
            self.draft.limitations,
        )
        self._labeled(
            "Реабилитационный потенциал",
            self.draft.rehabilitation_potential,
        )
        self._labeled(
            "Цель, поставленная на этап медицинской реабилитации",
            self.draft.goal_result,
        )
        self._labeled(
            "Трудоспособность, листок нетрудоспособности",
            self.draft.work_capacity,
        )
        self._labeled("Лучевая нагрузка", self.draft.radiation_exposure)
        self._labeled("Рекомендации", self.draft.recommendations)

    def _closing(self) -> None:
        self._multiline(ACKNOWLEDGEMENT_TEMPLATE)
        self._multiline(self.draft.signatures)
        self._multiline(SIGNING_DATE_TEMPLATE)

    def _section(self, title: str) -> None:
        paragraph = self.document.add_paragraph(style=STYLE_SECTION)
        paragraph.paragraph_format.keep_with_next = True
        run = paragraph.add_run(title)
        run.bold = True

    def _labeled(self, label: str, value: str) -> None:
        paragraph = self.document.add_paragraph(style=STYLE_BODY)
        paragraph.paragraph_format.keep_with_next = not value.strip()
        run = paragraph.add_run(f"{label}: ")
        run.bold = True
        if value.strip():
            lines = value.splitlines()
            paragraph.add_run(lines[0].strip())
            for line in lines[1:]:
                if line.strip():
                    paragraph.add_run().add_break()
                    paragraph.add_run(line.strip())

    def _manual_block(self, label: str, value: str, *, blank_lines: int) -> None:
        self._labeled(label, value)
        if value.strip():
            return
        for _ in range(blank_lines):
            self.document.add_paragraph("", style=STYLE_BODY)

    def _multiline(self, value: str) -> None:
        lines = [line.strip() for line in value.splitlines() if line.strip()]
        if not lines:
            self.document.add_paragraph("", style=STYLE_BODY)
            return
        for line in lines:
            self.document.add_paragraph(line, style=STYLE_BODY)

    def _scale_table(self, rows: Sequence[DischargeScaleRow]) -> None:
        table = self.document.add_table(rows=max(1, len(rows)) + 1, cols=3)
        configure_table(table, SCALE_WIDTHS)
        for cell, text in zip(
            table.rows[0].cells,
            ("Специалист", "Шкала/опросник", "Результат"),
            strict=True,
        ):
            set_cell_text(
                cell,
                text,
                style=STYLE_TABLE_HEADER,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                keep_with_next=True,
            )
        mark_header_row(table.rows[0])
        if not rows:
            empty = table.rows[1].cells[0].merge(table.rows[1].cells[-1])
            set_cell_text(
                empty,
                "",
                style=STYLE_TABLE,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
            )
            return
        for table_row, scale_row in zip(table.rows[1:], rows, strict=True):
            values = (
                scale_row.role.display_name,
                scale_row.name,
                scale_row.value,
            )
            for cell, value in zip(table_row.cells, values, strict=True):
                set_cell_text(
                    cell,
                    value,
                    style=STYLE_TABLE,
                    alignment=WD_ALIGN_PARAGRAPH.LEFT,
                    vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP,
                )


def _retain_letterhead(document: DocxDocument) -> None:
    body = document._element.body
    letterhead = next((child for child in body if child.tag == qn("w:tbl")), None)
    if letterhead is None:
        raise ValueError("discharge-summary template has no letterhead table")
    for child in list(body):
        if child is not letterhead and child.tag != qn("w:sectPr"):
            body.remove(child)


def _configure_body_styles(document: DocxDocument) -> None:
    for style_name in ("Normal", STYLE_BODY, STYLE_SECTION, STYLE_TITLE, STYLE_WARNING):
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(BODY_FONT_SIZE_PT)
        fonts = style.element.get_or_add_rPr().get_or_add_rFonts()
        for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
            fonts.set(qn(f"w:{attribute}"), "Times New Roman")
    for style_name in (STYLE_TABLE, STYLE_TABLE_HEADER, STYLE_MCF_CODE):
        document.styles[style_name].font.size = Pt(BODY_FONT_SIZE_PT)
    document.styles[STYLE_LABEL].font.size = Pt(BODY_FONT_SIZE_PT)


def _display_record_number(value: str) -> str:
    compact = value.replace("№", "").replace(" ", "").strip()
    if compact and not compact.casefold().startswith("скп"):
        compact = f"СКП{compact}"
    return compact


def _set_footer(document: DocxDocument, draft: DischargeSummaryDraft) -> None:
    value = (
        f"{draft.identity.full_name}, МКП №"
        f"{_display_record_number(draft.identity.medical_record_number)}"
    )
    for section in document.sections:
        footer = section.footer
        for table in list(footer.tables):
            table._element.getparent().remove(table._element)
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.clear()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(value)
        run.font.name = "Times New Roman"
        run.font.size = Pt(BODY_FONT_SIZE_PT)
        run.font.color.rgb = RGBColor(31, 78, 121)
        fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
        for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
            fonts.set(qn(f"w:{attribute}"), "Times New Roman")
        for extra in footer.paragraphs[1:]:
            extra._element.getparent().remove(extra._element)


def _set_metadata(document: DocxDocument) -> None:
    properties = document.core_properties
    properties.author = "MDRK Builder"
    properties.last_modified_by = "MDRK Builder"
    properties.title = "Выписной эпикриз"
    properties.subject = "Локально сформированный редактируемый документ"
    properties.comments = ""
    properties.keywords = ""
