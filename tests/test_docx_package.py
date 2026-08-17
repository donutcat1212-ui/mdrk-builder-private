from __future__ import annotations

from zipfile import ZipFile

import pytest
from docx import Document
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from lxml import etree

from mdrk_builder.infrastructure import docx_output
from mdrk_builder.infrastructure.document_metadata import (
    GENERATED_DOCUMENT_IDENTIFIER,
)


WORDPROCESSINGML_NAMESPACE = (
    "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
)


def _revision_element(tag: str, text: str, *, deleted_text: bool = False):
    text_tag = "delText" if deleted_text else "t"
    return parse_xml(
        f'<w:{tag} xmlns:w="{WORDPROCESSINGML_NAMESPACE}">'
        f"<w:r><w:{text_tag}>{text}</w:{text_tag}></w:r>"
        f"</w:{tag}>"
    )


def test_sanitizer_accepts_inserted_and_moved_to_content(tmp_path) -> None:
    document = Document()
    paragraph = document.add_paragraph("BEFORE ")
    commented_run = paragraph.add_run("COMMENTED ")
    commented_run._r.get_or_add_rPr().append(OxmlElement("w:vanish"))
    document.add_comment(
        commented_run,
        text="HIDDEN COMMENT",
        author="Sample Author",
    )
    paragraph._p.set(qn("w:rsidR"), "00112233")
    paragraph._p.append(_revision_element("ins", "INSERTED "))
    paragraph._p.append(_revision_element("del", "DELETED ", deleted_text=True))
    paragraph._p.append(_revision_element("moveFrom", "MOVED_FROM "))
    paragraph._p.append(_revision_element("moveTo", "MOVED_TO "))
    paragraph.add_run("AFTER")
    output = tmp_path / "tracked-changes.docx"

    docx_output.save_sanitized_docx_atomically(document, output)

    reopened = Document(output)
    assert reopened.paragraphs[0].text == (
        "BEFORE COMMENTED INSERTED MOVED_TO AFTER"
    )
    assert reopened.core_properties.identifier == GENERATED_DOCUMENT_IDENTIFIER
    with ZipFile(output) as package:
        names = package.namelist()
        xml_parts = [
            package.read(name)
            for name in names
            if name.endswith((".xml", ".rels"))
        ]
    roots = [etree.fromstring(data) for data in xml_parts]
    local_names = {
        etree.QName(element).localname for root in roots for element in root.iter()
    }
    attribute_names = {
        etree.QName(attribute).localname
        for root in roots
        for element in root.iter()
        for attribute in element.attrib
    }
    assert not any(name.casefold().startswith("word/comments") for name in names)
    assert "HIDDEN COMMENT" not in b"\n".join(xml_parts).decode(
        "utf-8", errors="ignore"
    )
    assert local_names.isdisjoint(
        {
            "commentRangeEnd",
            "commentRangeStart",
            "commentReference",
            "del",
            "ins",
            "moveFrom",
            "moveTo",
            "rsid",
            "rsidRoot",
            "rsids",
            "vanish",
        }
    )
    assert not any(name.casefold().startswith("rsid") for name in attribute_names)


def test_sanitizer_failure_preserves_existing_output(tmp_path, monkeypatch) -> None:
    output = tmp_path / "existing.docx"
    original = b"existing output must survive"
    output.write_bytes(original)
    document = Document()
    document.add_paragraph("replacement")

    def fail_sanitization(_path) -> None:
        raise RuntimeError("simulated sanitizer failure")

    monkeypatch.setattr(docx_output, "sanitize_docx_package", fail_sanitization)

    with pytest.raises(RuntimeError, match="simulated sanitizer failure"):
        docx_output.save_sanitized_docx_atomically(document, output)

    assert output.read_bytes() == original
    assert not list(tmp_path.glob(".existing-*.docx"))
