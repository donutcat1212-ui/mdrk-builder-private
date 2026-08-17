from datetime import datetime
from pathlib import Path

from docx import Document

from mdrk_builder.application.discharge_defaults import RECOMMENDATIONS_TEMPLATE
from mdrk_builder.application.discharge_source_selection import (
    source_scan_for_episode,
)
from mdrk_builder.application.discharge_summary import scan_discharge_summary
from mdrk_builder.application.episode_identity import DischargeEpisodeKey
from mdrk_builder.application.source_scan import ScannedDocument, SourceScanResult
from mdrk_builder.domain import SpecialistRole
from mdrk_builder.infrastructure.classifier import DocumentClassification
from mdrk_builder.infrastructure.ooxml_reader import (
    BodyItem,
    ParsedCell,
    ParsedDocument,
    ParsedRow,
    ParsedTable,
)
from mdrk_builder.infrastructure.discharge_summary_writer import (
    write_discharge_summary_docx,
)


EPISODE_ROOT = Path("/episode")


def _write_document(path, lines: tuple[str, ...]) -> None:
    document = Document()
    for line in lines:
        document.add_paragraph(line)
    document.save(path)


def _scanned_source(
    name: str,
    *lines: str,
    document_type: str = "follow_up",
    tables: tuple[ParsedTable, ...] = (),
) -> ScannedDocument:
    path = EPISODE_ROOT / name
    document = ParsedDocument(
        source_path=path,
        normalized_path=path,
        paragraphs=list(lines),
        tables=list(tables),
        body_items=[
            *(BodyItem("paragraph", index) for index in range(len(lines))),
            *(BodyItem("table", index) for index in range(len(tables))),
        ],
    )
    return ScannedDocument(
        document=document,
        classification=DocumentClassification(
            SpecialistRole.NEUROPSYCHOLOGIST,
            document_type,
        ),
    )


def _row(values: dict[int, str], logical_cols: int) -> ParsedRow:
    return ParsedRow(
        tuple(
            ParsedCell(column, 1, value)
            for column, value in sorted(values.items())
        ),
        logical_cols,
    )


def _episode_source_scan(*documents: ScannedDocument) -> SourceScanResult:
    return SourceScanResult(
        source_files=tuple(item.document.source_path for item in documents),
        documents=documents,
        failures=(),
        root=EPISODE_ROOT,
    )


def _episode_key() -> DischargeEpisodeKey:
    return DischargeEpisodeKey(
        normalized_full_name="альфа бета гамма",
        medical_record_number="5906/26",
        admission_at=datetime(2026, 8, 10, 10),
        discharge_at=datetime(2026, 8, 17, 12),
        episode_root=EPISODE_ROOT,
    )


def _primary_lines(
    record: str = "СКП5906/26",
    *,
    full_name: str = "Пациент Тестовый Пример",
    admission: str = "10.08.2026 10:00",
) -> tuple[str, ...]:
    return (
        "ПЕРВИЧНЫЙ ОСМОТР НЕВРОЛОГА",
        f"ФИО пациента: {full_name}",
        f"Номер ИБ: {record}",
        f"Дата и время поступления: {admission}",
        "Жалобы: PRIMARY COMPLAINTS",
        "Анамнез заболевания: PRIMARY DISEASE HISTORY",
        "Анамнез жизни: PRIMARY LIFE HISTORY",
        "Пациентом представлены необходимые для госпитализации документы: PRIMARY DOCS",
        "Физикальное обследование: PRIMARY PHYSICAL",
        "Неврологический статус: PRIMARY NEURO",
        "Локальный статус: PRIMARY LOCAL",
        "Шкалы при поступлении:",
        "Заключительный клинический диагноз: PRIMARY DIAGNOSIS",
        "Факторы риска проведения реабилитационных мероприятий: PRIMARY RISKS",
        "Факторы, ограничивающие проведение реабилитационных мероприятий: PRIMARY LIMITS",
        "Двигательный режим: палатный",
        "Диета: стол № 9",
        "Лечащий врач, врач-невролог",
    )


def _discharge_lines(
    record: str = "СКП5906/26",
    radiation: str | None = "4,2 мЗв",
    *,
    full_name: str = "Пациент Тестовый Пример",
    admission: str = "10.08.2026 10:00",
) -> tuple[str, ...]:
    values = [
        "Выписной эпикриз",
        "Отделение медицинской реабилитации",
        "Сведения о пациенте",
        f"ФИО пациента: {full_name}",
        f"Номер медицинской карты пациента №{record}",
        f"Дата и время поступления: {admission}",
        "Дата и время выписки: 17.08.2026 12:00",
        "Заключительный клинический диагноз: DISCHARGE DIAGNOSIS",
        "Лабораторные исследования: CURRENT LAB",
        "Инструментальные исследования: CURRENT INSTRUMENTAL",
        "Консультация оториноларинголога 13.08.2026 09:00",
        "CURRENT ENT CONCLUSION",
        "Консультация хирурга 29.07.2028 09:00",
        "IMPOSSIBLE FUTURE CONSULTATION",
    ]
    if radiation is not None:
        values.append(f"Лучевая нагрузка - {radiation}")
    values.append("Лечащий врач")
    return tuple(values)


def test_discharge_scan_applies_explicit_field_authority_and_chronology(tmp_path) -> None:
    current = tmp_path / "невролог"
    current.mkdir()
    primary_path = current / "первичный осмотр невролога.docx"
    discharge_path = current / "выписной эпикриз.docx"
    old_discharge_path = tmp_path / "старый выписной эпикриз.docx"
    _write_document(primary_path, _primary_lines())
    _write_document(discharge_path, _discharge_lines())
    _write_document(
        old_discharge_path,
        _discharge_lines("СКП5799/26", "12,7 мЗв"),
    )

    draft = scan_discharge_summary(tmp_path)

    assert draft.discharge_source == discharge_path
    assert draft.primary_neurologist_source == primary_path
    assert draft.identity.medical_record_number == "СКП5906/26"
    assert draft.admission_datetime == datetime(2026, 8, 10, 10, 0)
    assert draft.discharge_datetime == datetime(2026, 8, 17, 12, 0)
    assert not hasattr(draft, "episode")
    assert set(draft.source_paths) == {
        primary_path.resolve(),
        discharge_path.resolve(),
        old_discharge_path.resolve(),
    }
    assert old_discharge_path.resolve() in draft.immutable_sources()
    assert draft.clinical_diagnosis == "PRIMARY DIAGNOSIS"
    assert draft.complaints == "PRIMARY COMPLAINTS"
    assert draft.disease_history == "PRIMARY DISEASE HISTORY"
    assert draft.laboratory_results == "CURRENT LAB"
    assert draft.instrumental_results == "CURRENT INSTRUMENTAL"
    assert "CURRENT ENT CONCLUSION" in draft.other_consultations
    assert "IMPOSSIBLE FUTURE CONSULTATION" not in draft.other_consultations
    assert draft.radiation_exposure == "4,2 мЗв"
    assert draft.medications == ""
    assert draft.transfusions == ""
    assert draft.final_mdrk_source is None
    assert draft.rehabilitation_potential == ""
    assert draft.goal_result == ""
    assert draft.recommendations == RECOMMENDATIONS_TEMPLATE
    assert draft.field_sources["clinical_diagnosis"] == primary_path
    assert draft.field_sources["radiation_exposure"] == discharge_path
    assert any(issue.code == "consultation_outside_episode" for issue in draft.issues)
    assert not any(
        issue.code == "identity_conflict_medical_record_number"
        for issue in draft.issues
    )


def test_discharge_scan_writes_reopens_and_ignores_its_output(tmp_path) -> None:
    primary_path = tmp_path / "первичный осмотр невролога.docx"
    discharge_path = tmp_path / "выписной эпикриз.docx"
    output_path = tmp_path / "готовый выписной эпикриз.docx"
    _write_document(primary_path, _primary_lines())
    _write_document(discharge_path, _discharge_lines())

    draft = scan_discharge_summary(tmp_path)
    created = write_discharge_summary_docx(draft, output_path)

    reopened = Document(created)
    assert "ВЫПИСНОЙ ЭПИКРИЗ" in "\n".join(
        paragraph.text for paragraph in reopened.paragraphs
    )
    rescanned = scan_discharge_summary(tmp_path)
    assert rescanned.discharge_source == discharge_path
    assert rescanned.primary_neurologist_source == primary_path
    assert rescanned.clinical_diagnosis == draft.clinical_diagnosis
    assert output_path.resolve() not in rescanned.immutable_sources()
    rewritten = write_discharge_summary_docx(rescanned, output_path)
    assert rewritten == output_path.resolve()
    assert Document(rewritten).paragraphs


def test_discharge_scan_uses_profile_primary_not_admission_department(
    tmp_path,
) -> None:
    primary_path = tmp_path / "первичный осмотр невролога.docx"
    admission_path = tmp_path / "первичный осмотр в приёмном.docx"
    _write_document(primary_path, _primary_lines())
    _write_document(
        admission_path,
        (
            "ПЕРВИЧНЫЙ ОСМОТР НЕВРОЛОГА ПРИЁМНОГО ОТДЕЛЕНИЯ",
            *_primary_lines()[1:],
        ),
    )
    _write_document(tmp_path / "выписной эпикриз.docx", _discharge_lines())

    draft = scan_discharge_summary(tmp_path)

    assert draft.primary_neurologist_source == primary_path
    assert draft.clinical_diagnosis == "PRIMARY DIAGNOSIS"
    assert not any(
        issue.code == "episode_source_selection_ambiguous"
        for issue in draft.blocking_issues()
    )


def test_profile_primary_can_mention_admission_department_in_history(
    tmp_path,
) -> None:
    primary_lines = tuple(
        (
            "Анамнез заболевания: Пациент переведён из приёмного "
            "отделения в профильное."
            if line.startswith("Анамнез заболевания")
            else line
        )
        for line in _primary_lines()
    )
    primary_path = tmp_path / "первичный осмотр невролога.docx"
    _write_document(primary_path, primary_lines)
    _write_document(tmp_path / "выписной эпикриз.docx", _discharge_lines())

    draft = scan_discharge_summary(tmp_path)

    assert draft.primary_neurologist_source == primary_path
    assert not any(
        issue.code == "primary_neurologist_source_missing"
        for issue in draft.blocking_issues()
    )


def test_discharge_scan_defaults_radiation_to_zero(tmp_path) -> None:
    _write_document(tmp_path / "первичный осмотр невролога.docx", _primary_lines())
    _write_document(
        tmp_path / "выписной эпикриз.docx",
        _discharge_lines(radiation=None),
    )

    draft = scan_discharge_summary(tmp_path)

    assert draft.radiation_exposure == "0 мЗв"
    assert "radiation_exposure" not in draft.field_sources


def test_discharge_scan_blocks_without_current_discharge_source(tmp_path) -> None:
    _write_document(tmp_path / "первичный осмотр невролога.docx", _primary_lines())

    draft = scan_discharge_summary(tmp_path)

    assert any(
        issue.code == "discharge_summary_source_missing"
        for issue in draft.blocking_issues()
    )


def test_discharge_scan_fails_closed_when_header_boundary_is_missing(tmp_path) -> None:
    _write_document(tmp_path / "первичный осмотр невролога.docx", _primary_lines())
    without_diagnosis = tuple(
        line
        for line in _discharge_lines()
        if not line.startswith("Заключительный клинический диагноз")
    )
    _write_document(tmp_path / "выписной эпикриз.docx", without_diagnosis)

    draft = scan_discharge_summary(tmp_path)

    assert draft.header_text == ""
    assert "header_text" not in draft.field_sources
    assert any(
        issue.code == "discharge_header_missing"
        for issue in draft.blocking_issues()
    )


def test_discharge_scan_blocks_without_discharge_datetime(tmp_path) -> None:
    _write_document(tmp_path / "первичный осмотр невролога.docx", _primary_lines())
    without_discharge_datetime = tuple(
        line
        for line in _discharge_lines()
        if not line.startswith("Дата и время выписки")
    )
    _write_document(
        tmp_path / "выписной эпикриз.docx",
        without_discharge_datetime,
    )

    draft = scan_discharge_summary(tmp_path)

    assert any(
        issue.code == "discharge_datetime_missing"
        for issue in draft.blocking_issues()
    )


def test_discharge_scan_never_pairs_conflicting_patients(tmp_path) -> None:
    _write_document(
        tmp_path / "первичный осмотр невролога.docx",
        _primary_lines(
            "СКП7777/26",
            full_name="Другой Пациент Тестовый",
            admission="11.08.2026 10:00",
        ),
    )
    discharge_path = tmp_path / "выписной эпикриз.docx"
    _write_document(discharge_path, _discharge_lines())

    draft = scan_discharge_summary(tmp_path)

    assert draft.discharge_source == discharge_path
    assert draft.primary_neurologist_source is None
    assert draft.clinical_diagnosis == ""
    blocking_codes = {issue.code for issue in draft.blocking_issues()}
    assert "episode_source_identity_conflict" in blocking_codes
    assert "primary_neurologist_source_missing" in blocking_codes


def test_discharge_scan_does_not_pair_sources_without_patient_identity(tmp_path) -> None:
    _write_document(
        tmp_path / "первичный осмотр невролога.docx",
        _primary_lines("", full_name="нет данных"),
    )
    _write_document(
        tmp_path / "выписной эпикриз.docx",
        _discharge_lines("", full_name="нет данных"),
    )

    draft = scan_discharge_summary(tmp_path)

    assert draft.primary_neurologist_source is None
    blocking_codes = {issue.code for issue in draft.blocking_issues()}
    assert "episode_source_identity_insufficient" in blocking_codes
    assert "primary_neurologist_source_missing" in blocking_codes


def test_discharge_scan_does_not_pair_same_name_without_episode_anchor(
    tmp_path,
) -> None:
    _write_document(
        tmp_path / "первичный осмотр невролога.docx",
        _primary_lines("", admission=""),
    )
    _write_document(
        tmp_path / "выписной эпикриз.docx",
        _discharge_lines("", admission=""),
    )

    draft = scan_discharge_summary(tmp_path)

    assert draft.primary_neurologist_source is None
    blocking_codes = {issue.code for issue in draft.blocking_issues()}
    assert "episode_source_identity_insufficient" in blocking_codes
    assert "primary_neurologist_source_missing" in blocking_codes


def test_episode_source_projection_excludes_pre_admission_document() -> None:
    old_source = _scanned_source(
        "old-consultation.docx",
        "Осмотр нейропсихолога 09.08.2026 09:00",
        "ФИО пациента: АЛЬФА БЕТА ГАММА",
        "Номер ИБ: СКП5906/26",
    )
    current_source = _scanned_source(
        "current-consultation.docx",
        "Осмотр нейропсихолога 11.08.2026 09:00",
        "ФИО пациента: АЛЬФА БЕТА ГАММА",
        "Номер ИБ: СКП5906/26",
    )
    issues = []

    projected = source_scan_for_episode(
        _episode_source_scan(old_source, current_source),
        _episode_key(),
        issues=issues,
    )

    assert projected.documents == (current_source,)
    assert [issue.code for issue in issues] == [
        "episode_source_before_admission_excluded"
    ]
    assert issues[0].source == old_source.document.source_path


def test_episode_source_projection_rejects_undated_unidentified_root_match() -> None:
    unsupported_source = _scanned_source(
        "unsupported.docx",
        "Заключение специалиста без идентификаторов и даты",
    )
    dated_source = _scanned_source(
        "dated.docx",
        "Осмотр нейропсихолога 11.08.2026 09:00",
    )
    issues = []

    projected = source_scan_for_episode(
        _episode_source_scan(unsupported_source, dated_source),
        _episode_key(),
        issues=issues,
    )

    assert projected.documents == (dated_source,)
    assert [issue.code for issue in issues] == [
        "episode_source_identity_and_date_missing"
    ]
    assert issues[0].source == unsupported_source.document.source_path


def test_episode_source_projection_rejects_same_name_without_dated_anchor() -> None:
    same_name_only = _scanned_source(
        "same-name-only.docx",
        "Заключение специалиста без даты",
        "ФИО пациента: АЛЬФА БЕТА ГАММА",
    )
    issues = []

    projected = source_scan_for_episode(
        _episode_source_scan(same_name_only),
        _episode_key(),
        issues=issues,
    )

    assert projected.documents == ()
    assert [issue.code for issue in issues] == [
        "episode_source_identity_and_date_missing"
    ]
    assert issues[0].source == same_name_only.document.source_path


def test_episode_source_projection_excludes_post_discharge_document() -> None:
    future_source = _scanned_source(
        "future-consultation.docx",
        "Осмотр нейропсихолога 18.08.2026 09:00",
        "ФИО пациента: АЛЬФА БЕТА ГАММА",
        "Номер ИБ: СКП5906/26",
    )
    current_source = _scanned_source(
        "current-consultation.docx",
        "Осмотр нейропсихолога 17.08.2026 09:00",
        "ФИО пациента: АЛЬФА БЕТА ГАММА",
        "Номер ИБ: СКП5906/26",
    )
    issues = []

    projected = source_scan_for_episode(
        _episode_source_scan(future_source, current_source),
        _episode_key(),
        issues=issues,
    )

    assert projected.documents == (current_source,)
    assert [issue.code for issue in issues] == [
        "episode_source_after_discharge_excluded"
    ]
    assert issues[0].source == future_source.document.source_path


def test_episode_source_projection_requires_dated_assignment_sheet() -> None:
    identity_lines = (
        "ФИО пациента: АЛЬФА БЕТА ГАММА",
        "Номер ИБ: СКП5906/26",
        "Дата и время поступления: 10.08.2026 10:00",
    )
    undated_table = ParsedTable(
        (
            _row({0: "Назначения", 1: "время", 2: "кабинет", 3: "выполнено"}, 4),
            _row({0: "A13.23.011 Нейропсихологическая коррекция", 3: "+"}, 4),
        )
    )
    dated_table = ParsedTable(
        (
            _row({0: "Назначения", 1: "время", 2: "кабинет", 3: "11"}, 4),
            _row({0: "A13.23.011 Нейропсихологическая коррекция", 3: "+"}, 4),
        )
    )
    undated_source = _scanned_source(
        "undated-assignment.docx",
        *identity_lines,
        document_type="assignment_sheet",
        tables=(undated_table,),
    )
    dated_source = _scanned_source(
        "dated-assignment.docx",
        *identity_lines,
        document_type="assignment_sheet",
        tables=(dated_table,),
    )
    issues = []

    projected = source_scan_for_episode(
        _episode_source_scan(undated_source, dated_source),
        _episode_key(),
        issues=issues,
    )

    assert projected.documents == (dated_source,)
    assert [issue.code for issue in issues] == [
        "episode_assignment_sheet_date_missing"
    ]
    assert issues[0].source == undated_source.document.source_path
