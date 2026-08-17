from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document

from mdrk_builder.domain import (
    DischargeScaleRow,
    DischargeSummaryDraft,
    IcfDomain,
    IcfQualifier,
    PatientIdentity,
    Procedure,
    ReviewIssue,
    ReviewSeverity,
    SpecialistRole,
)
from mdrk_builder.infrastructure.discharge_summary_template import (
    discharge_summary_template_path,
)
from mdrk_builder.infrastructure.discharge_summary_writer import (
    DischargeSummaryGenerationBlockedError,
    write_discharge_summary_docx,
)


def _draft(tmp_path: Path) -> DischargeSummaryDraft:
    return DischargeSummaryDraft(
        folder=tmp_path,
        identity=PatientIdentity(
            full_name="Тестов Игорь Юрьевич",
            medical_record_number="СКП5906/26",
        ),
        header_text="Сведения о пациенте\nДата поступления: 10.08.2026",
        clinical_diagnosis="Тестовый диагноз",
        medications="",
        transfusions="",
        radiation_exposure="4,2 мЗв",
        recommendations="Тестовая рекомендация",
    )


def test_writer_preserves_branding_page_field_and_dynamic_footer(tmp_path) -> None:
    output = tmp_path / "Выписной эпикриз.docx"

    write_discharge_summary_docx(_draft(tmp_path), output)

    document = Document(output)
    assert document.tables[0].cell(0, 1).text.startswith(
        "ФЕДЕРАЛЬНОЕ МЕДИКО - БИОЛОГИЧЕСКОЕ АГЕНТСТВО"
    )
    assert document.sections[0].footer.paragraphs[0].text == (
        "Тестов Игорь Юрьевич, МКП №СКП5906/26"
    )
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "ВЫПИСНОЙ ЭПИКРИЗ" in text
    assert "Тестовый диагноз" in text
    assert "Лучевая нагрузка: 4,2 мЗв" in text
    assert "Тестовая рекомендация" in text
    with ZipFile(output) as package:
        header_xml = package.read("word/header1.xml").decode("utf-8")
        assert "PAGE" in header_xml


def test_writer_composes_shared_clinical_tables(tmp_path) -> None:
    draft = _draft(tmp_path)
    draft.final_mdrk_source = tmp_path / "final-mdrk.docx"
    draft.icf_domains = (
        IcfDomain(
            "d450",
            "Ходьба",
            SpecialistRole.PHYSICAL_THERAPIST,
            initial=IcfQualifier(2),
            final=IcfQualifier(1),
        ),
        IcfDomain(
            "d640",
            "Ведение домашнего хозяйства",
            SpecialistRole.OCCUPATIONAL_THERAPIST,
            initial=IcfQualifier(3),
        ),
    )
    draft.completed_procedures = (
        Procedure("Тренировка ходьбы", "ФТ", 5, 30, "ежедневно", "ST-150"),
    )
    draft.admission_scale_rows = (
        DischargeScaleRow(
            SpecialistRole.PHYSICAL_THERAPIST,
            "Индекс мобильности Ривермид",
            "6",
        ),
    )
    draft.discharge_scale_rows = (
        DischargeScaleRow(
            SpecialistRole.PHYSICAL_THERAPIST,
            "Индекс мобильности Ривермид",
            "9",
        ),
    )
    output = tmp_path / "Выписной эпикриз.docx"

    write_discharge_summary_docx(draft, output)

    document = Document(output)
    icf_table = next(
        table
        for table in document.tables
        if table.rows[0].cells[0].text == "МКФ категориальный профиль"
    )
    domain_row = next(row for row in icf_table.rows if row.cells[0].text == "d450")
    assert domain_row.cells[11].text == "2"
    assert domain_row.cells[12].text == "1"
    missing_final_row = next(
        row for row in icf_table.rows if row.cells[0].text == "d640"
    )
    assert missing_final_row.cells[11].text == "3"
    assert missing_final_row.cells[12].text == ""

    program = next(
        table
        for table in document.tables
        if table.rows[0].cells[0].text == "Реабилитационные процедуры"
    )
    assert [cell.text for cell in program.rows[1].cells] == [
        "ST-150 Тренировка ходьбы",
        "ФТ",
        "5",
        "30",
        "ежедневно",
    ]
    scale_tables = [
        table
        for table in document.tables
        if [cell.text for cell in table.rows[0].cells]
        == ["Специалист", "Шкала/опросник", "Результат"]
    ]
    assert scale_tables[0].rows[1].cells[2].text == "6"
    assert scale_tables[1].rows[1].cells[2].text == "9"


def test_writer_refuses_to_overwrite_template_or_source(tmp_path) -> None:
    draft = _draft(tmp_path)
    source = tmp_path / "source.docx"
    source.write_bytes(b"immutable")
    draft.source_paths = (source,)

    with pytest.raises(ValueError, match="immutable source"):
        write_discharge_summary_docx(draft, source)
    with pytest.raises(ValueError, match="canonical template"):
        write_discharge_summary_docx(draft, discharge_summary_template_path())


@pytest.mark.parametrize("acknowledged", [False, True])
def test_writer_never_allows_blocking_issue(
    tmp_path,
    acknowledged: bool,
) -> None:
    draft = _draft(tmp_path)
    draft.issues.append(
        ReviewIssue(
            "missing_required_source",
            "Нет обязательного источника",
            ReviewSeverity.BLOCKING,
            acknowledged=acknowledged,
        )
    )

    with pytest.raises(DischargeSummaryGenerationBlockedError):
        write_discharge_summary_docx(draft, tmp_path / "blocked.docx")
