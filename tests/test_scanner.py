from datetime import date, datetime
from pathlib import Path
from zipfile import ZipFile

from docx import Document

from mdrk_builder.application.scanner import (
    ScannedRecord,
    _initial_mdrk_day,
    _latest_clinical_sections,
    _merge_dates,
    _merge_icf,
    _merge_identity,
    _merge_mdrk1_baseline,
    _records_for_selected_medical_record,
    scan_patient_folder,
)
from mdrk_builder.application.source_scan import (
    ScannedDocument,
    SourceReadFailure,
    SourceScanResult,
    scan_source_documents,
)
from mdrk_builder.application.snapshot import build_snapshot
from mdrk_builder.domain import (
    Episode,
    MdrkKind,
    ReviewSeverity,
    ScaleMeasurement,
    SourceDocument,
    SpecialistFinding,
    SpecialistRole,
)
from mdrk_builder.infrastructure.classifier import DocumentClassification
from mdrk_builder.infrastructure.docx_output import save_sanitized_docx_atomically
from mdrk_builder.infrastructure.ooxml_reader import (
    BodyItem,
    ParsedCell,
    ParsedDocument,
    ParsedRow,
    ParsedTable,
)


def _document(path: str, text: str = "", tables: list[ParsedTable] | None = None) -> ParsedDocument:
    paragraphs = [text] if text else []
    parsed_tables = tables or []
    return ParsedDocument(
        source_path=Path(path),
        normalized_path=Path(path),
        paragraphs=paragraphs,
        tables=parsed_tables,
        body_items=[
            *([BodyItem("paragraph", 0)] if paragraphs else []),
            *(BodyItem("table", index) for index in range(len(parsed_tables))),
        ],
    )


def _record(
    path: str,
    text: str,
    *,
    role: SpecialistRole = SpecialistRole.NEUROLOGIST,
    document_type: str = "initial",
    clinical_datetime: datetime | None = None,
    tables: list[ParsedTable] | None = None,
) -> ScannedRecord:
    return ScannedRecord(
        _document(path, text, tables),
        DocumentClassification(role, document_type),
        clinical_datetime,
    )


def _row(values: dict[int, str], logical_cols: int = 15) -> ParsedRow:
    return ParsedRow(
        tuple(ParsedCell(column, 1, value) for column, value in sorted(values.items())),
        logical_cols,
    )


def _icf_table(*rows: ParsedRow) -> ParsedTable:
    return ParsedTable((_row({0: "МКФ", 13: "Ответственный специалист"}), *rows))


def test_source_scan_records_read_failures_and_preserves_caller_normalizer(
    tmp_path,
) -> None:
    good_path = tmp_path / "a-good.docx"
    document = Document()
    document.add_paragraph("Первичный осмотр невролога")
    document.save(good_path)
    broken_path = tmp_path / "b-broken.docx"
    with ZipFile(broken_path, "w"):
        pass

    class TrackingNormalizer:
        def __init__(self) -> None:
            self.normalized: list[Path] = []
            self.closed = False

        def normalize(self, source: Path) -> Path:
            self.normalized.append(source)
            return source

        def close(self) -> None:
            self.closed = True

    normalizer = TrackingNormalizer()

    result = scan_source_documents(tmp_path, normalizer=normalizer)

    assert isinstance(result, SourceScanResult)
    assert result.source_files == (good_path, broken_path)
    assert len(result.documents) == 1
    assert isinstance(result.documents[0], ScannedDocument)
    assert result.documents[0].document.source_path == good_path
    assert len(result.failures) == 1
    assert isinstance(result.failures[0], SourceReadFailure)
    assert result.failures[0].source_path == broken_path
    assert normalizer.normalized == [good_path, broken_path]
    assert not normalizer.closed


def test_source_scan_isolates_invalid_zip_and_xml_failures(tmp_path) -> None:
    good_path = tmp_path / "a-good.docx"
    document = Document()
    document.add_paragraph("Первичный осмотр невролога")
    document.save(good_path)

    bad_zip_path = tmp_path / "b-bad-zip.docx"
    bad_zip_path.write_bytes(b"not a zip package")

    bad_xml_path = tmp_path / "c-bad-xml.docx"
    with ZipFile(bad_xml_path, "w") as package:
        package.writestr("docProps/core.xml", "<coreProperties")
        package.writestr(
            "word/document.xml",
            '<w:document xmlns:w="http://schemas.openxmlformats.org/'
            'wordprocessingml/2006/main"><w:body/></w:document>',
        )

    result = scan_source_documents(tmp_path)

    assert [item.document.source_path for item in result.documents] == [good_path]
    assert [failure.source_path for failure in result.failures] == [
        bad_zip_path,
        bad_xml_path,
    ]


def test_scan_preserves_source_read_and_systematic_failure_issues(tmp_path) -> None:
    source_paths = [tmp_path / f"broken-{index}.docx" for index in range(3)]
    for source_path in source_paths:
        with ZipFile(source_path, "w"):
            pass

    episode = scan_patient_folder(tmp_path)

    read_failures = [
        issue for issue in episode.issues if issue.code == "source_read_failed"
    ]
    assert [issue.source for issue in read_failures] == source_paths
    assert all(issue.severity is ReviewSeverity.WARNING for issue in read_failures)
    assert any(
        issue.code == "systematic_read_failure"
        and issue.severity is ReviewSeverity.BLOCKING
        for issue in episode.issues
    )


def test_scan_patient_folder_accepts_preloaded_source_scan(tmp_path) -> None:
    source_path = tmp_path / "neurologist.docx"
    document = Document()
    for paragraph in (
        "Первичный осмотр невролога",
        "ФИО пациента: АЛЬФА БЕТА ГАММА",
        "Номер ИБ: 123/26",
    ):
        document.add_paragraph(paragraph)
    document.save(source_path)
    source_scan = scan_source_documents(tmp_path)
    source_path.unlink()

    episode = scan_patient_folder(tmp_path, source_scan=source_scan)

    assert episode.identity.full_name == "АЛЬФА БЕТА ГАММА"
    assert episode.identity.medical_record_number == "123/26"


def test_scan_ignores_generated_outputs_and_canonical_discharge_summaries(
    tmp_path,
) -> None:
    primary_path = tmp_path / "primary.docx"
    primary = Document()
    for paragraph in (
        "Первичный осмотр невролога",
        "ФИО пациента: АЛЬФА БЕТА ГАММА",
        "Номер ИБ: 123/26",
        "Дата и время поступления: 10.08.2026 10:00",
        "Клинический диагноз: ИСХОДНЫЙ_ДИАГНОЗ",
    ):
        primary.add_paragraph(paragraph)
    primary.save(primary_path)

    discharge_path = tmp_path / "discharge.docx"
    discharge = Document()
    for paragraph in (
        "Выписной эпикриз",
        "Сведения о пациенте",
        "ФИО пациента: ЧУЖОЙ ПАЦИЕНТ ТЕСТОВЫЙ",
        "Номер медицинской карты пациента №999/26",
        "Клинический диагноз: НЕ_ИСТОЧНИК",
    ):
        discharge.add_paragraph(paragraph)
    discharge.save(discharge_path)

    generated_path = tmp_path / "generated-mdrk.docx"
    generated = Document()
    for paragraph in (
        "Первичный осмотр невролога",
        "ФИО пациента: ЧУЖОЙ ПАЦИЕНТ ТЕСТОВЫЙ",
        "Номер ИБ: 888/26",
        "Клинический диагноз: НЕ_ИСТОЧНИК",
    ):
        generated.add_paragraph(paragraph)
    save_sanitized_docx_atomically(generated, generated_path)

    episode = scan_patient_folder(tmp_path)

    assert episode.identity.full_name == "АЛЬФА БЕТА ГАММА"
    assert episode.identity.medical_record_number == "123/26"
    assert episode.sections.clinical_diagnosis == "ИСХОДНЫЙ_ДИАГНОЗ"
    assert [source.path for source in episode.sources] == [primary_path]
    assert not any(
        issue.code == "identity_conflict_medical_record_number"
        for issue in episode.issues
    )


def test_initial_mdrk_day_rule_for_each_admission_weekday() -> None:
    cases = {
        date(2026, 8, 3): date(2026, 8, 4),   # Monday -> next day
        date(2026, 8, 4): date(2026, 8, 5),
        date(2026, 8, 5): date(2026, 8, 6),
        date(2026, 8, 6): date(2026, 8, 7),
        date(2026, 8, 7): date(2026, 8, 10),  # Friday -> Monday
        date(2026, 8, 8): date(2026, 8, 11),  # weekend -> Tuesday
        date(2026, 8, 9): date(2026, 8, 11),
    }

    assert {value: _initial_mdrk_day(value) for value in cases} == cases


def test_meeting_dates_use_clinic_weekend_rule_and_ignore_discharge_boundary() -> None:
    record = _record(
        "/patient/source.docx",
        "Дата и время поступления: 05.06.2026 12:12\n"
        "Дата и время выписки: 21.06.2026 12:00",
        role=SpecialistRole.PHYSICAL_THERAPIST,
        document_type="follow_up",
        clinical_datetime=datetime(2026, 6, 19, 11, 30),
    )
    episode = Episode(Path("/patient"))

    _merge_dates(episode, [record])

    assert episode.admission_datetime == datetime(2026, 6, 5, 12, 12)
    assert episode.initial_meeting_at == datetime(2026, 6, 8, 8, 0)
    assert episode.final_meeting_at == datetime(2026, 6, 19, 11, 30)
    assert episode.discharge_datetime is None
    assert episode.course_duration_days == 14


def test_explicit_final_mdrk_schedule_overrides_latest_specialist_time() -> None:
    clinical = _record(
        "/patient/ft.docx",
        "Дата и время поступления: 05.06.2026 12:12\n"
        "Дата и время выписки: 21.06.2026 12:00",
        role=SpecialistRole.PHYSICAL_THERAPIST,
        document_type="follow_up",
        clinical_datetime=datetime(2026, 6, 19, 11, 30),
    )
    schedule = _record(
        "/patient/turnaround.docx",
        "",
        role=SpecialistRole.OTHER,
        document_type="administrative",
        tables=[
            ParsedTable(
                (
                    _row({0: "Медицинское вмешательство", 1: "Дата", 2: "Исполнение"}, 3),
                    _row({0: "Консилиум МДРК", 1: "19.06.2026", 2: "19.06.2026 15:30"}, 3),
                )
            )
        ],
    )
    episode = Episode(Path("/patient"))

    _merge_dates(episode, [clinical, schedule])

    assert episode.initial_meeting_at == datetime(2026, 6, 8, 8)
    assert episode.final_meeting_at == datetime(2026, 6, 19, 15, 30)


def test_discharge_date_is_not_materialized_or_used_as_course_boundary() -> None:
    record = _record(
        "/patient/source.docx",
        "Дата и время поступления: 05.06.2026 12:12\n"
        "Дата и время выписки: 21.06.2026 12:00",
        document_type="final",
        clinical_datetime=None,
    )
    episode = Episode(Path("/patient"))

    _merge_dates(episode, [record])

    assert episode.discharge_datetime is None
    assert episode.course_duration_days is None
    assert episode.final_meeting_at is None


def test_scan_meeting_override_is_applied_before_sections_and_icf_materialize(
    tmp_path,
) -> None:
    def write_physician_source(
        path: Path,
        *,
        heading: str,
        examined_at: str,
        diagnosis: str,
        icf_code: str,
    ) -> None:
        document = Document()
        for value in (
            heading,
            f"Дата осмотра: {examined_at}",
            "ФИО пациента: АЛЬФА БЕТА ГАММА",
            "Номер ИБ: 123/26",
            "Дата и время поступления: 05.06.2026 12:00",
            f"Клинический диагноз: {diagnosis}",
        ):
            document.add_paragraph(value)
        table = document.add_table(rows=2, cols=15)
        table.cell(0, 0).text = "МКФ"
        table.cell(0, 13).text = "Ответственный специалист"
        table.cell(1, 0).text = icf_code
        table.cell(1, 1).text = "Тестовый домен"
        table.cell(1, 11).text = "3"
        table.cell(1, 13).text = "невролог"
        document.save(path)

    write_physician_source(
        tmp_path / "невролог первичный.docx",
        heading="Первичный осмотр невролога",
        examined_at="05.06.2026 13:00",
        diagnosis="исходный",
        icf_code="s110",
    )
    write_physician_source(
        tmp_path / "невролог повторный.docx",
        heading="Повторный осмотр невролога",
        examined_at="19.06.2026 16:00",
        diagnosis="слишком поздний",
        icf_code="s999",
    )

    episode = scan_patient_folder(
        tmp_path,
        final_meeting_at=datetime(2026, 6, 19, 15, 30),
    )

    assert episode.final_meeting_at == datetime(2026, 6, 19, 15, 30)
    assert episode.sections.clinical_diagnosis == "исходный"
    assert {domain.code for domain in episode.icf_domains} == {"s110"}


def test_scan_preserves_specialist_name_for_signature_roster(tmp_path) -> None:
    document = Document()
    for value in (
        "Первичная консультация медицинского логопеда",
        "Дата консультации: 05.06.2026 14:00",
        "ФИО пациента: АЛЬФА БЕТА ГАММА",
        "Номер ИБ: 123/26",
        "Медицинский логопед /________/ СОТРУДНИК Л.Г.",
    ):
        document.add_paragraph(value)
    document.save(tmp_path / "логопед первичный.docx")

    episode = scan_patient_folder(tmp_path)

    source = next(item for item in episode.sources if item.role is SpecialistRole.LOGOPEDIST)
    assert source.specialist_name == "СОТРУДНИК Л.Г."


def test_scan_aggregates_goal_and_tasks_from_latest_specialist_plans(tmp_path) -> None:
    def write_source(path: Path, *paragraphs: str) -> None:
        document = Document()
        for paragraph in paragraphs:
            document.add_paragraph(paragraph)
        document.save(path)

    write_source(
        tmp_path / "невролог первичный.docx",
        "Первичный осмотр невролога",
        "Дата осмотра: 14.08.2026 09:46",
        "ФИО пациента: АЛЬФА БЕТА ГАММА",
        "Номер ИБ: 123/26",
        "Дата и время поступления: 14.08.2026 08:40",
        "Клинический диагноз: ДИАГНОЗ_ТЕСТ",
    )
    write_source(
        tmp_path / "логопед первичный.docx",
        "Первичная консультация медицинского логопеда",
        "Дата консультации: 14.08.2026 12:00",
        "ФИО пациента: АЛЬФА БЕТА ГАММА",
        "Номер ИБ: 123/26",
        "Задача на этап МР: через 17 дней пациент ведёт диалог.",
        "Короткосрочная задача реабилитации №1: Улучшить артикуляцию.",
        "Короткосрочная задача реабилитации №2: Увеличить речевой выдох.",
    )
    write_source(
        tmp_path / "фт первичный.docx",
        "Первичный осмотр специалиста по физической реабилитации",
        "Дата осмотра: 14.08.2026 16:10",
        "ФИО пациента: АЛЬФА БЕТА ГАММА",
        "Номер ИБ: 123/26",
        "Реабилитационные задачи на этап МР:",
        "• Развитие силовой выносливости;",
        "• Улучшение функции равновесия;",
        "На основании данных обследования рекомендовано:",
        "• Индивидуальная лечебная физкультура;",
    )

    episode = scan_patient_folder(
        tmp_path,
        initial_meeting_at=datetime(2026, 8, 17, 8),
    )

    assert episode.initial_sections.goal == "через 17 дней пациент ведёт диалог."
    assert episode.initial_sections.tasks.splitlines() == [
        "Улучшить артикуляцию.",
        "Увеличить речевой выдох.",
        "Развитие силовой выносливости;",
        "Улучшение функции равновесия;",
    ]
    assert {
        path.name
        for key, path in episode.initial_field_sources.items()
        if key.startswith("sections.tasks")
    } == {"логопед первичный.docx", "фт первичный.docx"}
    assert not any(issue.field == "initial_sections.tasks" for issue in episode.issues)


def test_mixed_record_numbers_and_admission_dates_are_blocking() -> None:
    records = [
        _record(
            "/patient/stay-1.docx",
            "ФИО пациента: АЛЬФА БЕТА ГАММА\nНомер ИБ: СКП100/26\n"
            "Дата и время поступления: 05.06.2026 12:00",
        ),
        _record(
            "/patient/stay-2.docx",
            "ФИО пациента: АЛЬФА БЕТА ГАММА\nНомер ИБ: СКП200/26\n"
            "Дата и время поступления: 01.07.2026 09:00",
        ),
    ]
    episode = Episode(Path("/patient"))

    _merge_identity(episode, records)
    _merge_dates(episode, records)

    blocking_codes = {
        issue.code for issue in episode.issues if issue.severity is ReviewSeverity.BLOCKING
    }
    assert "identity_conflict_medical_record_number" in blocking_codes
    assert "mixed_hospitalizations_admission_date" in blocking_codes
    admission_conflict = next(
        issue
        for issue in episode.issues
        if issue.code == "mixed_hospitalizations_admission_date"
    )
    assert admission_conflict.source == Path("/patient/stay-2.docx")


def test_equivalent_record_number_formats_do_not_create_false_conflict() -> None:
    records = [
        _record(
            "/patient/a.docx",
            "Номер ИБ: СКП9002/99",
            clinical_datetime=datetime(2026, 8, 3, 9),
        ),
        _record(
            "/patient/b.docx",
            "Номер ИБ: 9002 / 99",
            clinical_datetime=datetime(2026, 8, 3, 8),
        ),
    ]
    episode = Episode(Path("/patient"))

    _merge_identity(episode, records)

    assert not any(
        issue.code == "identity_conflict_medical_record_number"
        for issue in episode.issues
    )


def test_record_number_conflict_points_to_genuinely_different_source() -> None:
    equivalent_path = Path("/patient/equivalent.docx")
    foreign_path = Path("/patient/foreign.docx")
    records = [
        _record(
            "/patient/chosen.docx",
            "Номер ИБ: СКП9002/99",
            clinical_datetime=datetime(2026, 8, 3, 10),
        ),
        _record(
            str(equivalent_path),
            "Номер ИБ: 9002 / 99",
            clinical_datetime=datetime(2026, 8, 3, 9),
        ),
        _record(
            str(foreign_path),
            "Номер ИБ: СКП9003/99",
            clinical_datetime=datetime(2026, 8, 3, 8),
        ),
    ]
    episode = Episode(Path("/patient"))

    _merge_identity(episode, records)

    conflict = next(
        issue
        for issue in episode.issues
        if issue.code == "identity_conflict_medical_record_number"
    )
    assert conflict.source == foreign_path
    assert conflict.source != equivalent_path


def test_scan_overrides_select_episode_before_materialization(tmp_path) -> None:
    def write_source(
        path: Path,
        *,
        record_number: str,
        admission: str,
        examined_at: str,
        diagnosis: str,
    ) -> None:
        document = Document()
        for value in (
            "Первичный осмотр невролога",
            f"Дата осмотра: {examined_at}",
            "ФИО пациента: АЛЬФА БЕТА ГАММА",
            f"Номер ИБ: {record_number}",
            f"Дата и время поступления: {admission}",
            f"Клинический диагноз: {diagnosis}",
        ):
            document.add_paragraph(value)
        document.save(path)

    selected_path = tmp_path / "невролог stay-a.docx"
    excluded_path = tmp_path / "невролог stay-b.docx"
    write_source(
        selected_path,
        record_number="СКП100/26",
        admission="05.06.2026 12:00",
        examined_at="05.06.2026 13:00",
        diagnosis="диагноз выбранного эпизода",
    )
    write_source(
        excluded_path,
        record_number="СКП200/26",
        admission="01.07.2026 09:00",
        examined_at="01.07.2026 10:00",
        diagnosis="диагноз другого эпизода",
    )

    episode = scan_patient_folder(
        tmp_path,
        medical_record_number_override="СКП 100 / 26",
        admission_datetime_override=datetime(2026, 6, 5, 12, 34),
    )

    assert episode.identity.medical_record_number == "СКП 100 / 26"
    assert episode.materialized_medical_record_number == "СКП 100 / 26"
    assert episode.admission_datetime == datetime(2026, 6, 5, 12, 34)
    assert episode.materialized_admission_datetime == datetime(2026, 6, 5, 12, 34)
    assert episode.initial_meeting_at == datetime(2026, 6, 8, 8)
    assert episode.initial_sections.clinical_diagnosis == "диагноз выбранного эпизода"
    assert excluded_path.resolve() in episode.excluded_source_paths
    record_conflict = next(
        issue
        for issue in episode.issues
        if issue.code == "identity_conflict_medical_record_number"
    )
    assert record_conflict.source == excluded_path.resolve()
    assert not any(
        issue.code == "record_number_override_without_source"
        for issue in episode.issues
    )


def test_unknown_record_number_override_is_blocking(tmp_path) -> None:
    document = Document()
    for value in (
        "Первичный осмотр невролога",
        "Дата осмотра: 05.06.2026 13:00",
        "ФИО пациента: АЛЬФА БЕТА ГАММА",
        "Номер ИБ: СКП100/26",
        "Дата и время поступления: 05.06.2026 12:00",
        "Клинический диагноз: тестовый диагноз",
    ):
        document.add_paragraph(value)
    document.save(tmp_path / "невролог.docx")

    episode = scan_patient_folder(
        tmp_path,
        medical_record_number_override="СКП999/26",
    )

    assert any(
        issue.code == "record_number_override_without_source"
        and issue.severity is ReviewSeverity.BLOCKING
        for issue in episode.issues
    )


def test_different_medical_record_stays_visible_but_cannot_supply_episode_data() -> None:
    physician_path = "/patient/rehab-primary.docx"
    cardiology_path = "/patient/cardiology-discharge.docx"
    records = [
        _record(
            physician_path,
            "ФИО пациента: АЛЬФА БЕТА ГАММА\n"
            "Номер ИБ: СКП9002/99\n"
            "Дата и время поступления: 01.08.2026 09:11\n"
            "Клинический диагноз: реабилитационный диагноз\n"
            "Лабораторные исследования: анализы эпизода",
            role=SpecialistRole.NEUROLOGIST,
            document_type="initial",
            clinical_datetime=datetime(2026, 8, 3, 8, 39),
        ),
        _record(
            cardiology_path,
            "ФИО пациента: АЛЬФА БЕТА ГАММА\n"
            "Номер ИБ: СКП9003/99\n"
            "Дата и время поступления: 28.07.2026 14:17\n"
            "Дата и время выписки: 01.08.2026 12:00\n"
            "Клинический диагноз: чужой кардиологический диагноз\n"
            "Лабораторные исседования: чужие анализы",
            role=SpecialistRole.OTHER,
            document_type="final",
            clinical_datetime=datetime(2026, 8, 1, 12),
        ),
    ]
    episode = Episode(Path("/patient"))

    _merge_identity(episode, records)
    active = _records_for_selected_medical_record(episode, records)
    _merge_dates(episode, active)
    _latest_clinical_sections(episode, active)

    assert episode.identity.medical_record_number == "СКП9002/99"
    assert [item.document.source_path for item in active] == [Path(physician_path)]
    assert episode.excluded_source_paths == {Path(cardiology_path)}
    assert episode.admission_datetime == datetime(2026, 8, 1, 9, 11)
    assert episode.discharge_datetime is None
    assert episode.initial_meeting_at == datetime(2026, 8, 4, 8)
    assert episode.initial_sections.clinical_diagnosis == "реабилитационный диагноз"
    assert episode.initial_sections.laboratory_results == "анализы эпизода"
    assert any(issue.code == "source_medical_record_mismatch" for issue in episode.issues)
    assert not any(
        issue.code == "physician_source_after_meeting" for issue in episode.issues
    )
    assert episode.initial_sections.rehabilitation_potential == "средний"
    assert episode.sections.rehabilitation_potential == "средний"


def test_future_physician_previews_missing_fields_even_when_older_physician_is_empty() -> None:
    older_empty = _record(
        "/patient/older.docx",
        "Клинический диагноз:",
        clinical_datetime=datetime(2026, 8, 1, 12),
    )
    future_primary = _record(
        "/patient/primary.docx",
        "Клинический диагноз: диагноз для проверки\n"
        "Анамнез заболевания: анамнез для проверки",
        clinical_datetime=datetime(2026, 8, 3, 8, 39),
    )
    episode = Episode(Path("/patient"))
    episode.initial_meeting_at = datetime(2026, 8, 2, 8)
    episode.final_meeting_at = datetime(2026, 8, 4, 8)

    _latest_clinical_sections(episode, [older_empty, future_primary])

    assert episode.initial_sections.clinical_diagnosis == "диагноз для проверки"
    assert episode.initial_sections.disease_history == "анамнез для проверки"
    assert any(
        issue.code == "physician_source_after_meeting"
        and issue.severity is ReviewSeverity.BLOCKING
        and issue.field == "initial_meeting_at"
        for issue in episode.issues
    )


def test_other_consilium_cannot_supply_latest_physician_sections() -> None:
    physician = _record(
        "/patient/initial.docx",
        "Клинический диагноз: верный диагноз",
        clinical_datetime=datetime(2026, 6, 5, 13),
    )
    peg = _record(
        "/patient/консилиум гастростома.docx",
        "Клинический диагноз: не брать из PEG-протокола",
        role=SpecialistRole.OTHER,
        document_type="other_consilium",
        clinical_datetime=datetime(2026, 6, 20, 13),
    )
    episode = Episode(Path("/patient"))

    _latest_clinical_sections(episode, [physician, peg])

    assert episode.sections.clinical_diagnosis == "верный диагноз"


def test_clinical_sections_are_selected_per_field_as_of_each_meeting() -> None:
    initial = _record(
        "/patient/initial.docx",
        "Клинический диагноз: исходный\nАнамнез заболевания: первичный анамнез",
        clinical_datetime=datetime(2026, 6, 5, 13),
    )
    follow_up = _record(
        "/patient/follow-up.docx",
        "Клинический диагноз: уточнённый",
        document_type="follow_up",
        clinical_datetime=datetime(2026, 6, 15, 9),
    )
    after_final_meeting = _record(
        "/patient/discharge.docx",
        "Клинический диагноз: слишком поздний",
        document_type="final",
        clinical_datetime=datetime(2026, 6, 21, 12),
    )
    episode = Episode(Path("/patient"))
    episode.initial_meeting_at = datetime(2026, 6, 6, 8)
    episode.final_meeting_at = datetime(2026, 6, 19, 11, 30)

    _latest_clinical_sections(episode, [initial, follow_up, after_final_meeting])

    assert episode.initial_sections.clinical_diagnosis == "исходный"
    assert episode.initial_sections.disease_history == "первичный анамнез"
    assert episode.sections.clinical_diagnosis == (
        "исходный\n"
        "(Дополнение от 15.06.2026 09:00): уточнённый"
    )
    assert episode.sections.disease_history == "первичный анамнез"
    assert episode.initial_field_sources["sections.clinical_diagnosis"] == Path(
        "/patient/initial.docx"
    )
    assert episode.field_sources["sections.clinical_diagnosis"] == Path(
        "/patient/follow-up.docx"
    )
    assert not any(
        issue.code == "physician_source_after_meeting"
        and issue.source == Path("/patient/discharge.docx")
        for issue in episode.issues
    )


def test_mdrk2_sections_ignore_empty_diaries_and_deduplicate_copied_diagnostics() -> None:
    initial = _record(
        "/patient/initial.docx",
        "Анамнез заболевания: Острое начало.\n"
        "Лабораторные исследования: Анализ крови Hb 120 г/л. "
        "Анализ мочи без патологии.",
        clinical_datetime=datetime(2026, 8, 3, 8, 20),
    )
    empty_diary = _record(
        "/patient/diary-empty.docx",
        "Анамнез заболевания: Без изменений.\n"
        "Лабораторные исследования: По ИПМР.",
        document_type="follow_up",
        clinical_datetime=datetime(2026, 8, 5, 9),
    )
    copied_with_addition = _record(
        "/patient/diary-update.docx",
        "Анамнез заболевания: Острое начало. Самочувствие улучшилось.\n"
        "Лабораторные исследования: Анализ крови - Hb 120 г/л. "
        "Анализ мочи без патологии. СРБ 12 мг/л.",
        document_type="follow_up",
        clinical_datetime=datetime(2026, 8, 7, 10, 15),
    )
    episode = Episode(Path("/patient"))
    episode.initial_meeting_at = datetime(2026, 8, 4, 8)
    episode.final_meeting_at = datetime(2026, 8, 10, 8)

    _latest_clinical_sections(episode, [initial, empty_diary, copied_with_addition])

    assert episode.initial_sections.disease_history == "Острое начало."
    assert episode.sections.disease_history == (
        "Острое начало.\n"
        "(Дополнение от 07.08.2026 10:15): Самочувствие улучшилось."
    )
    assert episode.sections.laboratory_results == (
        "Анализ крови Hb 120 г/л. Анализ мочи без патологии.\n"
        "(Дополнение от 07.08.2026 10:15): СРБ 12 мг/л."
    )


def test_specialist_copy_cannot_fill_physician_owned_clinical_fields() -> None:
    physician = _record(
        "/patient/physician.docx",
        "Клинический диагноз:",
        clinical_datetime=datetime(2026, 6, 5, 13),
    )
    physical_therapist = _record(
        "/patient/ft.docx",
        "Клинический диагноз: копия диагноза\n"
        "Медикаментозная терапия: копия лечения",
        role=SpecialistRole.PHYSICAL_THERAPIST,
        clinical_datetime=datetime(2026, 6, 5, 14),
    )
    episode = Episode(Path("/patient"))
    episode.initial_meeting_at = datetime(2026, 6, 6, 8)
    episode.final_meeting_at = datetime(2026, 6, 19, 11)

    _latest_clinical_sections(episode, [physician, physical_therapist])

    assert episode.initial_sections.clinical_diagnosis == ""
    assert episode.initial_sections.medication == ""
    assert episode.sections.clinical_diagnosis == ""
    assert episode.sections.medication == ""


def test_icf_ownership_filters_physician_copy_and_keeps_dynamics_and_pf() -> None:
    physician_initial = _record(
        "/patient/physician-initial.docx",
        "",
        clinical_datetime=datetime(2026, 6, 5, 13),
        tables=[
            _icf_table(
                _row({0: "b2351", 1: "Равновесие", 11: "2", 13: "ФТ"}),
                _row({0: "s110", 1: "Структура головного мозга", 11: "3", 13: "МРТ"}),
                ParsedRow(
                    (ParsedCell(0, 1, "Pf"), ParsedCell(1, 14, "Мужчина 58 лет")),
                    15,
                ),
            )
        ],
    )
    physician_final = _record(
        "/patient/physician-final.docx",
        "",
        document_type="final",
        clinical_datetime=datetime(2026, 6, 21, 12),
        tables=[
            _icf_table(
                _row({0: "b2351", 1: "Равновесие", 11: "2", 12: "1", 13: "ФТ"}),
                _row({0: "s110", 1: "Структура головного мозга", 11: "3", 12: "4", 13: "МРТ"}),
                _row({0: "s999", 1: "Слишком поздний домен", 11: "4", 12: "4", 13: "МРТ"}),
                ParsedRow(
                    (ParsedCell(0, 1, "Pf"), ParsedCell(1, 14, "Мужчина 58 лет")),
                    15,
                ),
            )
        ],
    )
    physician_follow_up = _record(
        "/patient/physician-follow-up.docx",
        "",
        document_type="follow_up",
        clinical_datetime=datetime(2026, 6, 15, 9),
        tables=[
            _icf_table(
                _row({0: "s110", 1: "Структура головного мозга", 11: "3", 12: "3", 13: "МРТ"})
            )
        ],
    )
    psychologist = _record(
        "/patient/pathopsychologist.docx",
        "",
        role=SpecialistRole.PATHOPSYCHOLOGIST,
        document_type="consultation",
        clinical_datetime=datetime(2026, 6, 5, 15),
        tables=[
            _icf_table(
                ParsedRow(
                    (ParsedCell(0, 1, "Pf"), ParsedCell(1, 14, "Мужской, 58 лет")),
                    15,
                ),
                _row({0: "e310", 1: "Семья и близкие родственники", 11: "0", 13: "патопсихолог"}),
            )
        ],
    )
    ft_initial = _record(
        "/patient/ft-initial.docx",
        "",
        role=SpecialistRole.PHYSICAL_THERAPIST,
        clinical_datetime=datetime(2026, 6, 5, 14, 30),
        tables=[_icf_table(_row({0: "b2351", 1: "Равновесие", 11: "2"}))],
    )
    ft_final = _record(
        "/patient/ft-final.docx",
        "",
        role=SpecialistRole.PHYSICAL_THERAPIST,
        document_type="follow_up",
        clinical_datetime=datetime(2026, 6, 19, 11, 30),
        tables=[_icf_table(_row({0: "b2351", 1: "Равновесие", 11: "2", 12: "1"}))],
    )
    episode = Episode(Path("/patient"))
    episode.initial_meeting_at = datetime(2026, 6, 6, 8)
    episode.final_meeting_at = datetime(2026, 6, 19, 15, 30)
    episode.initial_sections.medication = "Лозартан 25 мг"
    episode.sections.medication = "Лозартан 25 мг"
    episode.initial_field_sources["sections.medication"] = Path(
        "/patient/physician-initial.docx"
    )
    episode.field_sources["sections.medication"] = Path(
        "/patient/physician-initial.docx"
    )

    _merge_icf(
        episode,
        [
            physician_initial,
            physician_follow_up,
            physician_final,
            psychologist,
            ft_initial,
            ft_final,
        ],
    )

    b_domains = [item for item in episode.icf_domains if item.code == "b2351"]
    assert len(b_domains) == 1
    assert b_domains[0].specialist is SpecialistRole.PHYSICAL_THERAPIST
    assert b_domains[0].initial and b_domains[0].initial.value == 2
    assert b_domains[0].final and b_domains[0].final.value == 1
    assert b_domains[0].dynamic_marker == "+"

    structure = next(item for item in episode.icf_domains if item.code == "s110")
    assert structure.specialist is SpecialistRole.OTHER
    assert structure.initial and structure.final
    assert structure.initial.value == structure.final.value == 3
    assert structure.final_source == Path("/patient/physician-follow-up.docx")
    assert not any(item.code == "s999" for item in episode.icf_domains)

    personal = next(item for item in episode.icf_domains if item.code == "Pf")
    assert personal.description == "Мужской, 58 лет"
    assert personal.specialist is SpecialistRole.OTHER
    assert personal.initial is None and personal.final is None
    assert any(issue.code == "personal_factor_conflict" for issue in episode.issues)

    zero_domain = next(item for item in episode.icf_domains if item.code == "e310")
    assert zero_domain.initial and zero_domain.initial.value == 0

    medication = next(item for item in episode.icf_domains if item.code == "e1101")
    assert medication.initial and medication.initial.display() == "4+"
    assert medication.final and medication.final.display() == "4+"
    assert medication.initial_source == Path("/patient/physician-initial.docx")
    assert medication.final_source == Path("/patient/physician-initial.docx")


def test_pf_uses_physician_source_when_no_psychologist_source_exists() -> None:
    physician = _record(
        "/patient/physician-initial.docx",
        "",
        clinical_datetime=datetime(2026, 8, 3, 8, 39),
        tables=[
                _icf_table(
                    _row({0: "b730", 1: "Мышечная сила", 11: "3"}),
                    _row({0: "e1101", 1: "Лекарственные препараты", 11: "4+"}),
                    _row({0: "e310", 1: "Семья и ближайшие родственники", 11: "4+"}),
                _row({0: "d550", 1: "Приём пищи", 11: "2"}),
                ParsedRow(
                    (
                        ParsedCell(0, 1, "Pf"),
                        ParsedCell(1, 14, "Мужчина 65 лет, не работает"),
                    ),
                    15,
                )
            )
        ],
    )
    episode = Episode(Path("/patient"))
    episode.initial_meeting_at = datetime(2026, 8, 4, 8)
    episode.final_meeting_at = datetime(2026, 8, 14, 8)

    _merge_icf(episode, [physician])

    personal = next(item for item in episode.icf_domains if item.code == "Pf")
    assert personal.description == "Мужчина 65 лет, не работает"
    assert personal.specialist is SpecialistRole.OTHER
    assert personal.initial is None and personal.final is None
    assert personal.initial_source == Path("/patient/physician-initial.docx")
    assert personal.final_source is None
    # Copied FT/FZT domains from a neurologist's primary SHRM table are excluded.
    assert {item.code for item in episode.icf_domains} == {"b730", "e1101", "Pf"}


def test_pf_is_merged_once_across_roles_and_prefers_authoritative_source() -> None:
    ft = _record(
        "/patient/ft.docx",
        "",
        role=SpecialistRole.PHYSICAL_THERAPIST,
        clinical_datetime=datetime(2026, 6, 22, 10),
        tables=[
            _icf_table(
                ParsedRow(
                    (ParsedCell(0, 1, "Pf"), ParsedCell(1, 15, "женщина, возраст 78 года.")),
                    16,
                )
            )
        ],
    )
    frm = _record(
        "/patient/frm.docx",
        "",
        role=SpecialistRole.FRM,
        clinical_datetime=datetime(2026, 6, 22, 9),
        tables=[
            _icf_table(
                ParsedRow(
                    (ParsedCell(0, 1, "Pf"), ParsedCell(1, 14, "женщина, 78 лет")),
                    15,
                )
            )
        ],
    )
    episode = Episode(Path("/patient"))
    episode.initial_meeting_at = datetime(2026, 6, 23, 8)
    episode.final_meeting_at = datetime(2026, 7, 6, 8)

    _merge_icf(episode, [ft, frm])

    personal = [item for item in episode.icf_domains if item.code.casefold() == "pf"]
    assert len(personal) == 1
    assert personal[0].description == "женщина, 78 лет"
    assert personal[0].specialist is SpecialistRole.OTHER
    assert personal[0].initial_source == Path("/patient/frm.docx")
    assert any(issue.code == "personal_factor_conflict" for issue in episode.issues)


def test_pf_first_seen_after_mdrk1_does_not_leak_into_initial_snapshot() -> None:
    later = _record(
        "/patient/ergo-follow-up.docx",
        "",
        role=SpecialistRole.OCCUPATIONAL_THERAPIST,
        clinical_datetime=datetime(2026, 6, 10, 12),
        tables=[
            _icf_table(
                ParsedRow(
                    (ParsedCell(0, 1, "Pf"), ParsedCell(1, 14, "Женщина 70 лет")),
                    15,
                )
            )
        ],
    )
    episode = Episode(Path("/patient"))
    episode.initial_meeting_at = datetime(2026, 6, 6, 8)
    episode.final_meeting_at = datetime(2026, 6, 19, 8)

    _merge_icf(episode, [later])

    personal = next(item for item in episode.icf_domains if item.code == "Pf")
    assert personal.initial_source is None
    assert personal.final_source == Path("/patient/ergo-follow-up.docx")
    assert not any(
        item.code == "Pf" for item in build_snapshot(episode, MdrkKind.INITIAL).icf_domains
    )
    assert any(item.code == "Pf" for item in build_snapshot(episode, MdrkKind.FINAL).icf_domains)


def test_follow_up_only_domain_uses_first_and_last_points_but_stays_out_of_mdrk1() -> None:
    follow_up = _record(
        "/patient/ft-follow-up.docx",
        "",
        role=SpecialistRole.PHYSICAL_THERAPIST,
        document_type="follow_up",
        clinical_datetime=datetime(2026, 6, 11, 13),
        tables=[
            _icf_table(
                _row({0: "b999", 1: "Новый домен", 11: "2", 12: "1"})
            )
        ],
    )
    episode = Episode(Path("/patient"))
    episode.initial_meeting_at = datetime(2026, 6, 6, 8)
    episode.final_meeting_at = datetime(2026, 6, 19, 15, 30)

    _merge_icf(episode, [follow_up])

    domain = next(item for item in episode.icf_domains if item.code == "b999")
    assert domain.initial and domain.initial.value == 2
    assert domain.final and domain.final.value == 1
    assert domain.initial_source == Path("/patient/ft-follow-up.docx")
    assert domain.final_source == Path("/patient/ft-follow-up.docx")
    assert domain.initial_measured_at == datetime(2026, 6, 11, 13)
    assert domain.final_measured_at == datetime(2026, 6, 11, 13)
    assert not any(
        item.code == "b999"
        for item in build_snapshot(episode, MdrkKind.INITIAL).icf_domains
    )


def test_icf_fuzzy_merges_near_duplicate_description_and_keeps_first_last_only() -> None:
    role = SpecialistRole.PHYSICAL_THERAPIST
    records = [
        _record(
            "/patient/ft-initial.docx",
            "",
            role=role,
            document_type="initial",
            clinical_datetime=datetime(2026, 6, 5, 10),
            tables=[
                _icf_table(
                    _row({0: "d450", 1: "Ходьба пациента", 11: "3"})
                )
            ],
        ),
        _record(
            "/patient/ft-diary-1.docx",
            "",
            role=role,
            document_type="follow_up",
            clinical_datetime=datetime(2026, 6, 10, 10),
            tables=[_icf_table(_row({0: "d450", 1: "Ходьба пациета", 11: "2"}))],
        ),
        _record(
            "/patient/ft-diary-2.docx",
            "",
            role=role,
            document_type="follow_up",
            clinical_datetime=datetime(2026, 6, 18, 10),
            tables=[_icf_table(_row({0: "d450", 1: "Ходьба пациента.", 11: "1"}))],
        ),
        _record(
            "/patient/ft-too-late.docx",
            "",
            role=role,
            document_type="final",
            clinical_datetime=datetime(2026, 6, 21, 10),
            tables=[_icf_table(_row({0: "d450", 1: "Ходьба пациента", 11: "0"}))],
        ),
    ]
    episode = Episode(Path("/patient"))
    episode.initial_meeting_at = datetime(2026, 6, 6, 8)
    episode.final_meeting_at = datetime(2026, 6, 20, 11)

    _merge_icf(episode, records)

    domains = [item for item in episode.icf_domains if item.code == "d450"]
    assert len(domains) == 1
    domain = domains[0]
    assert domain.description == "Ходьба пациента"
    assert domain.initial and domain.initial.value == 3
    assert domain.final and domain.final.value == 1
    assert domain.initial_source == Path("/patient/ft-initial.docx")
    assert domain.final_source == Path("/patient/ft-diary-2.docx")
    assert domain.initial_measured_at == datetime(2026, 6, 5, 10)
    assert domain.final_measured_at == datetime(2026, 6, 18, 10)


def test_single_new_icf_point_becomes_final_snapshot_baseline_without_fake_repeat() -> None:
    role = SpecialistRole.OCCUPATIONAL_THERAPIST
    follow_up = _record(
        "/patient/ot-new-domain.docx",
        "",
        role=role,
        document_type="follow_up",
        clinical_datetime=datetime(2026, 6, 12, 14),
        tables=[_icf_table(_row({0: "d640", 1: "Домашняя работа", 11: "2"}))],
    )
    episode = Episode(Path("/patient"))
    episode.initial_meeting_at = datetime(2026, 6, 6, 8)
    episode.final_meeting_at = datetime(2026, 6, 20, 11)

    _merge_icf(episode, [follow_up])

    domain = next(item for item in episode.icf_domains if item.code == "d640")
    assert domain.initial and domain.initial.value == 2
    assert domain.final is None
    assert build_snapshot(episode, MdrkKind.INITIAL).icf_domains == ()
    assert [item.code for item in build_snapshot(episode, MdrkKind.FINAL).icf_domains] == [
        "d640"
    ]


def test_icf_copies_at_same_timestamp_are_one_point_despite_wording_variant() -> None:
    role = SpecialistRole.PHYSICAL_THERAPIST
    occurred_at = datetime(2026, 6, 5, 10)
    records = [
        _record(
            "/patient/ft-copy-a.docx",
            "",
            role=role,
            clinical_datetime=occurred_at,
            tables=[_icf_table(_row({0: "d450", 1: "Ходьба пациента", 11: "3"}))],
        ),
        _record(
            "/patient/ft-copy-b.docx",
            "",
            role=role,
            clinical_datetime=occurred_at,
            tables=[
                _icf_table(
                    _row({0: "d450", 1: "Ходьба пациета", 11: "3"})
                )
            ],
        ),
    ]
    episode = Episode(Path("/patient"))
    episode.initial_meeting_at = datetime(2026, 6, 6, 8)
    episode.final_meeting_at = datetime(2026, 6, 20, 11)

    _merge_icf(episode, records)

    domain = next(item for item in episode.icf_domains if item.code == "d450")
    assert domain.initial and domain.initial.value == 3
    assert domain.final is None
    assert domain.initial_source == Path("/patient/ft-copy-a.docx")


def test_same_icf_code_for_left_and_right_limb_remains_two_domains() -> None:
    role = SpecialistRole.PHYSICAL_THERAPIST
    record = _record(
        "/patient/ft-limbs.docx",
        "",
        role=role,
        clinical_datetime=datetime(2026, 6, 5, 10),
        tables=[
            _icf_table(
                _row({0: "b7301", 1: "Сила мышц левой руки", 11: "3"}),
                _row({0: "b7301", 1: "Сила мышц правой руки", 11: "1"}),
            )
        ],
    )
    episode = Episode(Path("/patient"))
    episode.initial_meeting_at = datetime(2026, 6, 6, 8)
    episode.final_meeting_at = datetime(2026, 6, 20, 11)

    _merge_icf(episode, [record])

    domains = [item for item in episode.icf_domains if item.code == "b7301"]
    assert len(domains) == 2
    assert {item.description for item in domains} == {
        "Сила мышц левой руки",
        "Сила мышц правой руки",
    }


def test_mdrk1_fills_only_missing_baseline_icf_and_scales() -> None:
    physical_role = SpecialistRole.PHYSICAL_THERAPIST
    primary_at = datetime(2026, 6, 7, 10)
    mdrk1_at = datetime(2026, 6, 8, 8)
    repeat_at = datetime(2026, 6, 18, 10)
    primary = _record(
        "/patient/ft-primary.docx",
        "",
        role=physical_role,
        clinical_datetime=primary_at,
        tables=[_icf_table(_row({0: "b730", 1: "Сила мышц", 11: "2", 13: "ФТ"}))],
    )
    repeat = _record(
        "/patient/ft-repeat.docx",
        "",
        role=physical_role,
        document_type="follow_up",
        clinical_datetime=repeat_at,
        tables=[_icf_table(_row({0: "d450", 1: "Ходьба", 11: "1", 13: "ФТ"}))],
    )
    mdrk_icf = _icf_table(
        _row({0: "b730", 1: "Сила мышц", 11: "4", 13: "ФТ"}),
        _row({0: "d450", 1: "Ходьба", 11: "3", 13: "ФТ"}),
        _row({0: "e310", 1: "Поддержка семьи", 11: "4+"}),
        _row({0: "Pf", 1: "Мужчина трудоспособного возраста"}),
    )
    mdrk_scales = ParsedTable(
        (
            _row({0: "Шкала/опросник", 1: "Исходно 08.06.2026 07:45"}, 2),
            _row({0: "Шкала баланса Берга", 1: "99"}, 2),
            _row({0: "Шкала Тинетти", 1: "10"}, 2),
        )
    )
    paragraphs = [
        "Консилиум мультидисциплинарной реабилитационной команды",
        '"08" июня 2026 г. время: 08 час. 00 мин.',
        "1. Клинический диагноз: текст МДРК не должен использоваться",
        "Результат осмотра специалиста по физической реабилитации:",
    ]
    mdrk_document = ParsedDocument(
        source_path=Path("/patient/mdrk-1.docx"),
        normalized_path=Path("/patient/mdrk-1.docx"),
        paragraphs=paragraphs,
        tables=[mdrk_icf, mdrk_scales],
        body_items=[
            BodyItem("paragraph", 0),
            BodyItem("paragraph", 1),
            BodyItem("paragraph", 2),
            BodyItem("table", 0),
            BodyItem("paragraph", 3),
            BodyItem("table", 1),
        ],
    )
    mdrk_record = ScannedRecord(
        mdrk_document,
        DocumentClassification(
            SpecialistRole.OTHER,
            "mdrk",
            is_mdrk=True,
            confidence=1.0,
            mdrk_kind=MdrkKind.INITIAL,
        ),
        mdrk1_at,
    )
    episode = Episode(Path("/patient"))
    episode.initial_meeting_at = mdrk1_at
    episode.final_meeting_at = datetime(2026, 6, 19, 15, 30)
    episode.sources.extend(
        (
            SourceDocument(primary.document.source_path, physical_role, primary_at),
            SourceDocument(repeat.document.source_path, physical_role, repeat_at),
            SourceDocument(mdrk_document.source_path, SpecialistRole.OTHER, mdrk1_at, "mdrk_initial"),
        )
    )
    episode.findings.extend(
        (
            SpecialistFinding(
                physical_role,
                source_datetime=primary_at,
                source=primary.document.source_path,
                scales=[
                    ScaleMeasurement(
                        "Шкала баланса Берга",
                        "30",
                        primary_at,
                        physical_role,
                        primary.document.source_path,
                    )
                ],
            ),
            SpecialistFinding(
                physical_role,
                source_datetime=repeat_at,
                source=repeat.document.source_path,
                scales=[
                    ScaleMeasurement(
                        "Шкала Тинетти",
                        "20",
                        repeat_at,
                        physical_role,
                        repeat.document.source_path,
                    )
                ],
            ),
        )
    )

    _merge_icf(episode, [primary, repeat])
    _merge_mdrk1_baseline(episode, [mdrk_record])

    by_code = {item.code: item for item in episode.icf_domains}
    assert by_code["b730"].initial and by_code["b730"].initial.value == 2
    assert by_code["b730"].initial_source == primary.document.source_path
    assert by_code["d450"].initial and by_code["d450"].initial.value == 3
    assert by_code["d450"].final and by_code["d450"].final.value == 1
    assert by_code["d450"].initial_source == mdrk_document.source_path
    assert by_code["d450"].final_source == repeat.document.source_path
    assert by_code["e310"].initial and by_code["e310"].initial.display() == "4+"
    assert by_code["Pf"].initial_source == mdrk_document.source_path
    assert episode.sections.clinical_diagnosis == ""

    rows = {row.name: row for row in build_snapshot(episode, MdrkKind.FINAL).scale_rows}
    assert rows["Шкала баланса Берга"].initial.value == "30"
    assert rows["Шкала баланса Берга"].current is None
    assert rows["Шкала Тинетти"].initial.value == "10"
    assert rows["Шкала Тинетти"].current.value == "20"
    assert any(issue.code == "mdrk1_baseline_fallback" for issue in episode.issues)


def test_mdrk1_fallback_rejects_document_from_non_initial_meeting_day() -> None:
    episode = Episode(Path("/patient"))
    episode.initial_meeting_at = datetime(2026, 6, 8, 8)
    episode.final_meeting_at = datetime(2026, 6, 19, 15, 30)
    wrong_day = _record(
        "/patient/not-really-mdrk1.docx",
        "",
        role=SpecialistRole.OTHER,
        document_type="mdrk",
        clinical_datetime=datetime(2026, 6, 18, 8),
        tables=[_icf_table(_row({0: "d450", 1: "Ходьба", 11: "3"}))],
    )

    _merge_mdrk1_baseline(episode, [wrong_day])

    assert episode.icf_domains == []
    assert episode.findings == []
