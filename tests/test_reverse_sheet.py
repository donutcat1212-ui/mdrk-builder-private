from datetime import date, datetime
from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document
from docx.oxml.ns import qn

from mdrk_builder.application.reverse_sheet import scan_reverse_sheet
from mdrk_builder.domain import (
    PatientIdentity,
    ReverseSheetDraft,
    ReverseSheetRow,
    ReviewSeverity,
)
from mdrk_builder.infrastructure.reverse_sheet_writer import write_reverse_sheet_docx
from mdrk_builder.ui import reverse_sheet_dialog as reverse_dialog_module
from mdrk_builder.ui.reverse_sheet_dialog import confirm_incomplete_reverse_dates


def _write_docx(path: Path, *paragraphs: str) -> None:
    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)


def test_reverse_sheet_preserves_source_read_failure_issue(tmp_path) -> None:
    broken_path = tmp_path / "broken.docx"
    with ZipFile(broken_path, "w"):
        pass

    draft = scan_reverse_sheet(tmp_path)

    read_failure = next(
        issue for issue in draft.issues if issue.code == "reverse_source_read_failed"
    )
    assert read_failure.source == broken_path
    assert read_failure.severity is ReviewSeverity.WARNING
    assert broken_path.name in read_failure.message
    assert any(
        issue.code == "reverse_primary_neurologist_missing"
        for issue in draft.issues
    )


def test_reverse_sheet_uses_strict_header_dates_and_consultation_chronology(tmp_path) -> None:
    _write_docx(
        tmp_path / "невролог первичный.docx",
        "Первичный осмотр невролога",
        "Дата осмотра: 05.08.2026 09:00",
        "ФИО пациента: ПАЦИЕНТ ТЕСТОВЫЙ ПРИМЕР",
        "Дата рождения: 01.02.1970",
        "Номер ИБ: 123/26",
        "Дата и время поступления: 05.08.2026 08:15",
        "Консультация специалиста по физической реабилитации назначена на 06.08.2026",
        "АЛЬФА А.А., лечащий врач, врач-невролог /________/ ОМЕГА О.О., заведующий отделением",
    )
    _write_docx(
        tmp_path / "фт консультация.docx",
        "Первичный осмотр специалиста по физической реабилитации",
        "Дата консультации: 07.08.2026 10:30",
        "Специалист: АЛЬФА А.А.",
    )
    _write_docx(
        tmp_path / "логопед повторный.docx",
        "Повторная консультация медицинского логопеда",
        "Дата консультации: 08.08.2026 11:45",
        "Медицинский логопед: БЕТА Б.Б.",
    )
    _write_docx(
        tmp_path / "невролог дневник.docx",
        "Повторная консультация невролога",
        "Дата консультации: 08.08.2026 12:15",
        "Врач-невролог: ДЕЛЬТА Д.Д.",
    )
    _write_docx(
        tmp_path / "мдрк.docx",
        "Консилиум мультидисциплинарной реабилитационной команды",
        "10.08.2026 12:00",
        "Лечащий врач: ГАММА Г.Г.",
    )
    _write_docx(
        tmp_path / "узи.docx",
        "УЗИ органов брюшной полости",
        "Дата исследования: 06.08.2026 15:00",
    )

    draft = scan_reverse_sheet(tmp_path)

    assert draft.identity.full_name == "ПАЦИЕНТ ТЕСТОВЫЙ ПРИМЕР"
    assert draft.identity.medical_record_number == "123/26"
    assert draft.admission_datetime == datetime(2026, 8, 5, 8, 15)
    assert [row.intervention for row in draft.rows] == [
        "Консультация специалиста по физической реабилитации",
        "Консультация медицинского логопеда",
        "Консилиум МДРК",
    ]
    assert draft.rows[0].appointment_date == date(2026, 8, 5)
    assert draft.rows[1].appointment_date == date(2026, 8, 8)
    assert draft.rows[2].appointment_date == date(2026, 8, 10)
    assert all("невролог" not in row.intervention.casefold() for row in draft.rows)
    assert draft.rows[2].performer == "АЛЬФА А.А."
    assert [row.performed_at for row in draft.rows] == sorted(
        row.performed_at for row in draft.rows if row.performed_at is not None
    )


def test_reverse_sheet_writer_preserves_reference_grid_and_source(tmp_path) -> None:
    source = tmp_path / "source.docx"
    _write_docx(source, "SOURCE SENTINEL")
    original = source.read_bytes()
    draft = ReverseSheetDraft(
        folder=tmp_path,
        identity=PatientIdentity(
            "ПАЦИЕНТ ТЕСТОВЫЙ ПРИМЕР",
            date(1970, 2, 1),
            "",
            "123/26",
        ),
        admission_datetime=datetime(2026, 8, 5, 8, 15),
        header_source=source,
        rows=[
            ReverseSheetRow(
                "Консультация невролога",
                None,
                datetime(2026, 8, 6, 10, 30),
                "АЛЬФА А.А.",
                source,
            )
        ],
    )
    output = tmp_path / "result.docx"

    write_reverse_sheet_docx(draft, output)

    assert source.read_bytes() == original
    document = Document(output)
    assert len(document.tables) == 1
    table = document.tables[0]
    assert len(table.rows) == 27
    widths = [int(node.get(qn("w:w"))) for node in table._tbl.tblGrid]
    assert widths == [2739, 1420, 1826, 81, 1746, 1826]
    assert "оборотная сторона раздела" in table.cell(0, 4).text
    assert "ПАЦИЕНТ ТЕСТОВЫЙ ПРИМЕР" in table.cell(1, 0).text
    assert table.cell(3, 0).text == "Консультация невролога"
    assert table.cell(3, 1).text == ""
    assert table.cell(3, 3).text == "06.08.2026 10:30"
    assert table._tbl.tblPr.first_child_found_in("w:tblpPr") is not None
    assert table.cell(0, 4)._tc.tcPr.first_child_found_in("w:tcBorders") is None
    assert table.cell(2, 0)._tc.tcPr.first_child_found_in("w:tcBorders") is not None


def test_existing_reverse_sheet_supplies_only_mdrk_dates_as_reviewable_evidence(tmp_path) -> None:
    _write_docx(
        tmp_path / "невролог первичный.docx",
        "Первичный осмотр невролога",
        "Дата осмотра: 05.08.2026 09:00",
        "ФИО пациента: ПАЦИЕНТ ТЕСТОВЫЙ ПРИМЕР",
        "Номер ИБ: 123/26",
        "Дата и время поступления: 05.08.2026 08:15",
    )
    _write_docx(
        tmp_path / "мдрк.docx",
        "Консилиум мультидисциплинарной реабилитационной команды",
        "Лечащий врач: АЛЬФА А.А.",
    )
    existing = Document()
    existing.add_paragraph('Оборотная сторона раздела "Лист назначений и их выполнение"')
    table = existing.add_table(rows=1, cols=6)
    values = (
        "Консилиум МДРК",
        "06.08.2026",
        "",
        "07.08.2026 08:00",
        "",
        "АЛЬФА А.А.",
    )
    for index, value in enumerate(values):
        table.cell(0, index).text = value
    existing.save(tmp_path / "оборотная сторона раздела.docx")

    draft = scan_reverse_sheet(tmp_path)

    mdrk = next(row for row in draft.rows if row.intervention == "Консилиум МДРК")
    assert mdrk.appointment_date == date(2026, 8, 7)
    assert mdrk.performed_at == datetime(2026, 8, 7, 8)
    assert any(
        issue.code == "reverse_mdrk_date_carried_from_existing_sheet"
        for issue in draft.issues
    )


def test_source_mdrk_time_wins_and_planned_date_matches_execution_date(tmp_path) -> None:
    _write_docx(
        tmp_path / "невролог первичный.docx",
        "Первичный осмотр невролога",
        "Дата осмотра: 05.08.2026 09:00",
        "ФИО пациента: ПАЦИЕНТ ТЕСТОВЫЙ ПРИМЕР",
        "Номер ИБ: 123/26",
    )
    _write_docx(
        tmp_path / "мдрк.docx",
        "Консилиум мультидисциплинарной реабилитационной команды",
        "10.08.2026 08-46",
    )
    existing = Document()
    table = existing.add_table(rows=1, cols=6)
    for index, value in enumerate(
        ("Консилиум МДРК", "09.08.2026", "", "09.08.2026 07:00", "", "АЛЬФА А.А.")
    ):
        table.cell(0, index).text = value
    existing.save(tmp_path / "оборотная сторона раздела.docx")

    draft = scan_reverse_sheet(tmp_path)
    mdrk = next(row for row in draft.rows if row.intervention == "Консилиум МДРК")

    assert mdrk.appointment_date == date(2026, 8, 10)
    assert mdrk.performed_at == datetime(2026, 8, 10, 8, 46)
    assert not any(
        issue.code == "reverse_mdrk_date_carried_from_existing_sheet"
        for issue in draft.issues
    )


def test_initial_examination_uses_primary_neurologist_date_as_appointment(tmp_path) -> None:
    _write_docx(
        tmp_path / "невролог первичный.docx",
        "Первичный осмотр невролога",
        "Дата осмотра: 05.08.2026 09:00",
        "ФИО пациента: ПАЦИЕНТ ТЕСТОВЫЙ ПРИМЕР",
        "Номер ИБ: 123/26",
    )
    _write_docx(
        tmp_path / "патопсихолог.docx",
        "Первичное обследование медицинского психолога (патопсихолога)",
        "Дата приема: 06.08.2026 08:00",
        "Медицинский психолог/патопсихолог: БЕТА Б.Б.",
    )

    draft = scan_reverse_sheet(tmp_path)
    row = next(item for item in draft.rows if "патопсихолога" in item.intervention)

    assert row.appointment_date == date(2026, 8, 5)
    assert row.performed_at == datetime(2026, 8, 6, 8)


@pytest.mark.parametrize(
    ("noise", "signature", "expected"),
    [
        ("ФАМИЛИЯ А.Д.", "АЛЬФА А.Д.", "АЛЬФА А.Д."),
        ("Оценка работоспособности по А.Ю. ШКАЛОВОЙ", "БЕТА В.Г.", "БЕТА В.Г."),
    ],
)
def test_neuropsych_performer_comes_from_professional_signature(
    tmp_path,
    noise: str,
    signature: str,
    expected: str,
) -> None:
    _write_docx(
        tmp_path / "невролог первичный.docx",
        "Первичный осмотр невролога",
        "Дата осмотра: 05.08.2026 09:00",
        "ФИО пациента: ПАЦИЕНТ ТЕСТОВЫЙ ПРИМЕР",
        "Номер ИБ: 123/26",
    )
    _write_docx(
        tmp_path / "нейропсихолог.docx",
        "Повторная консультация медицинского психолога (нейропсихолога)",
        "Дата консультации: 08.08.2026 11:45",
        noise,
        f"Медицинский психолог (нейропсихолог)____________________________{signature}",
    )

    draft = scan_reverse_sheet(tmp_path)
    row = next(item for item in draft.rows if "нейропсихолога" in item.intervention)

    assert row.performer == expected


def test_writer_places_manual_fill_rows_between_appointment_blocks(tmp_path) -> None:
    draft = ReverseSheetDraft(
        folder=tmp_path,
        rows=[
            ReverseSheetRow("Первая", date(2026, 8, 6), datetime(2026, 8, 6, 9)),
            ReverseSheetRow("Вторая", date(2026, 8, 6), datetime(2026, 8, 6, 10)),
            ReverseSheetRow("Третья", date(2026, 8, 15), datetime(2026, 8, 15, 9)),
        ],
    )
    output = write_reverse_sheet_docx(draft, tmp_path / "blocks.docx")

    table = Document(output).tables[0]
    assert table.cell(3, 0).text == "Первая"
    assert table.cell(4, 0).text == "Вторая"
    assert table.cell(5, 0).text == ""
    assert table.cell(6, 0).text == ""
    assert table.cell(7, 0).text == "Третья"


def test_missing_reverse_dates_can_be_accepted_globally(monkeypatch) -> None:
    calls: list[str] = []

    def accept(_title: str, message: str, **_kwargs) -> bool:
        calls.append(message)
        return True

    monkeypatch.setattr(reverse_dialog_module.messagebox, "askyesno", accept)
    rows = [ReverseSheetRow("КОНСУЛЬТАЦИЯ_ТЕСТ", None, None)]

    assert confirm_incomplete_reverse_dates(object(), rows)
    assert len(calls) == 1
    assert "дата назначения" in calls[0]
    assert "дата исполнения" in calls[0]


def test_complete_reverse_dates_need_no_confirmation(monkeypatch) -> None:
    monkeypatch.setattr(
        reverse_dialog_module.messagebox,
        "askyesno",
        lambda *_args, **_kwargs: (_ for _ in ()).throw(AssertionError("unexpected dialog")),
    )

    assert confirm_incomplete_reverse_dates(
        object(),
        [ReverseSheetRow("КОНСУЛЬТАЦИЯ_ТЕСТ", date(2026, 8, 5), datetime(2026, 8, 5, 9))],
    )
