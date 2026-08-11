from __future__ import annotations

from datetime import date, datetime
from pathlib import Path

import pytest
from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn

from mdrk_builder.domain import (
    Episode,
    IcfDomain,
    IcfQualifier,
    MdrkKind,
    Procedure,
    ReviewIssue,
    ReviewSeverity,
    ScaleMeasurement,
    SourceDocument,
    SpecialistFinding,
    SpecialistRole,
)
from mdrk_builder.infrastructure.docx_layout import TABLE_WIDTH_DXA
from mdrk_builder.infrastructure.docx_writer import (
    DocumentGenerationBlockedError,
    SignatoryRow,
    write_mdrk_docx,
)


def _representative_episode(folder: Path) -> Episode:
    initial_at = datetime(2026, 6, 5, 16)
    final_at = datetime(2026, 6, 19, 13)
    episode = Episode(folder=folder)
    episode.identity.full_name = "ПАЦИЕНТ_ТЕСТ"
    episode.identity.birth_date = date(1968, 2, 20)
    episode.identity.sex = "муж"
    episode.identity.medical_record_number = "СКП0000/00"
    episode.department = "Отделение медицинской реабилитации для пациентов с нарушением функции ЦНС №2"
    episode.admission_datetime = datetime(2026, 6, 4, 12, 15)
    episode.discharge_datetime = datetime(2026, 6, 20, 10)
    episode.initial_meeting_at = initial_at
    episode.final_meeting_at = final_at
    episode.course_duration_days = 16
    initial_sections = episode.initial_sections
    initial_sections.clinical_diagnosis = "ДИАГНОЗ_ИСХОДНЫЙ"
    initial_sections.disease_history = "АНАМНЕЗ_ЗАБОЛЕВАНИЯ_ИСХОДНЫЙ"
    initial_sections.life_history = "АНАМНЕЗ_ЖИЗНИ_ИСХОДНЫЙ"
    initial_sections.laboratory_results = "ЛАБОРАТОРНЫЙ_МАРКЕР_ИСХОДНЫЙ"
    initial_sections.instrumental_results = "ИНСТРУМЕНТАЛЬНЫЙ_МАРКЕР_ИСХОДНЫЙ"
    initial_sections.rehabilitation_potential = "средний"
    initial_sections.limitations = "МАРКЕР_ОГРАНИЧЕНИЙ"
    initial_sections.risks = "МАРКЕР_РИСКОВ"
    initial_sections.movement_regimen = "палатный"
    initial_sections.diet = "ОВД"
    initial_sections.medication = "ТЕРАПИЯ_СТРОКА_1\nТЕРАПИЯ_СТРОКА_2"
    initial_sections.goal = "ЦЕЛЬ_ИСХОДНАЯ"
    initial_sections.tasks = "ЗАДАЧА_1\nЗАДАЧА_2"

    final_sections = episode.sections
    final_sections.clinical_diagnosis = "ДИАГНОЗ_ИТОГОВЫЙ"
    final_sections.disease_history = "АНАМНЕЗ_ЗАБОЛЕВАНИЯ_ИТОГОВЫЙ"
    final_sections.life_history = "АНАМНЕЗ_ЖИЗНИ_ИТОГОВЫЙ"
    final_sections.laboratory_results = "ЛАБОРАТОРНЫЙ_МАРКЕР_ИТОГОВЫЙ"
    final_sections.instrumental_results = "ИНСТРУМЕНТАЛЬНЫЙ_МАРКЕР_ИТОГОВЫЙ"
    final_sections.rehabilitation_potential = "высокий"
    final_sections.limitations = "МАРКЕР_ОГРАНИЧЕНИЙ"
    final_sections.risks = "МАРКЕР_РИСКОВ_ИТОГОВЫЙ"
    final_sections.movement_regimen = "свободный"
    final_sections.diet = "ОВД"
    final_sections.medication = "ТЕРАПИЯ_ИТОГОВАЯ"
    final_sections.goal = "ЦЕЛЬ_ИТОГОВАЯ_НЕ_ИСПОЛЬЗУЕТСЯ"
    final_sections.tasks = "ЗАДАЧИ_ИТОГОВЫЕ_НЕ_ИСПОЛЬЗУЮТСЯ"
    episode.sources.append(
        SourceDocument(
            folder / "невролог.docx",
            role=SpecialistRole.NEUROLOGIST,
            clinical_datetime=initial_at,
        )
    )

    initial_scale = ScaleMeasurement(
        "Шкала Тинетти",
        "14 баллов",
        initial_at,
        SpecialistRole.PHYSICAL_THERAPIST,
    )
    final_scale = ScaleMeasurement(
        "Шкала Тинетти",
        "24 балла",
        final_at,
        SpecialistRole.PHYSICAL_THERAPIST,
    )
    episode.findings.extend(
        (
            SpecialistFinding(
                SpecialistRole.NEUROLOGIST,
                "ЗАКЛЮЧЕНИЕ_НЕВРОЛОГА",
                initial_at,
                scales=[
                    ScaleMeasurement(
                        "СКФ",
                        "63,73",
                        initial_at,
                        SpecialistRole.NEUROLOGIST,
                    )
                ],
            ),
            SpecialistFinding(
                SpecialistRole.PHYSICAL_THERAPIST,
                "ЗАКЛЮЧЕНИЕ_ФТ_ИСХОДНОЕ",
                initial_at,
                scales=[initial_scale],
            ),
            SpecialistFinding(
                SpecialistRole.PHYSICAL_THERAPIST,
                "ЗАКЛЮЧЕНИЕ_ФТ_ИТОГОВОЕ",
                final_at,
                scales=[final_scale],
            ),
            SpecialistFinding(
                SpecialistRole.NEUROPSYCHOLOGIST,
                "ЗАКЛЮЧЕНИЕ_НЕЙРОПСИХОЛОГА",
                initial_at,
            ),
        )
    )
    episode.icf_domains.extend(
        (
            IcfDomain(
                "b730",
                "Функции мышечной силы",
                SpecialistRole.PHYSICAL_THERAPIST,
                initial=IcfQualifier(2),
                final=IcfQualifier(1),
            ),
            IcfDomain(
                "e310",
                "Семья и ближайшие родственники",
                SpecialistRole.OTHER,
                initial=IcfQualifier(4, facilitator=True),
                final=IcfQualifier(3, facilitator=True),
            ),
            IcfDomain(
                "d450",
                "Ходьба",
                SpecialistRole.PHYSICAL_THERAPIST,
                initial=IcfQualifier(1),
                final=IcfQualifier(2),
            ),
            IcfDomain(
                "Pf",
                "ПЕРСОНАЛЬНЫЙ_ФАКТОР_ТЕСТ",
                SpecialistRole.OTHER,
            ),
        )
    )
    episode.procedures.extend(
        (
            Procedure("Тренировка ходьбы", "ФТ", 5, 30, "ежедневно", "ST-150"),
            Procedure("Нейропсихологическая коррекция", "Нейропсихолог", 6, 45, "5 раз/нед."),
        )
    )
    return episode


def _find_table(document, first_cell_text: str):
    return next(
        table
        for table in document.tables
        if table.rows and table.rows[0].cells[0].text == first_cell_text
    )


def _find_domain_row(table, code: str):
    return next(row for row in table.rows if row.cells[0].text == code)


def _assert_signature_table_page_separator(document) -> None:
    procedures = _find_table(document, "Реабилитационные процедуры")
    signatures = _find_table(document, "Специалист МДРК")
    body_children = list(document._element.body)
    procedure_index = body_children.index(procedures._tbl)
    signature_index = body_children.index(signatures._tbl)

    assert signature_index == procedure_index + 3
    separators = body_children[procedure_index + 1 : signature_index]
    assert all(separator.tag == qn("w:p") for separator in separators)
    assert all(
        separator.find("./w:pPr/w:pageBreakBefore", separator.nsmap) is None
        for separator in separators
    )
    assert separators[-1].find("./w:pPr/w:keepNext", separators[-1].nsmap) is not None

    for row in signatures.rows[:-1]:
        assert all(
            cell.paragraphs[0]._p.get_or_add_pPr().find(qn("w:keepNext")) is not None
            for cell in row.cells
        )
    for cell in signatures.rows[-1].cells:
        marker = cell.paragraphs[0]._p.get_or_add_pPr().find(qn("w:keepNext"))
        assert marker is None or marker.get(qn("w:val")) in {"0", "false", "off"}


def _cell_fill(cell) -> str | None:
    shading = cell._tc.get_or_add_tcPr().find(qn("w:shd"))
    return None if shading is None else shading.get(qn("w:fill"))


def _assert_compact_header_cell(cell) -> None:
    properties = cell._tc.get_or_add_tcPr()
    assert properties.find(qn("w:noWrap")) is not None
    margins = properties.find(qn("w:tcMar"))
    assert margins is not None
    assert margins.find(qn("w:left")).get(qn("w:w")) == "0"
    assert margins.find(qn("w:right")).get(qn("w:w")) == "0"


def _assert_table_geometry_and_pagination(document) -> None:
    table_property_order = (
        "w:tblStyle",
        "w:tblpPr",
        "w:tblOverlap",
        "w:bidiVisual",
        "w:tblStyleRowBandSize",
        "w:tblStyleColBandSize",
        "w:tblW",
        "w:jc",
        "w:tblCellSpacing",
        "w:tblInd",
        "w:tblBorders",
        "w:shd",
        "w:tblLayout",
        "w:tblCellMar",
        "w:tblLook",
        "w:tblCaption",
        "w:tblDescription",
        "w:tblPrChange",
    )
    cell_property_order = (
        "w:cnfStyle",
        "w:tcW",
        "w:gridSpan",
        "w:hMerge",
        "w:vMerge",
        "w:tcBorders",
        "w:shd",
        "w:noWrap",
        "w:tcMar",
        "w:textDirection",
        "w:tcFitText",
        "w:vAlign",
        "w:hideMark",
        "w:headers",
        "w:cellIns",
        "w:cellDel",
        "w:cellMerge",
        "w:tcPrChange",
    )
    for table in document.tables:
        properties = table._tbl.tblPr
        _assert_ooxml_child_order(properties, table_property_order)
        width = properties.find(qn("w:tblW"))
        indent = properties.find(qn("w:tblInd"))
        layout = properties.find(qn("w:tblLayout"))
        assert width is not None
        assert width.get(qn("w:type")) == "dxa"
        assert int(width.get(qn("w:w"))) == TABLE_WIDTH_DXA
        assert indent is not None and int(indent.get(qn("w:w"))) == 0
        assert layout is not None and layout.get(qn("w:type")) == "fixed"

        grid_widths = [
            int(column.get(qn("w:w")))
            for column in table._tbl.tblGrid.findall(qn("w:gridCol"))
        ]
        assert sum(grid_widths) == TABLE_WIDTH_DXA
        for row in table.rows:
            row_properties = row._tr.get_or_add_trPr()
            assert row_properties.find(qn("w:cantSplit")) is not None
            assert row_properties.find(qn("w:trHeight")) is None
            for cell in row.cells:
                _assert_ooxml_child_order(cell._tc.get_or_add_tcPr(), cell_property_order)

        header_count = 2 if table.rows[0].cells[0].text == "МКФ категориальный профиль" else 1
        for row in table.rows[:header_count]:
            assert row._tr.get_or_add_trPr().find(qn("w:tblHeader")) is not None


def _assert_ooxml_child_order(element, expected_order: tuple[str, ...]) -> None:
    rank = {qn(tag): index for index, tag in enumerate(expected_order)}
    observed = [rank[child.tag] for child in element if child.tag in rank]
    assert observed == sorted(observed)


def test_writer_renders_initial_and_final_from_one_template(tmp_path) -> None:
    episode = _representative_episode(tmp_path)
    signatories = (
        SignatoryRow(SpecialistRole.NEUROLOGIST, "СОТРУДНИК_1"),
        SignatoryRow(SpecialistRole.PHYSICAL_THERAPIST, "СОТРУДНИК_2"),
    )
    initial_path = write_mdrk_docx(
        episode,
        MdrkKind.INITIAL,
        tmp_path / "МДРК1.docx",
        signatories=signatories,
    )
    final_path = write_mdrk_docx(
        episode,
        MdrkKind.FINAL,
        tmp_path / "МДРК2.docx",
    )

    initial = Document(initial_path)
    final = Document(final_path)
    initial_text = "\n".join(paragraph.text for paragraph in initial.paragraphs)
    final_text = "\n".join(paragraph.text for paragraph in final.paragraphs)
    for number in range(1, 13):
        assert f"{number}. " in initial_text
        assert f"{number}. " in final_text
    assert "ЦЕЛЬ_ИСХОДНАЯ" in initial_text
    assert "ДИАГНОЗ_ИСХОДНЫЙ" in initial_text
    assert "ДИАГНОЗ_ИТОГОВЫЙ" not in initial_text
    assert "Дата и время выписки:" not in initial_text
    assert "ДИАГНОЗ_ИТОГОВЫЙ" in final_text
    assert "ДИАГНОЗ_ИСХОДНЫЙ" not in final_text
    assert "Дата и время выписки:" not in final_text
    assert "Достигнута в полном объёме" in final_text
    assert "Выполнены в полном объёме" in final_text

    assert [paragraph.text for paragraph in initial.paragraphs[:11]] == [
        "Консилиум мультидисциплинарной реабилитационной команды в составе "
        "заведующего отделением, врача фрм, специалиста по физической реабилитации, "
        "медицинского психолога/нейропсихолога, медицинского логопеда и специалиста "
        "по эргореабилитации",
        '"05" июня 2026 г. время: 16 час. 00 мин.',
        "",
        "Номер ИБ: СКП0000/00",
        "Отделение медицинской реабилитации для пациентов с нарушением функции ЦНС №2.",
        "",
        "ФИО пациента: ПАЦИЕНТ_ТЕСТ",
        "Дата рождения: «20» февраля 1968г. (58 лет)",
        "Пол: муж",
        "",
        "1. Клинический диагноз",
    ]
    assert initial.paragraphs[3].runs[-1].bold
    assert initial.paragraphs[3].runs[-1].underline
    assert initial.paragraphs[6].runs[-1].bold
    assert initial.paragraphs[6].runs[-1].underline

    initial_mcf = _find_table(initial, "МКФ категориальный профиль")
    final_mcf = _find_table(final, "МКФ категориальный профиль")
    assert len(initial_mcf.columns) == len(final_mcf.columns) == 15
    assert initial_mcf.rows[1].cells[14].text == ""
    assert final_mcf.rows[1].cells[14].text == "+/-"

    initial_b730 = _find_domain_row(initial_mcf, "b730")
    final_b730 = _find_domain_row(final_mcf, "b730")
    assert (initial_b730.cells[11].text, initial_b730.cells[12].text, initial_b730.cells[14].text) == (
        "2",
        "",
        "",
    )
    assert (final_b730.cells[11].text, final_b730.cells[12].text, final_b730.cells[14].text) == (
        "2",
        "1",
        "+",
    )
    assert [_cell_fill(initial_b730.cells[index]) for index in range(2, 11)] == [
        None,
        None,
        None,
        None,
        "BFBFBF",
        "BFBFBF",
        "BFBFBF",
        None,
        None,
    ]
    initial_e310 = _find_domain_row(initial_mcf, "e310")
    assert initial_e310.cells[13].text == ""
    assert [_cell_fill(initial_e310.cells[index]) for index in range(2, 11)] == [
        "BFBFBF",
        "BFBFBF",
        "BFBFBF",
        "BFBFBF",
        "BFBFBF",
        None,
        None,
        None,
        None,
    ]
    final_d450 = _find_domain_row(final_mcf, "d450")
    assert final_d450.cells[14].text == "-"
    personal_factor = _find_domain_row(final_mcf, "Pf")
    physical_cells = personal_factor._tr.findall(qn("w:tc"))
    assert len(physical_cells) == 2
    grid_span = physical_cells[1].find("./w:tcPr/w:gridSpan", physical_cells[1].nsmap)
    assert grid_span is not None and grid_span.get(qn("w:val")) == "14"
    assert personal_factor.cells[1].text == "ПЕРСОНАЛЬНЫЙ_ФАКТОР_ТЕСТ"
    for cell in final_mcf.rows[2].cells[2:13]:
        _assert_compact_header_cell(cell)

    problem_headers = [
        row for row in initial_mcf.rows if row.cells[6].text == "Проблемы"
    ]
    assert len(problem_headers) == 2
    for header_row in problem_headers:
        assert [header_row.cells[index].text for index in range(2, 6)] == ["", "", "", ""]
        assert header_row.cells[6].text == "Проблемы"
    environment_index, environment_header = next(
        (index, row)
        for index, row in enumerate(initial_mcf.rows)
        if row.cells[2].text == "Позитивные\nфакторы"
    )
    assert environment_header.cells[7].text == "Барьеры"
    assert environment_header.cells[14].text == ""
    environment_scale = initial_mcf.rows[environment_index + 1]
    assert [environment_scale.cells[index].text for index in range(2, 11)] == [
        "4+", "3+", "2+", "1+", "0", "1", "2", "3", "4"
    ]
    final_environment_header = next(
        row
        for row in final_mcf.rows
        if row.cells[2].text == "Позитивные\nфакторы"
    )
    assert final_environment_header.cells[14].text == "+/-"

    initial_scale = _find_table(initial, "Шкала/опросник")
    final_scale = next(
        table
        for table in final.tables
        if any(row.cells[0].text == "Шкала Тинетти" for row in table.rows)
    )
    assert len(initial_scale.columns) == 2
    assert len(final_scale.columns) == 3
    assert "05.06.2026 16:00" in initial_scale.rows[0].cells[1].text
    assert "05.06.2026 16:00" in final_scale.rows[0].cells[1].text
    assert "19.06.2026 13:00" in final_scale.rows[0].cells[2].text
    assert initial_scale.rows[1].cells[1].text == "14 баллов"
    assert final_scale.rows[1].cells[1].text == "14 баллов"
    assert final_scale.rows[1].cells[2].text == "24 балла"

    physician_table = _find_table(initial, "Дата и время\nрасчета шкалы")
    assert len(physician_table.columns) == 3
    assert [cell.text for cell in physician_table.rows[0].cells] == [
        "Дата и время\nрасчета шкалы",
        "Шкала/опросник",
        "Результат расчета",
    ]
    physician_heading = next(
        paragraph
        for paragraph in initial.paragraphs
        if paragraph.text.startswith("Результат осмотра врача физической")
    )
    assert physician_heading.text == (
        "Результат осмотра врача физической и реабилитационной медицины "
        "(05 июня 2026 16:00):"
    )
    assert not physician_heading.runs[0].bold

    procedures = _find_table(initial, "Реабилитационные процедуры")
    _assert_compact_header_cell(procedures.rows[0].cells[2])
    assert procedures.rows[1].cells[0].text == "ST-150 Тренировка ходьбы"
    assert procedures.rows[1].cells[2].text == "5"
    assert procedures.rows[2].cells[2].text == "6"
    signatures = _find_table(initial, "Специалист МДРК")
    final_signatures = _find_table(final, "Специалист МДРК")
    expected_roles = [
        "Врач ФРМ",
        "Специалист по физической реабилитации",
        "Медицинский психолог/нейропсихолог",
        "Медицинский психолог/патопсихолог",
        "Медицинский логопед",
        "Специалист по эргореабилитации",
        "Консультанты",
        "Заведующий отделением",
    ]
    assert [row.cells[0].text for row in signatures.rows[1:]] == expected_roles
    assert [row.cells[0].text for row in final_signatures.rows[1:]] == expected_roles
    assert signatures.rows[1].cells[1].text == "СОТРУДНИК_1"
    assert signatures.rows[2].cells[1].text == "СОТРУДНИК_2"
    assert signatures.rows[1].cells[2].text == ""
    assert all(row.cells[1].text == "" for row in final_signatures.rows[1:])
    _assert_signature_table_page_separator(initial)
    _assert_signature_table_page_separator(final)
    conclusion_labels = [
        paragraph
        for paragraph in initial.paragraphs
        if paragraph.text.startswith("Заключение: ")
    ]
    assert conclusion_labels

    plan_heading = next(
        paragraph for paragraph in initial.paragraphs if paragraph.text == "Режим и питание:"
    )
    medication_heading = next(
        paragraph
        for paragraph in initial.paragraphs
        if paragraph.text == "Медикаментозное лечение:"
    )
    procedures_heading = next(
        paragraph
        for paragraph in initial.paragraphs
        if paragraph.text == "Реабилитационные мероприятия:"
    )
    assert all(
        paragraph.runs[0].bold and paragraph.runs[0].underline
        for paragraph in (plan_heading, medication_heading, procedures_heading)
    )
    assert "ТЕРАПИЯ_СТРОКА_1" in [p.text for p in initial.paragraphs]
    assert "ТЕРАПИЯ_СТРОКА_2" in [p.text for p in initial.paragraphs]

    _assert_table_geometry_and_pagination(initial)
    _assert_table_geometry_and_pagination(final)


def test_writer_recomputes_blockers_after_manual_fill(tmp_path) -> None:
    episode = _representative_episode(tmp_path)
    episode.issues.append(
        ReviewIssue(
            "required_full_name",
            "Старая ошибка до ручного заполнения",
            ReviewSeverity.BLOCKING,
            field="identity.full_name",
        )
    )

    assert write_mdrk_docx(
        episode, MdrkKind.INITIAL, tmp_path / "manual-override.docx"
    ).is_file()


def test_writer_blocks_on_current_required_issue(tmp_path) -> None:
    episode = _representative_episode(tmp_path)
    episode.sources.clear()

    with pytest.raises(DocumentGenerationBlockedError) as error:
        write_mdrk_docx(episode, MdrkKind.FINAL, tmp_path / "blocked.docx")

    assert any(issue.code == "required_physician_source" for issue in error.value.issues)


def test_writer_never_overwrites_an_episode_source(tmp_path) -> None:
    episode = _representative_episode(tmp_path)
    source_path = episode.sources[0].path

    with pytest.raises(ValueError, match="immutable source"):
        write_mdrk_docx(episode, MdrkKind.INITIAL, source_path)

    assert not source_path.exists()


def test_writer_cleans_atomic_temporary_file_on_save_failure(
    tmp_path, monkeypatch
) -> None:
    episode = _representative_episode(tmp_path)
    output = tmp_path / "failed.docx"

    def fail_save(_document, _path) -> None:
        raise RuntimeError("simulated save failure")

    monkeypatch.setattr(DocxDocument, "save", fail_save)
    with pytest.raises(RuntimeError, match="simulated save failure"):
        write_mdrk_docx(episode, MdrkKind.INITIAL, output)

    assert not output.exists()
    assert not list(tmp_path.glob(".failed-*.docx"))
