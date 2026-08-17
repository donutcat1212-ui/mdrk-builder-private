from zipfile import ZipFile

from docx import Document

from mdrk_builder.infrastructure.discharge_summary_template import (
    FOOTER_PATIENT_PLACEHOLDER,
    FOOTER_RECORD_PLACEHOLDER,
    discharge_summary_template_path,
)


def test_runtime_template_keeps_layout_without_sample_patient_data() -> None:
    path = discharge_summary_template_path()
    document = Document(path)

    assert len(document.tables) == 1
    assert "ФЕДЕРАЛЬНЫЙ ЦЕНТР МОЗГА И НЕЙРОТЕХНОЛОГИЙ" in document.tables[0].cell(
        0, 1
    ).text
    footer = document.sections[0].footer.paragraphs[0].text
    assert FOOTER_PATIENT_PLACEHOLDER in footer
    assert FOOTER_RECORD_PLACEHOLDER in footer
    assert document.core_properties.author == "MDRK Builder"
    assert document.core_properties.last_modified_by == "MDRK Builder"

    with ZipFile(path) as package:
        xml = "\n".join(
            package.read(name).decode("utf-8", errors="ignore")
            for name in package.namelist()
            if name.endswith((".xml", ".rels"))
        )
        assert "PAGE" in package.read("word/header1.xml").decode("utf-8")
    for forbidden in (
        "Гирфанова",
        "1008/23",
        "Лиджиева",
        "/Users/",
    ):
        assert forbidden not in xml
